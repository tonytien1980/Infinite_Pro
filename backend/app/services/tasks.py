from __future__ import annotations

from dataclasses import dataclass
import hashlib
import logging
import re
from io import BytesIO

from fastapi import HTTPException
from sqlalchemy import select
from sqlalchemy.orm import Session, object_session, selectinload

from app.domain import models, schemas
from app.domain.enums import (
    CapabilityArchetype,
    DeliverableClass,
    EngagementContinuityMode,
    ExternalDataStrategy,
    InputEntryMode,
    PresenceState,
    TaskStatus,
    WritebackDepth,
)
from app.extensions.registry import ExtensionRegistry
from app.extensions.resolver import AgentResolver, PackResolver, resolve_runtime_agent_binding
from app.extensions.schemas import AgentResolverInput, AgentSpec, PackResolverInput, PackSpec, PackType
from app.services.artifact_storage import load_artifact_content
from app.services.content_revisions import (
    CONTENT_REVISION_SOURCE_MANUAL_EDIT,
    CONTENT_REVISION_SOURCE_ROLLBACK,
    CONTENT_REVISION_SOURCE_RUNTIME_BACKFILL,
    create_deliverable_content_revision,
    create_matter_content_revision,
    ensure_deliverable_content_revisions,
    ensure_matter_content_revisions,
)
from app.services.deliverable_records import (
    DELIVERABLE_ARTIFACT_KIND_EXPORT,
    DELIVERABLE_ARTIFACT_KIND_RELEASE,
    DELIVERABLE_EVENT_CONTENT_ROLLED_BACK,
    DELIVERABLE_EVENT_CONTENT_UPDATED,
    DELIVERABLE_EVENT_EXPORTED,
    DELIVERABLE_EVENT_NOTE_ADDED,
    DELIVERABLE_EVENT_PUBLISHED,
    DELIVERABLE_EVENT_SOURCE_CONTENT_ROLLBACK,
    DELIVERABLE_EVENT_SOURCE_CONTENT_UPDATE,
    DELIVERABLE_EVENT_SOURCE_EXPORT,
    DELIVERABLE_EVENT_SOURCE_METADATA_UPDATE,
    DELIVERABLE_EVENT_SOURCE_PUBLISH,
    DELIVERABLE_EVENT_STATUS_CHANGED,
    DELIVERABLE_EVENT_VERSION_TAG_UPDATED,
    create_deliverable_artifact_record,
    create_deliverable_publish_record,
    default_deliverable_version_tag,
    ensure_deliverable_release_records,
    ensure_deliverable_version_events,
    label_for_deliverable_status,
    record_deliverable_version_event,
)
from app.services.source_materials import (
    build_source_material_summary,
    ensure_source_chain_participation_links,
    infer_artifact_type,
    resolve_participation_canonical_key,
)

logger = logging.getLogger(__name__)
EXTERNAL_DATA_STRATEGY_CONSTRAINT_TYPE = "external_data_strategy"
UNSPECIFIED_LABEL = "未指定"
DEFAULT_CLIENT_NAME = "尚未明確標示客戶"
DEFAULT_DOMAIN_LENS = "綜合"
EXTERNAL_DATA_STRATEGY_LABELS = {
    ExternalDataStrategy.STRICT: "僅使用我提供的資料（嚴格模式）",
    ExternalDataStrategy.SUPPLEMENTAL: "視需要補充外部資料",
    ExternalDataStrategy.LATEST: "優先使用最新外部資料（研究模式）",
}
PACK_OVERRIDE_CONSTRAINT_TYPE = "pack_override"
AGENT_OVERRIDE_CONSTRAINT_TYPE = "agent_override"
LEGACY_TASK_SLICE_IDENTITY_SCOPE = "task_slice"
SLICE_OVERLAY_IDENTITY_SCOPE = "slice_overlay"
WORLD_AUTHORITY_IDENTITY_SCOPE = "world_authority"
SLICE_PARTICIPATION_CONTINUITY_SCOPE = "slice_participation"
WORLD_SHARED_CONTINUITY_SCOPE = "world_shared"
OBJECT_TYPE_SOURCE_DOCUMENT = "source_document"
OBJECT_TYPE_SOURCE_MATERIAL = "source_material"
OBJECT_TYPE_ARTIFACT = "artifact"
OBJECT_TYPE_EVIDENCE = "evidence"
PACK_REASON_DOMAIN_MATCHES = {
    "operations_pack": {"營運"},
    "finance_fundraising_pack": {"財務", "募資"},
    "legal_risk_pack": {"法務"},
    "marketing_sales_pack": {"行銷", "銷售"},
    "business_development_pack": {"銷售", "商務開發"},
    "research_intelligence_pack": {"研究", "情報", "綜合"},
    "organization_people_pack": {"組織人力", "組織", "人力"},
    "product_service_pack": {"產品服務", "產品", "服務"},
}
INDUSTRY_PACK_REASON_HINTS = {
    "online_education_pack": {
        "線上教育",
        "線上課程",
        "課程",
        "教學",
        "招生",
        "cohort",
        "bootcamp",
        "edtech",
    },
    "ecommerce_pack": {
        "電商",
        "ecommerce",
        "shopify",
        "商城",
        "蝦皮",
        "momo",
        "商品",
        "sku",
    },
    "gaming_pack": {
        "遊戲",
        "gaming",
        "game",
        "玩家",
        "live ops",
        "steam",
        "retention",
        "發行",
    },
    "funeral_services_pack": {
        "殯葬",
        "禮儀",
        "喪葬",
        "funeral",
        "memorial",
        "生前契約",
        "殯儀",
    },
    "health_supplements_pack": {
        "保健",
        "保健食品",
        "健康食品",
        "supplement",
        "supplements",
        "nutraceutical",
        "維他命",
        "益生菌",
    },
    "energy_pack": {
        "energy",
        "能源",
        "電力",
        "儲能",
        "solar",
        "renewable",
        "ppa",
        "epc",
    },
    "saas_pack": {
        "saas",
        "software",
        "subscription",
        "arr",
        "mrr",
        "plg",
        "churn",
    },
    "media_creator_pack": {
        "creator",
        "media",
        "自媒體",
        "內容創作者",
        "youtube",
        "podcast",
        "newsletter",
        "業配",
    },
    "professional_services_pack": {
        "consulting",
        "agency",
        "顧問",
        "專業服務",
        "retainer",
        "managed service",
        "代操",
        "代管",
    },
    "manufacturing_pack": {
        "manufacturing",
        "工廠",
        "製造",
        "oem",
        "odm",
        "產能",
        "良率",
        "供應鏈",
    },
    "healthcare_clinic_pack": {
        "clinic",
        "healthcare",
        "medical",
        "診所",
        "門診",
        "醫療",
        "病患",
        "療程",
    },
}
CLIENT_STAGE_KEYWORDS = {
    "創業階段": ("創業", "早期", "新創", "起步", "初期", "pmf", "驗證"),
    "制度化階段": ("制度化", "流程", "sop", "管理", "內控", "團隊", "穩定化"),
    "規模化階段": ("規模化", "擴張", "成長", "授權", "跨部門", "複製", "scale"),
}
CLIENT_TYPE_KEYWORDS = {
    "中小企業": ("中小企業", "公司", "企業", "工廠", "零售", "品牌方"),
    "個人品牌與服務": ("個人品牌", "顧問", "教練", "講師", "服務型", "freelance"),
    "自媒體": ("自媒體", "內容", "社群", "頻道", "podcast", "newsletter", "youtuber"),
    "大型企業": ("大型企業", "集團", "上市", "總部", "跨國", "enterprise"),
}
DOMAIN_LENS_KEYWORDS = {
    "營運": ("營運", "流程", "效率", "供應鏈", "交付", "執行"),
    "財務": ("財務", "現金流", "損益", "預算", "成本", "資金"),
    "法務": ("法務", "合約", "條款", "責任", "權利", "compliance", "liability", "nda"),
    "行銷": ("行銷", "品牌", "內容", "流量", "廣告", "社群", "campaign"),
    "銷售": ("銷售", "業務", "客戶開發", "pipeline", "成交", "proposal", "提案"),
    "募資": ("募資", "投資人", "term sheet", "融資", "fundraising", "cap table"),
    "組織人力": ("組織", "人力", "團隊", "招募", "人才", "管理", "ownership", "org"),
    "產品服務": ("產品", "服務", "方案", "定價", "sku", "offer", "package", "value proposition"),
}
TASK_TYPE_DOMAIN_HINTS = {
    "contract_review": ["法務"],
    "document_restructuring": ["綜合"],
    "research_synthesis": ["綜合"],
    "complex_convergence": ["綜合"],
}
TASK_TYPE_WORKSTREAM_HINTS = {
    "contract_review": "法務審閱工作流",
    "document_restructuring": "文件重組工作流",
    "research_synthesis": "研究綜整工作流",
    "complex_convergence": "決策收斂工作流",
}
EXTERNAL_EVENT_KEYWORDS = (
    "地緣政治",
    "關稅",
    "制裁",
    "匯率",
    "升息",
    "降息",
    "市場崩跌",
    "市場衝擊",
    "產業策略",
    "政策",
    "法規",
    "監管",
    "戰爭",
    "地緣",
    "景氣",
    "產業趨勢",
    "競爭格局",
    "macro",
    "geopolit",
    "regulation",
    "market shock",
    "tariff",
    "sanction",
)
EXTENSION_REGISTRY = ExtensionRegistry()
PACK_RESOLVER = PackResolver(EXTENSION_REGISTRY)
AGENT_RESOLVER = AgentResolver(EXTENSION_REGISTRY)


@dataclass
class InitialMaterialSnapshot:
    url_count: int
    file_count: int
    pasted_text_present: bool

    @property
    def total_count(self) -> int:
        return self.url_count + self.file_count + (1 if self.pasted_text_present else 0)


@dataclass
class CompiledCaseWorldSeed:
    identity: dict[str, object]
    client: schemas.ClientRead
    engagement: schemas.EngagementRead
    workstream: schemas.WorkstreamRead
    decision_context: schemas.DecisionContextRead
    domain_lenses: list[str]
    background_brief: str
    goal_description: str
    success_criteria: str
    inferred_constraints: list[schemas.ConstraintCreate]
    input_entry_mode: InputEntryMode
    external_research_heavy_candidate: bool
    deliverable_class_hint: DeliverableClass
    pack_resolution: schemas.PackResolutionRead
    agent_selection: schemas.AgentSelectionRead
    canonical_intake_summary: dict[str, object]
    task_interpretation: dict[str, object]
    extracted_objects: list[dict[str, object]]
    inferred_links: list[dict[str, object]]
    facts: list[dict[str, str]]
    assumptions_payload: list[dict[str, str]]
    evidence_gaps_payload: list[dict[str, object]]
    next_best_actions: list[str]


def resolve_external_data_strategy_from_constraints(
    constraints: list[models.Constraint] | list[schemas.ConstraintRead],
) -> ExternalDataStrategy:
    for constraint in constraints:
        if constraint.constraint_type != EXTERNAL_DATA_STRATEGY_CONSTRAINT_TYPE:
            continue

        try:
            return ExternalDataStrategy(constraint.description.strip())
        except ValueError:
            logger.warning(
                "Unknown external data strategy in constraint payload: %s",
                constraint.description,
            )

    return ExternalDataStrategy.SUPPLEMENTAL


def get_external_data_strategy_for_task(task: models.Task) -> ExternalDataStrategy:
    return resolve_external_data_strategy_from_constraints(task.constraints)


def _normalize_whitespace(value: str | None) -> str:
    return re.sub(r"\s+", " ", (value or "")).strip()


def _split_multiline_items(value: str | None) -> list[str]:
    return [item.strip() for item in (value or "").splitlines() if item.strip()]


def _unique_preserve_order(values: list[str]) -> list[str]:
    seen: set[str] = set()
    unique_values: list[str] = []
    for item in values:
        normalized = item.strip()
        if not normalized or normalized in seen:
            continue
        seen.add(normalized)
        unique_values.append(normalized)
    return unique_values


def _default_continuity_mode(entry_preset: InputEntryMode) -> EngagementContinuityMode:
    if entry_preset == InputEntryMode.ONE_LINE_INQUIRY:
        return EngagementContinuityMode.ONE_OFF
    return EngagementContinuityMode.FOLLOW_UP


def _default_writeback_depth(entry_preset: InputEntryMode) -> WritebackDepth:
    if entry_preset == InputEntryMode.ONE_LINE_INQUIRY:
        return WritebackDepth.MINIMAL
    return WritebackDepth.MILESTONE


def _normalize_continuity_mode(
    value: str | EngagementContinuityMode | None,
    *,
    fallback: InputEntryMode = InputEntryMode.ONE_LINE_INQUIRY,
) -> EngagementContinuityMode:
    if isinstance(value, EngagementContinuityMode):
        return value
    if value:
        try:
            return EngagementContinuityMode(value)
        except ValueError:
            logger.warning("Unknown engagement continuity mode: %s", value)
    return _default_continuity_mode(fallback)


def _normalize_writeback_depth(
    value: str | WritebackDepth | None,
    *,
    fallback: InputEntryMode = InputEntryMode.ONE_LINE_INQUIRY,
) -> WritebackDepth:
    if isinstance(value, WritebackDepth):
        return value
    if value:
        try:
            return WritebackDepth(value)
        except ValueError:
            logger.warning("Unknown writeback depth: %s", value)
    return _default_writeback_depth(fallback)


def _contains_any_keyword(text: str, keywords: tuple[str, ...]) -> bool:
    normalized = text.lower()
    return any(keyword.lower() in normalized for keyword in keywords)


def _infer_stage_from_text(text: str) -> str:
    normalized = text.lower()
    for stage, keywords in CLIENT_STAGE_KEYWORDS.items():
        if any(keyword.lower() in normalized for keyword in keywords):
            return stage
    return UNSPECIFIED_LABEL


def _infer_client_type_from_text(text: str) -> str:
    normalized = text.lower()
    for client_type, keywords in CLIENT_TYPE_KEYWORDS.items():
        if any(keyword.lower() in normalized for keyword in keywords):
            return client_type
    return UNSPECIFIED_LABEL


def _infer_domain_lenses_from_text(task_type: str, text: str) -> list[str]:
    normalized = text.lower()
    detected = list(TASK_TYPE_DOMAIN_HINTS.get(task_type, []))
    for lens, keywords in DOMAIN_LENS_KEYWORDS.items():
        if any(keyword.lower() in normalized for keyword in keywords):
            detected.append(lens)
    detected = [item for item in detected if item != DEFAULT_DOMAIN_LENS]
    return _unique_preserve_order(detected) or [DEFAULT_DOMAIN_LENS]


def _infer_client_name(payload: schemas.TaskCreateRequest) -> str:
    return _normalize_whitespace(payload.client_name) or DEFAULT_CLIENT_NAME


def _infer_client_stage(payload: schemas.TaskCreateRequest) -> str:
    explicit = _normalize_whitespace(payload.client_stage)
    if explicit:
        return explicit
    text = " ".join(
        filter(
            None,
            [
                payload.title,
                payload.description,
                payload.background_text,
                payload.goal_description,
                payload.assumptions,
                payload.notes,
            ],
        )
    )
    return _infer_stage_from_text(text)


def _infer_client_type(payload: schemas.TaskCreateRequest) -> str:
    explicit = _normalize_whitespace(payload.client_type)
    if explicit:
        return explicit
    text = " ".join(
        filter(
            None,
            [
                payload.title,
                payload.description,
                payload.background_text,
                payload.subject_name,
                payload.notes,
            ],
        )
    )
    return _infer_client_type_from_text(text)


def _infer_domain_lenses(payload: schemas.TaskCreateRequest) -> list[str]:
    explicit = _unique_preserve_order(payload.domain_lenses)
    if explicit:
        return explicit
    text = " ".join(
        filter(
            None,
            [
                payload.title,
                payload.description,
                payload.background_text,
                payload.goal_description,
                payload.subject_name,
                payload.notes,
            ],
        )
    )
    return _infer_domain_lenses_from_text(payload.task_type, text)


def _infer_engagement_name(payload: schemas.TaskCreateRequest) -> str:
    return _normalize_whitespace(payload.engagement_name) or payload.title


def _infer_workstream_name(payload: schemas.TaskCreateRequest) -> str:
    return (
        _normalize_whitespace(payload.workstream_name)
        or _normalize_whitespace(payload.subject_name)
        or TASK_TYPE_WORKSTREAM_HINTS.get(payload.task_type, "顧問工作流")
    )


def _build_source_priority_from_payload(payload: schemas.TaskCreateRequest) -> str:
    source_hints: list[str] = []
    if payload.description.strip():
        source_hints.append("原始問題")
    if payload.background_text.strip():
        source_hints.append("背景脈絡")
    if payload.notes and payload.notes.strip():
        source_hints.append("既有資料整理")
    if payload.subject_name:
        source_hints.append("分析對象說明")

    if not source_hints:
        return "目前主要只能依賴任務名稱與少量上下文進行第一輪判斷。"

    if len(source_hints) == 1:
        joined = source_hints[0]
    else:
        joined = "、".join(source_hints[:-1]) + f" 與 {source_hints[-1]}"
    return f"系統會先依據 {joined} 形成第一輪判斷脈絡，再視需要補充其他來源。"


def _build_external_data_policy_text(strategy: ExternalDataStrategy) -> str:
    if strategy == ExternalDataStrategy.STRICT:
        return "這輪判斷只使用你提供的資料，不主動補充外部來源。"
    if strategy == ExternalDataStrategy.LATEST:
        return "這輪判斷會優先補充最新外部資料，再與你提供的材料一起評估。"
    return "這輪判斷會先使用你提供的資料，若證據不足再由 Host 視需要補充外部來源。"


def _infer_decision_title(payload: schemas.TaskCreateRequest) -> str:
    return _normalize_whitespace(payload.decision_title) or f"{payload.title}｜Decision Context"


def _infer_background_brief(payload: schemas.TaskCreateRequest) -> str:
    if payload.background_text.strip():
        return payload.background_text

    lines = [
        f"工作流程：{payload.mode.value}",
        f"任務名稱：{payload.title}",
    ]
    if payload.description.strip():
        lines.append(f"核心問題：{payload.description.strip()}")
    if payload.subject_name:
        lines.append(f"分析對象：{payload.subject_name}")
    lines.append(
        f"外部資料使用方式：{EXTERNAL_DATA_STRATEGY_LABELS[payload.external_data_strategy]}"
    )
    if payload.assumptions:
        lines.append(f"已確定假設：{payload.assumptions}")

    return "\n".join(lines)


def _infer_goal_description(payload: schemas.TaskCreateRequest) -> str:
    if payload.goal_description:
        return payload.goal_description

    subject = payload.subject_name or "目前議題"
    if payload.task_type == "contract_review":
        return f"針對「{subject}」整理高風險條款、主要風險與下一步審閱建議。"
    if payload.task_type == "document_restructuring":
        return f"針對「{subject}」提出可直接使用的新版結構、重組策略與改寫重點。"
    if payload.mode.value == "multi_agent":
        return f"針對「{subject}」收斂主要建議、主要風險與下一步決策方向。"
    return f"針對「{subject}」整理摘要、關鍵發現、建議與缺漏資訊。"


def _infer_success_criteria(payload: schemas.TaskCreateRequest) -> str:
    if payload.success_criteria:
        return payload.success_criteria

    if payload.task_type == "contract_review":
        return "能清楚標示高風險條款、主要風險、建議處理方式與待補文件。"
    if payload.task_type == "document_restructuring":
        return "能提出可直接重組的新版骨架、改寫方向與後續行動。"
    if payload.mode.value == "multi_agent":
        return "能形成可供決策討論的收斂摘要、建議、風險、行動項目與未收斂處。"
    return "能清楚整理摘要、關鍵發現、洞察、建議與研究缺口。"


def _infer_constraints(payload: schemas.TaskCreateRequest) -> list[schemas.ConstraintCreate]:
    if payload.constraints:
        return payload.constraints

    if payload.task_type == "contract_review":
        description = "在正式法務審閱前，本結果應視為內部 issue spotting / redline 草稿。"
    elif payload.task_type == "document_restructuring":
        description = "若原始草稿或上下文不足，重構建議應視為工作骨架而非最終定稿。"
    elif payload.mode.value == "multi_agent":
        description = "若證據厚度不足，多代理結果應視為收斂骨架而非最終決策。"
    else:
        description = "若資料來源仍偏薄，這份分析應視為第一輪顧問草稿。"

    return [
        schemas.ConstraintCreate(
            description=description,
            constraint_type="system_inferred",
            severity="medium",
        )
    ]


def _build_external_data_strategy_constraint(
    strategy: ExternalDataStrategy,
) -> schemas.ConstraintCreate:
    return schemas.ConstraintCreate(
        description=strategy.value,
        constraint_type=EXTERNAL_DATA_STRATEGY_CONSTRAINT_TYPE,
        severity="low",
    )


def task_load_options():
    return (
        selectinload(models.Task.matter_workspace_links)
        .selectinload(models.MatterWorkspaceTaskLink.matter_workspace)
        .selectinload(models.MatterWorkspace.case_world_state),
        selectinload(models.Task.clients),
        selectinload(models.Task.engagements),
        selectinload(models.Task.workstreams),
        selectinload(models.Task.decision_contexts),
        selectinload(models.Task.contexts),
        selectinload(models.Task.subjects),
        selectinload(models.Task.goals),
        selectinload(models.Task.constraints),
        selectinload(models.Task.uploads),
        selectinload(models.Task.source_materials),
        selectinload(models.Task.artifacts),
        selectinload(models.Task.evidence),
        selectinload(models.Task.insights),
        selectinload(models.Task.risks),
        selectinload(models.Task.options),
        selectinload(models.Task.recommendations),
        selectinload(models.Task.action_items),
        selectinload(models.Task.recommendation_evidence_links),
        selectinload(models.Task.risk_evidence_links),
        selectinload(models.Task.action_item_evidence_links),
        selectinload(models.Task.deliverables).selectinload(models.Deliverable.object_links),
        selectinload(models.Task.runs),
        selectinload(models.Task.case_world_drafts),
        selectinload(models.Task.evidence_gaps),
        selectinload(models.Task.research_runs).selectinload(models.ResearchRun.source_documents),
        selectinload(models.Task.decision_records),
        selectinload(models.Task.action_plans),
        selectinload(models.Task.action_executions),
        selectinload(models.Task.outcome_records),
    )


def _primary_item(items: list[models.Client] | list[models.Engagement] | list[models.Workstream] | list[models.DecisionContext]):
    if not items:
        return None

    for item in items:
        if getattr(item, "identity_scope", LEGACY_TASK_SLICE_IDENTITY_SCOPE) != WORLD_AUTHORITY_IDENTITY_SCOPE:
            return item
    return items[0]


def _task_reference_role_for_identity_scope(identity_scope: str | None) -> str:
    if identity_scope == WORLD_AUTHORITY_IDENTITY_SCOPE:
        return "compatibility_only"
    if identity_scope == SLICE_OVERLAY_IDENTITY_SCOPE:
        return "slice_linkage"
    return "legacy_owner"


def _canonical_workspace_value(value: str | None, fallback: str) -> str:
    normalized = _normalize_whitespace(value)
    return normalized or fallback


def _build_matter_workspace_key(
    client_name: str,
    engagement_name: str,
    workstream_name: str,
) -> str:
    raw_key = " :: ".join(
        [
            _normalize_whitespace(client_name).lower(),
            _normalize_whitespace(engagement_name).lower(),
            _normalize_whitespace(workstream_name).lower(),
        ]
    )
    return hashlib.sha1(raw_key.encode("utf-8")).hexdigest()


def _build_matter_workspace_title(engagement_name: str, workstream_name: str) -> str:
    if _normalize_whitespace(engagement_name) == _normalize_whitespace(workstream_name):
        return engagement_name
    return f"{engagement_name}｜{workstream_name}"


def _build_matter_workspace_object_path(
    client_name: str,
    engagement_name: str,
    workstream_name: str,
) -> str:
    return " / ".join([client_name, engagement_name, workstream_name])


def _derive_matter_workspace_identity_from_values(
    *,
    task_title: str,
    task_type: str,
    client_name: str | None,
    engagement_name: str | None,
    workstream_name: str | None,
    client_type: str | None,
    client_stage: str | None,
    domain_lenses: list[str],
    decision_title: str | None,
    decision_summary: str | None,
) -> dict[str, object]:
    canonical_client_name = _canonical_workspace_value(client_name, DEFAULT_CLIENT_NAME)
    canonical_engagement_name = _canonical_workspace_value(engagement_name, task_title)
    canonical_workstream_name = _canonical_workspace_value(
        workstream_name,
        TASK_TYPE_WORKSTREAM_HINTS.get(task_type, "顧問工作流"),
    )
    workspace_lenses = _unique_preserve_order(
        [item for item in domain_lenses if item and item != DEFAULT_DOMAIN_LENS]
    ) or [DEFAULT_DOMAIN_LENS]
    return {
        "matter_key": _build_matter_workspace_key(
            canonical_client_name,
            canonical_engagement_name,
            canonical_workstream_name,
        ),
        "title": _build_matter_workspace_title(canonical_engagement_name, canonical_workstream_name),
        "client_name": canonical_client_name,
        "engagement_name": canonical_engagement_name,
        "workstream_name": canonical_workstream_name,
        "client_type": (client_type or "").strip() or UNSPECIFIED_LABEL,
        "client_stage": (client_stage or "").strip() or UNSPECIFIED_LABEL,
        "domain_lenses": workspace_lenses,
        "current_decision_context_title": _normalize_whitespace(decision_title) or None,
        "current_decision_context_summary": _normalize_whitespace(decision_summary) or None,
    }


def _derive_matter_workspace_identity(
    task: models.Task,
    client: schemas.ClientRead | None,
    engagement: schemas.EngagementRead | None,
    workstream: schemas.WorkstreamRead | None,
    decision_context: schemas.DecisionContextRead | None,
    domain_lenses: list[str],
) -> dict[str, object]:
    client_type = (
        decision_context.client_type
        if decision_context and decision_context.client_type and decision_context.client_type != UNSPECIFIED_LABEL
        else client.client_type if client and client.client_type != UNSPECIFIED_LABEL else ""
    )
    client_stage = (
        decision_context.client_stage
        if decision_context and decision_context.client_stage and decision_context.client_stage != UNSPECIFIED_LABEL
        else client.client_stage if client and client.client_stage != UNSPECIFIED_LABEL else ""
    )
    return _derive_matter_workspace_identity_from_values(
        task_title=task.title,
        task_type=task.task_type,
        client_name=client.name if client else None,
        engagement_name=engagement.name if engagement else None,
        workstream_name=workstream.name if workstream else None,
        client_type=client_type,
        client_stage=client_stage,
        domain_lenses=domain_lenses,
        decision_title=decision_context.title if decision_context else None,
        decision_summary=decision_context.summary if decision_context else None,
    )


def _extract_assumptions(task: models.Task) -> list[str]:
    latest_context = task.contexts[-1] if task.contexts else None
    return _split_multiline_items(latest_context.assumptions if latest_context else None)


def _build_legacy_source_material_read(source_document: models.SourceDocument) -> schemas.SourceMaterialRead:
    return schemas.SourceMaterialRead(
        id=source_document.id,
        task_id=source_document.task_id,
        matter_workspace_id=source_document.matter_workspace_id,
        source_document_id=source_document.id,
        continuity_scope=source_document.continuity_scope,
        source_type=source_document.source_type,
        title=source_document.canonical_display_name or source_document.file_name,
        canonical_display_name=source_document.canonical_display_name or source_document.file_name,
        source_ref=source_document.storage_path,
        file_extension=source_document.file_extension,
        content_type=source_document.content_type,
        file_size=source_document.file_size,
        storage_key=source_document.storage_key,
        storage_kind=source_document.storage_kind,
        storage_provider=source_document.storage_provider,
        content_digest=source_document.content_digest,
        ingest_status=source_document.ingest_status,
        ingest_strategy=source_document.ingest_strategy,
        support_level=source_document.support_level,
        retention_policy=source_document.retention_policy,
        purge_at=source_document.purge_at,
        availability_state=source_document.availability_state,
        metadata_only=source_document.metadata_only,
        summary=_normalize_whitespace(build_source_material_summary(source_document)),
        participation=schemas.ObjectParticipationRead(
            canonical_object_id=None,
            canonical_owner_scope="slice_local",
            compatibility_task_id=source_document.task_id,
            current_task_participation=True,
            participation_type="slice_local_only",
            participation_task_count=1,
            mapping_mode="slice_local_only",
        ),
        created_at=source_document.created_at,
        updated_at=source_document.updated_at,
    )


def _build_client_read_from_row(
    client: models.Client,
    *,
    reference_task_id: str | None = None,
) -> schemas.ClientRead:
    identity_scope = client.identity_scope or LEGACY_TASK_SLICE_IDENTITY_SCOPE
    return schemas.ClientRead(
        id=client.id,
        task_id=(
            reference_task_id
            if identity_scope == WORLD_AUTHORITY_IDENTITY_SCOPE and reference_task_id
            else client.task_id
        ),
        matter_workspace_id=client.matter_workspace_id,
        identity_scope=identity_scope,
        task_reference_role=_task_reference_role_for_identity_scope(identity_scope),
        name=client.name,
        client_type=client.client_type,
        client_stage=client.client_stage,
        description=client.description,
        created_at=client.created_at,
    )


def _build_engagement_read_from_row(
    engagement: models.Engagement,
    *,
    reference_task_id: str | None = None,
) -> schemas.EngagementRead:
    identity_scope = engagement.identity_scope or LEGACY_TASK_SLICE_IDENTITY_SCOPE
    return schemas.EngagementRead(
        id=engagement.id,
        task_id=(
            reference_task_id
            if identity_scope == WORLD_AUTHORITY_IDENTITY_SCOPE and reference_task_id
            else engagement.task_id
        ),
        matter_workspace_id=engagement.matter_workspace_id,
        identity_scope=identity_scope,
        task_reference_role=_task_reference_role_for_identity_scope(identity_scope),
        client_id=engagement.client_id,
        name=engagement.name,
        description=engagement.description,
        created_at=engagement.created_at,
    )


def _build_workstream_read_from_row(
    workstream: models.Workstream,
    *,
    reference_task_id: str | None = None,
) -> schemas.WorkstreamRead:
    identity_scope = workstream.identity_scope or LEGACY_TASK_SLICE_IDENTITY_SCOPE
    return schemas.WorkstreamRead(
        id=workstream.id,
        task_id=(
            reference_task_id
            if identity_scope == WORLD_AUTHORITY_IDENTITY_SCOPE and reference_task_id
            else workstream.task_id
        ),
        matter_workspace_id=workstream.matter_workspace_id,
        identity_scope=identity_scope,
        task_reference_role=_task_reference_role_for_identity_scope(identity_scope),
        engagement_id=workstream.engagement_id,
        name=workstream.name,
        description=workstream.description,
        domain_lenses=workstream.domain_lenses,
        created_at=workstream.created_at,
    )


def _build_source_document_read(
    source_document: models.SourceDocument,
    *,
    task: models.Task,
    snapshot: _TaskObjectParticipationSnapshot,
) -> schemas.SourceDocumentRead:
    return schemas.SourceDocumentRead(
        **schemas.SourceDocumentRead.model_validate(source_document).model_dump(exclude={"participation"}),
        participation=_build_object_participation_read(
            object_type=OBJECT_TYPE_SOURCE_DOCUMENT,
            object_id=source_document.id,
            object_task_id=source_document.task_id,
            current_task_id=task.id,
            matter_workspace_id=source_document.matter_workspace_id,
            continuity_scope=source_document.continuity_scope,
            snapshot=snapshot,
        ),
    )


def _build_legacy_artifact_read(
    source_document: models.SourceDocument,
    source_material_id: str | None,
) -> schemas.ArtifactRead:
    return schemas.ArtifactRead(
        id=source_document.id,
        task_id=source_document.task_id,
        matter_workspace_id=source_document.matter_workspace_id,
        continuity_scope=source_document.continuity_scope,
        title=source_document.file_name,
        artifact_type=infer_artifact_type(source_document),
        source_document_id=source_document.id,
        source_material_id=source_material_id,
        description=_normalize_whitespace(build_source_material_summary(source_document)[:280]),
        participation=schemas.ObjectParticipationRead(
            canonical_object_id=None,
            canonical_owner_scope="slice_local",
            compatibility_task_id=source_document.task_id,
            current_task_participation=True,
            participation_type="slice_local_only",
            participation_task_count=1,
            mapping_mode="slice_local_only",
        ),
        created_at=source_document.created_at,
    )


@dataclass
class _TaskObjectParticipationSnapshot:
    current_links: dict[tuple[str, str], models.TaskObjectParticipationLink]
    participation_counts: dict[tuple[str, str], int]


def _load_task_object_participation_snapshot(
    db: Session | None,
    *,
    matter_workspace_id: str | None,
    task_id: str,
) -> _TaskObjectParticipationSnapshot:
    if db is None or matter_workspace_id is None:
        return _TaskObjectParticipationSnapshot(current_links={}, participation_counts={})

    rows = db.scalars(
        select(models.TaskObjectParticipationLink)
        .where(models.TaskObjectParticipationLink.matter_workspace_id == matter_workspace_id)
        .order_by(models.TaskObjectParticipationLink.created_at)
    ).all()
    current_links: dict[tuple[str, str], models.TaskObjectParticipationLink] = {}
    participation_counts: dict[tuple[str, str], int] = {}
    for row in rows:
        canonical_key = resolve_participation_canonical_key(row)
        key = (row.object_type, canonical_key)
        participation_counts[key] = participation_counts.get(key, 0) + 1
        if row.task_id == task_id:
            current_links[key] = row
    return _TaskObjectParticipationSnapshot(
        current_links=current_links,
        participation_counts=participation_counts,
    )


def _build_object_participation_read(
    *,
    object_type: str,
    object_id: str,
    object_task_id: str,
    current_task_id: str,
    matter_workspace_id: str | None,
    continuity_scope: str,
    snapshot: _TaskObjectParticipationSnapshot,
) -> schemas.ObjectParticipationRead | None:
    if continuity_scope == WORLD_SHARED_CONTINUITY_SCOPE and matter_workspace_id:
        key = (object_type, object_id)
        current_link = snapshot.current_links.get(key)
        participation_count = snapshot.participation_counts.get(key, 0)
        fallback_to_task_ref = participation_count == 0 and object_task_id == current_task_id
        return schemas.ObjectParticipationRead(
            canonical_object_id=object_id,
            canonical_owner_scope="world_canonical",
            compatibility_task_id=object_task_id,
            current_task_participation=current_link is not None or fallback_to_task_ref,
            participation_type=(
                current_link.participation_type
                if current_link is not None
                else "direct_ingest" if fallback_to_task_ref else None
            ),
            participation_task_count=max(participation_count, 1),
            mapping_mode=(
                "explicit_mapping"
                if participation_count > 0
                else "compatibility_task_ref" if fallback_to_task_ref else None
            ),
        )

    if object_task_id == current_task_id:
        return schemas.ObjectParticipationRead(
            canonical_object_id=None,
            canonical_owner_scope="slice_local",
            compatibility_task_id=object_task_id,
            current_task_participation=True,
            participation_type="slice_local_only",
            participation_task_count=1,
            mapping_mode="slice_local_only",
        )
    return None


def _backfill_world_shared_source_chain_participation_links(
    db: Session,
    *,
    task: models.Task,
    matter_workspace_id: str | None,
) -> None:
    if not matter_workspace_id:
        return

    for source_document in task.uploads:
        if source_document.continuity_scope != WORLD_SHARED_CONTINUITY_SCOPE:
            continue
        ensure_source_chain_participation_links(
            db,
            task_id=task.id,
            matter_workspace_id=matter_workspace_id,
            source_document_id=source_document.id,
            source_material_id=None,
            artifact_id=None,
            evidence_id=None,
            participation_type="direct_ingest",
        )
    for source_material in task.source_materials:
        if source_material.continuity_scope != WORLD_SHARED_CONTINUITY_SCOPE:
            continue
        ensure_source_chain_participation_links(
            db,
            task_id=task.id,
            matter_workspace_id=matter_workspace_id,
            source_document_id=source_material.source_document_id,
            source_material_id=source_material.id,
            artifact_id=None,
            evidence_id=None,
            participation_type="direct_ingest",
        )
    for artifact in task.artifacts:
        if artifact.continuity_scope != WORLD_SHARED_CONTINUITY_SCOPE:
            continue
        ensure_source_chain_participation_links(
            db,
            task_id=task.id,
            matter_workspace_id=matter_workspace_id,
            source_document_id=artifact.source_document_id,
            source_material_id=artifact.source_material_id,
            artifact_id=artifact.id,
            evidence_id=None,
            participation_type="direct_ingest",
        )
    for evidence in task.evidence:
        if evidence.continuity_scope != WORLD_SHARED_CONTINUITY_SCOPE:
            continue
        ensure_source_chain_participation_links(
            db,
            task_id=task.id,
            matter_workspace_id=matter_workspace_id,
            source_document_id=evidence.source_document_id,
            source_material_id=evidence.source_material_id,
            artifact_id=evidence.artifact_id,
            evidence_id=evidence.id,
            participation_type="direct_ingest",
        )


def _ensure_source_chain_participation_snapshot_ready(
    task: models.Task,
) -> tuple[Session | None, models.MatterWorkspace | None, str | None]:
    db = object_session(task)
    matter_workspace = _get_linked_matter_workspace(task)
    matter_workspace_id = matter_workspace.id if matter_workspace is not None else None
    if db is not None and matter_workspace_id is not None:
        _backfill_world_shared_source_chain_participation_links(
            db,
            task=task,
            matter_workspace_id=matter_workspace_id,
        )
        db.flush()
    return db, matter_workspace, matter_workspace_id


def _build_source_documents(task: models.Task) -> list[schemas.SourceDocumentRead]:
    db, matter_workspace, matter_workspace_id = _ensure_source_chain_participation_snapshot_ready(task)
    participation_snapshot = _load_task_object_participation_snapshot(
        db,
        matter_workspace_id=matter_workspace_id,
        task_id=task.id,
    )
    documents = [
        _build_source_document_read(item, task=task, snapshot=participation_snapshot)
        for item in task.uploads
    ]
    if db is not None and matter_workspace is not None:
        world_documents = db.scalars(
            select(models.SourceDocument)
            .where(models.SourceDocument.matter_workspace_id == matter_workspace.id)
            .where(models.SourceDocument.continuity_scope == WORLD_SHARED_CONTINUITY_SCOPE)
            .order_by(models.SourceDocument.created_at)
        ).all()
        covered_ids = {item.id for item in documents}
        for item in world_documents:
            if item.id in covered_ids:
                continue
            documents.append(
                _build_source_document_read(item, task=task, snapshot=participation_snapshot)
            )
            covered_ids.add(item.id)
    return sorted(documents, key=lambda item: item.created_at)


def build_upload_result_item_from_aggregate(
    aggregate: schemas.TaskAggregateResponse,
    *,
    source_document_id: str,
    evidence_id: str,
    source_material_id: str | None = None,
    artifact_id: str | None = None,
) -> schemas.UploadResultItem:
    source_documents_by_id = {item.id: item for item in aggregate.uploads}
    source_materials_by_id = {item.id: item for item in aggregate.source_materials}
    artifacts_by_id = {item.id: item for item in aggregate.artifacts}
    evidence_by_id = {item.id: item for item in aggregate.evidence}

    source_document = source_documents_by_id.get(source_document_id)
    evidence = evidence_by_id.get(evidence_id)
    source_material = source_materials_by_id.get(source_material_id) if source_material_id else None
    artifact = artifacts_by_id.get(artifact_id) if artifact_id else None

    if source_document is None or evidence is None:
        raise RuntimeError("Upload/source response could not resolve canonical aggregate rows.")

    if source_material is None and source_material_id is None and evidence.source_material_id:
        source_material = source_materials_by_id.get(evidence.source_material_id)
    if artifact is None and artifact_id is None and evidence.artifact_id:
        artifact = artifacts_by_id.get(evidence.artifact_id)

    return schemas.UploadResultItem(
        source_document=source_document,
        evidence=evidence,
        source_material=source_material,
        artifact=artifact,
    )


def _build_source_materials(task: models.Task) -> list[schemas.SourceMaterialRead]:
    db, matter_workspace, matter_workspace_id = _ensure_source_chain_participation_snapshot_ready(task)
    participation_snapshot = _load_task_object_participation_snapshot(
        db,
        matter_workspace_id=matter_workspace_id,
        task_id=task.id,
    )
    materials = [
        schemas.SourceMaterialRead(
            **schemas.SourceMaterialRead.model_validate(item).model_dump(exclude={"participation"}),
            participation=_build_object_participation_read(
                object_type=OBJECT_TYPE_SOURCE_MATERIAL,
                object_id=item.id,
                object_task_id=item.task_id,
                current_task_id=task.id,
                matter_workspace_id=item.matter_workspace_id,
                continuity_scope=item.continuity_scope,
                snapshot=participation_snapshot,
            ),
        )
        for item in task.source_materials
    ]
    if db is not None and matter_workspace is not None:
        world_materials = db.scalars(
            select(models.SourceMaterial)
            .where(models.SourceMaterial.matter_workspace_id == matter_workspace.id)
            .where(models.SourceMaterial.continuity_scope == WORLD_SHARED_CONTINUITY_SCOPE)
            .order_by(models.SourceMaterial.created_at)
        ).all()
        covered_ids = {item.id for item in materials}
        for item in world_materials:
            if item.id in covered_ids:
                continue
            materials.append(
                schemas.SourceMaterialRead(
                    **schemas.SourceMaterialRead.model_validate(item).model_dump(exclude={"participation"}),
                    participation=_build_object_participation_read(
                        object_type=OBJECT_TYPE_SOURCE_MATERIAL,
                        object_id=item.id,
                        object_task_id=item.task_id,
                        current_task_id=task.id,
                        matter_workspace_id=item.matter_workspace_id,
                        continuity_scope=item.continuity_scope,
                        snapshot=participation_snapshot,
                    ),
                )
            )
            covered_ids.add(item.id)
    covered_source_documents = {
        item.source_document_id for item in materials if item.source_document_id
    }
    for item in task.uploads:
        if item.id in covered_source_documents:
            continue
        materials.append(_build_legacy_source_material_read(item))
    return sorted(materials, key=lambda item: item.created_at)


def _build_artifacts(
    task: models.Task,
    source_materials: list[schemas.SourceMaterialRead],
) -> list[schemas.ArtifactRead]:
    db, matter_workspace, matter_workspace_id = _ensure_source_chain_participation_snapshot_ready(task)
    participation_snapshot = _load_task_object_participation_snapshot(
        db,
        matter_workspace_id=matter_workspace_id,
        task_id=task.id,
    )
    artifacts = [
        schemas.ArtifactRead(
            **schemas.ArtifactRead.model_validate(item).model_dump(exclude={"participation"}),
            participation=_build_object_participation_read(
                object_type=OBJECT_TYPE_ARTIFACT,
                object_id=item.id,
                object_task_id=item.task_id,
                current_task_id=task.id,
                matter_workspace_id=item.matter_workspace_id,
                continuity_scope=item.continuity_scope,
                snapshot=participation_snapshot,
            ),
        )
        for item in task.artifacts
    ]
    if db is not None and matter_workspace is not None:
        world_artifacts = db.scalars(
            select(models.Artifact)
            .where(models.Artifact.matter_workspace_id == matter_workspace.id)
            .where(models.Artifact.continuity_scope == WORLD_SHARED_CONTINUITY_SCOPE)
            .order_by(models.Artifact.created_at)
        ).all()
        covered_ids = {item.id for item in artifacts}
        for item in world_artifacts:
            if item.id in covered_ids:
                continue
            artifacts.append(
                schemas.ArtifactRead(
                    **schemas.ArtifactRead.model_validate(item).model_dump(exclude={"participation"}),
                    participation=_build_object_participation_read(
                        object_type=OBJECT_TYPE_ARTIFACT,
                        object_id=item.id,
                        object_task_id=item.task_id,
                        current_task_id=task.id,
                        matter_workspace_id=item.matter_workspace_id,
                        continuity_scope=item.continuity_scope,
                        snapshot=participation_snapshot,
                    ),
                )
            )
            covered_ids.add(item.id)
    source_material_by_document = {
        item.source_document_id: item.id
        for item in source_materials
        if item.source_document_id
    }
    covered_source_documents = {
        item.source_document_id for item in artifacts if item.source_document_id
    }
    for item in task.uploads:
        if item.id in covered_source_documents:
            continue
        artifacts.append(
            _build_legacy_artifact_read(
                item,
                source_material_by_document.get(item.id),
            )
        )
    return sorted(artifacts, key=lambda item: item.created_at)


def _meaningful_source_materials(
    source_materials: list[schemas.SourceMaterialRead],
) -> list[schemas.SourceMaterialRead]:
    return [
        item
        for item in source_materials
        if item.ingest_status in {"processed", "metadata_only"}
        and (
            _normalize_whitespace(item.summary)
            or item.metadata_only
            or _normalize_whitespace(item.canonical_display_name)
        )
    ]


def _meaningful_artifacts(artifacts: list[schemas.ArtifactRead]) -> list[schemas.ArtifactRead]:
    return [
        item
        for item in artifacts
        if _normalize_whitespace(item.title) or _normalize_whitespace(item.description)
    ]


def _material_unit_count(
    source_materials: list[schemas.SourceMaterialRead],
    artifacts: list[schemas.ArtifactRead],
) -> int:
    material_keys = {
        item.source_document_id or f"source_material:{item.id}"
        for item in _meaningful_source_materials(source_materials)
    }
    artifact_keys = {
        item.source_document_id or item.source_material_id or f"artifact:{item.id}"
        for item in _meaningful_artifacts(artifacts)
    }
    return len({*material_keys, *artifact_keys})


def _usable_evidence(task: models.Task) -> list[models.Evidence]:
    return [
        item
        for item in task.evidence
        if not item.evidence_type.endswith("ingestion_issue")
        and not item.evidence_type.endswith("unparsed")
        and _normalize_whitespace(item.excerpt_or_summary)
    ]


def _infer_input_entry_mode(
    task: models.Task,
    source_materials: list[schemas.SourceMaterialRead],
    artifacts: list[schemas.ArtifactRead],
) -> InputEntryMode:
    material_count = _material_unit_count(source_materials, artifacts)

    if material_count >= 2:
        return InputEntryMode.MULTI_MATERIAL_CASE
    if material_count == 1:
        return InputEntryMode.SINGLE_DOCUMENT_INTAKE
    return InputEntryMode.ONE_LINE_INQUIRY


def _is_external_research_heavy_candidate(
    task: models.Task,
    decision_context: schemas.DecisionContextRead | None,
    source_materials: list[schemas.SourceMaterialRead],
    artifacts: list[schemas.ArtifactRead],
    input_entry_mode: InputEntryMode,
) -> bool:
    if input_entry_mode != InputEntryMode.ONE_LINE_INQUIRY:
        return False

    if _meaningful_source_materials(source_materials) or _meaningful_artifacts(artifacts):
        return False

    signal_text = " ".join(
        filter(
            None,
            [
                task.title,
                task.description,
                decision_context.title if decision_context else "",
                decision_context.summary if decision_context else "",
                decision_context.judgment_to_make if decision_context else "",
                *[item.description for item in task.goals],
            ],
        )
    )
    return _contains_any_keyword(signal_text, EXTERNAL_EVENT_KEYWORDS)


def _build_presence_state_item(
    state: PresenceState,
    reason: str,
    display_value: str | None = None,
) -> schemas.PresenceStateItemRead:
    return schemas.PresenceStateItemRead(
        state=state,
        reason=reason,
        display_value=_normalize_whitespace(display_value),
    )


def _build_presence_state_summary(
    task: models.Task,
    client: schemas.ClientRead | None,
    engagement: schemas.EngagementRead | None,
    workstream: schemas.WorkstreamRead | None,
    decision_context: schemas.DecisionContextRead | None,
    domain_lenses: list[str],
    source_materials: list[schemas.SourceMaterialRead],
    artifacts: list[schemas.ArtifactRead],
    input_entry_mode: InputEntryMode,
    external_research_heavy_candidate: bool,
) -> schemas.PresenceStateSummaryRead:
    meaningful_source_materials = _meaningful_source_materials(source_materials)
    meaningful_artifacts = _meaningful_artifacts(artifacts)
    explicit_lenses = [item for item in domain_lenses if item and item != DEFAULT_DOMAIN_LENS]
    default_decision_title = f"{task.title}｜Decision Context"
    default_decision_summary = (
        f"本輪要圍繞「{task.description or task.title}」形成可採用的顧問判斷。"
    )
    goal_descriptions = {_normalize_whitespace(item.description) for item in task.goals}

    if external_research_heavy_candidate and (
        not client or client.name == DEFAULT_CLIENT_NAME
    ):
        client_state = _build_presence_state_item(
            PresenceState.NOT_APPLICABLE,
            "這輪先屬於外部事件導向判斷，尚未綁定到特定 client。",
        )
    elif client is None:
        client_state = _build_presence_state_item(
            PresenceState.MISSING,
            "目前尚未建立 client object。",
        )
    elif client.name != DEFAULT_CLIENT_NAME:
        client_state = _build_presence_state_item(
            PresenceState.EXPLICIT,
            "目前已存在可直接引用的 client 主體。",
            client.name,
        )
    elif client.client_type != UNSPECIFIED_LABEL or client.client_stage != UNSPECIFIED_LABEL:
        client_state = _build_presence_state_item(
            PresenceState.INFERRED,
            "Host 已根據現有內容推定 client 脈絡，但尚未有明確 client 名稱。",
            client.name,
        )
    else:
        client_state = _build_presence_state_item(
            PresenceState.PROVISIONAL,
            "目前 client 仍是系統為了建立工作世界而保留的暫定物件。",
            client.name if client else None,
        )

    if external_research_heavy_candidate and (
        not engagement or engagement.name == task.title
    ):
        engagement_state = _build_presence_state_item(
            PresenceState.NOT_APPLICABLE,
            "這輪尚未進入特定 engagement 範圍，先以外部態勢判斷為主。",
        )
    elif engagement is None:
        engagement_state = _build_presence_state_item(
            PresenceState.MISSING,
            "目前尚未建立 engagement。",
        )
    elif engagement.name != task.title:
        engagement_state = _build_presence_state_item(
            PresenceState.EXPLICIT,
            "已有可識別的 engagement。",
            engagement.name,
        )
    elif _normalize_whitespace(engagement.description):
        engagement_state = _build_presence_state_item(
            PresenceState.INFERRED,
            "engagement 已由任務內容推定，但尚未獨立命名。",
            engagement.name,
        )
    else:
        engagement_state = _build_presence_state_item(
            PresenceState.PROVISIONAL,
            "目前 engagement 仍是沿用任務主題建立的暫定工作殼層。",
            engagement.name if engagement else None,
        )

    if external_research_heavy_candidate and (
        not workstream or workstream.name in TASK_TYPE_WORKSTREAM_HINTS.values()
    ):
        workstream_state = _build_presence_state_item(
            PresenceState.NOT_APPLICABLE,
            "這輪先屬於外部態勢判斷，尚未落到明確 workstream。",
        )
    elif workstream is None:
        workstream_state = _build_presence_state_item(
            PresenceState.MISSING,
            "目前尚未建立 workstream。",
        )
    elif workstream.name not in TASK_TYPE_WORKSTREAM_HINTS.values():
        workstream_state = _build_presence_state_item(
            PresenceState.EXPLICIT,
            "已有可直接引用的 workstream。",
            workstream.name,
        )
    elif explicit_lenses or _normalize_whitespace(workstream.description):
        workstream_state = _build_presence_state_item(
            PresenceState.INFERRED,
            "workstream 已依據任務內容形成初步脈絡，但仍偏推定。",
            workstream.name,
        )
    else:
        workstream_state = _build_presence_state_item(
            PresenceState.PROVISIONAL,
            "目前 workstream 仍是系統保留的暫定工作流名稱。",
            workstream.name if workstream else None,
        )

    if decision_context is None:
        decision_context_state = _build_presence_state_item(
            PresenceState.MISSING,
            "目前尚未形成 decision context。",
        )
    elif (
        decision_context.title != default_decision_title
        or decision_context.summary != default_decision_summary
        or _normalize_whitespace(decision_context.judgment_to_make) not in goal_descriptions
    ):
        decision_context_state = _build_presence_state_item(
            PresenceState.EXPLICIT,
            "本輪 decision context 已形成可直接引用的判斷主軸。",
            decision_context.judgment_to_make or decision_context.title,
        )
    elif decision_context.judgment_to_make or decision_context.summary:
        decision_context_state = _build_presence_state_item(
            PresenceState.PROVISIONAL,
            "decision context 已建立，但仍偏向 Host 依任務文字先形成的 provisional framing。",
            decision_context.judgment_to_make or decision_context.title,
        )
    else:
        decision_context_state = _build_presence_state_item(
            PresenceState.MISSING,
            "目前尚未形成足以支撐判斷的 decision context。",
        )

    if meaningful_artifacts:
        artifact_state = _build_presence_state_item(
            PresenceState.EXPLICIT,
            f"目前已有 {len(meaningful_artifacts)} 份 artifact 可直接納入工作鏈。",
            meaningful_artifacts[0].title,
        )
    elif external_research_heavy_candidate:
        artifact_state = _build_presence_state_item(
            PresenceState.NOT_APPLICABLE,
            "這輪先屬於外部事件判斷，尚未期待有 company-specific artifact。",
        )
    else:
        artifact_state = _build_presence_state_item(
            PresenceState.MISSING,
            "目前尚未提供可直接引用的 artifact。",
        )

    if meaningful_source_materials:
        source_material_state = _build_presence_state_item(
            PresenceState.EXPLICIT,
            f"目前已有 {len(meaningful_source_materials)} 份 source material。",
            meaningful_source_materials[0].title,
        )
    elif input_entry_mode == InputEntryMode.ONE_LINE_INQUIRY:
        source_material_state = _build_presence_state_item(
            PresenceState.MISSING,
            "目前仍缺可支撐判斷的 source material。",
        )
    else:
        source_material_state = _build_presence_state_item(
            PresenceState.PROVISIONAL,
            "目前 source material 仍偏薄，後續可再補充。",
        )

    if explicit_lenses:
        domain_lens_state = _build_presence_state_item(
            PresenceState.INFERRED,
            "Host 已從現有內容建立可用的 DomainLens。",
            "、".join(explicit_lenses),
        )
    else:
        domain_lens_state = _build_presence_state_item(
            PresenceState.PROVISIONAL,
            "目前仍以綜合視角作為 provisional DomainLens。",
            DEFAULT_DOMAIN_LENS,
        )

    client_stage_value = (
        decision_context.client_stage
        if decision_context and decision_context.client_stage
        else client.client_stage if client else ""
    )
    if external_research_heavy_candidate and (
        not client_stage_value or client_stage_value == UNSPECIFIED_LABEL
    ):
        client_stage_state = _build_presence_state_item(
            PresenceState.NOT_APPLICABLE,
            "這輪尚未綁定到特定 company stage。",
        )
    elif client_stage_value and client_stage_value != UNSPECIFIED_LABEL:
        client_stage_state = _build_presence_state_item(
            PresenceState.INFERRED,
            "目前已有 client stage 脈絡，可用於限制判斷邊界。",
            client_stage_value,
        )
    else:
        client_stage_state = _build_presence_state_item(
            PresenceState.MISSING,
            "目前尚未清楚標示 client stage。",
        )

    client_type_value = (
        decision_context.client_type
        if decision_context and decision_context.client_type
        else client.client_type if client else ""
    )
    if external_research_heavy_candidate and (
        not client_type_value or client_type_value == UNSPECIFIED_LABEL
    ):
        client_type_state = _build_presence_state_item(
            PresenceState.NOT_APPLICABLE,
            "這輪尚未綁定到特定 client type。",
        )
    elif client_type_value and client_type_value != UNSPECIFIED_LABEL:
        client_type_state = _build_presence_state_item(
            PresenceState.INFERRED,
            "目前已有 client type 脈絡，可協助限制建議適用範圍。",
            client_type_value,
        )
    else:
        client_type_state = _build_presence_state_item(
            PresenceState.MISSING,
            "目前尚未清楚標示 client type。",
        )

    return schemas.PresenceStateSummaryRead(
        client=client_state,
        engagement=engagement_state,
        workstream=workstream_state,
        decision_context=decision_context_state,
        artifact=artifact_state,
        source_material=source_material_state,
        domain_lens=domain_lens_state,
        client_stage=client_stage_state,
        client_type=client_type_state,
    )


def _extract_explicit_pack_overrides(
    constraints: list[models.Constraint],
) -> list[str]:
    explicit_pack_ids: list[str] = []
    for constraint in constraints:
        if constraint.constraint_type != PACK_OVERRIDE_CONSTRAINT_TYPE:
            continue
        explicit_pack_ids.extend(
            item.strip()
            for item in re.split(r"[\s,，\n]+", constraint.description or "")
            if item.strip()
        )
    return _unique_preserve_order(explicit_pack_ids)


def _build_industry_hints(
    task: models.Task,
    client: schemas.ClientRead | None,
    engagement: schemas.EngagementRead | None,
    workstream: schemas.WorkstreamRead | None,
    decision_context: schemas.DecisionContextRead | None,
) -> list[str]:
    raw_hints = [
        task.title,
        task.description,
        client.name if client else "",
        client.client_type if client else "",
        client.client_stage if client else "",
        engagement.name if engagement else "",
        engagement.description if engagement else "",
        workstream.name if workstream else "",
        workstream.description if workstream else "",
        decision_context.title if decision_context else "",
        decision_context.summary if decision_context else "",
        decision_context.judgment_to_make if decision_context else "",
        *[item.name for item in task.subjects if item.name],
        *[item.description for item in task.goals if item.description],
    ]
    return _unique_preserve_order([item for item in raw_hints if item is not None])


def _build_selected_pack_reason(
    *,
    pack: PackSpec,
    explicit_pack_ids: list[str],
    domain_lenses: list[str],
    client_type: str | None,
    client_stage: str | None,
    industry_hints: list[str],
) -> str:
    reasons: list[str] = []
    if pack.pack_id in explicit_pack_ids:
        reasons.append("由使用者明確指定覆寫。")

    if pack.pack_type == PackType.DOMAIN:
        matched_lenses = [
            item
            for item in domain_lenses
            if item in PACK_REASON_DOMAIN_MATCHES.get(pack.pack_id, set())
        ]
        if matched_lenses:
            reasons.append(f"對齊 DomainLens：{join_natural_list(matched_lenses)}。")
    else:
        if client_type and client_type in pack.relevant_client_types:
            reasons.append(f"對齊客戶型態：{client_type}。")
        if client_stage and client_stage in pack.relevant_client_stages:
            reasons.append(f"對齊客戶階段：{client_stage}。")
        hint_text = " ".join(item.lower() for item in industry_hints if item)
        matched_hints = [
            item
            for item in INDUSTRY_PACK_REASON_HINTS.get(pack.pack_id, set())
            if item.lower() in hint_text
        ]
        if matched_hints:
            reasons.append(f"對齊產業線索：{join_natural_list(matched_hints[:2])}。")

    if not reasons:
        reasons.append("由 Host 依目前 DecisionContext 與 context spine 推定。")
    return " ".join(reasons)


def _serialize_selected_pack(
    *,
    pack: PackSpec,
    explicit_pack_ids: list[str],
    domain_lenses: list[str],
    client_type: str | None,
    client_stage: str | None,
    industry_hints: list[str],
) -> schemas.SelectedPackRead:
    return schemas.SelectedPackRead(
        pack_id=pack.pack_id,
        pack_type=pack.pack_type.value,
        pack_name=pack.pack_name,
        description=pack.description,
        domain_definition=pack.domain_definition,
        industry_definition=pack.industry_definition,
        common_business_models=pack.common_business_models,
        common_problem_patterns=pack.common_problem_patterns,
        stage_specific_heuristics=pack.stage_specific_heuristics,
        key_kpis_or_operating_signals=pack.key_kpis_or_operating_signals,
        key_kpis=pack.key_kpis,
        reason=_build_selected_pack_reason(
            pack=pack,
            explicit_pack_ids=explicit_pack_ids,
            domain_lenses=domain_lenses,
            client_type=client_type,
            client_stage=client_stage,
            industry_hints=industry_hints,
        ),
        status=pack.status.value,
        version=pack.version,
        evidence_expectations=pack.evidence_expectations,
        common_risks=pack.common_risks,
        decision_patterns=pack.decision_patterns,
        deliverable_presets=pack.deliverable_presets,
        routing_hints=pack.routing_hints,
        pack_notes=pack.pack_notes,
        scope_boundaries=pack.scope_boundaries,
        pack_rationale=pack.pack_rationale,
    )


def join_natural_list(items: list[str]) -> str:
    normalized = [item.strip() for item in items if item and item.strip()]
    if not normalized:
        return ""
    if len(normalized) == 1:
        return normalized[0]
    return f"{'、'.join(normalized[:-1])} 與 {normalized[-1]}"


def _extract_explicit_agent_overrides(
    constraints: list[models.Constraint],
) -> list[str]:
    explicit_agent_ids: list[str] = []
    for constraint in constraints:
        if constraint.constraint_type != AGENT_OVERRIDE_CONSTRAINT_TYPE:
            continue
        explicit_agent_ids.extend(
            item.strip()
            for item in re.split(r"[\s,，\n]+", constraint.description or "")
            if item.strip()
        )
    return _unique_preserve_order(explicit_agent_ids)


def _normalize_override_ids(values: list[str]) -> list[str]:
    normalized: list[str] = []
    for value in values:
        normalized.extend(
            item.strip()
            for item in re.split(r"[\s,，\n]+", value or "")
            if item.strip()
        )
    return _unique_preserve_order(normalized)


def _validate_pack_override_ids(pack_ids: list[str]) -> list[str]:
    unknown = [pack_id for pack_id in pack_ids if EXTENSION_REGISTRY.get_pack(pack_id) is None]
    if unknown:
        raise HTTPException(
            status_code=400,
            detail=f"Unknown pack override ids: {join_natural_list(unknown)}",
        )
    return pack_ids


def _validate_agent_override_ids(agent_ids: list[str]) -> list[str]:
    unknown = [agent_id for agent_id in agent_ids if EXTENSION_REGISTRY.get_agent(agent_id) is None]
    if unknown:
        raise HTTPException(
            status_code=400,
            detail=f"Unknown agent override ids: {join_natural_list(unknown)}",
        )
    return agent_ids


def _latest_deliverable(task: models.Task) -> models.Deliverable | None:
    return max(task.deliverables, key=lambda item: item.version, default=None)


def _infer_capability_from_deliverable(
    task: models.Task,
) -> CapabilityArchetype | None:
    latest_deliverable = _latest_deliverable(task)
    if not latest_deliverable or not isinstance(latest_deliverable.content_structure, dict):
        return None

    capability_frame = latest_deliverable.content_structure.get("capability_frame")
    if not isinstance(capability_frame, dict):
        return None

    capability = capability_frame.get("capability")
    if not isinstance(capability, str):
        return None

    try:
        return CapabilityArchetype(capability)
    except ValueError:
        return None


def _build_capability_signal_text(
    task: models.Task,
    decision_context: schemas.DecisionContextRead | None,
    domain_lenses: list[str],
) -> str:
    parts = [
        task.title,
        task.description,
        decision_context.title if decision_context else "",
        decision_context.summary if decision_context else "",
        decision_context.judgment_to_make if decision_context else "",
        *(goal.description for goal in task.goals),
        *(constraint.description for constraint in task.constraints),
        *domain_lenses,
    ]
    return " ".join(part.strip() for part in parts if part and part.strip()).lower()


def _infer_capability_for_task(
    task: models.Task,
    decision_context: schemas.DecisionContextRead | None,
    domain_lenses: list[str],
) -> CapabilityArchetype:
    existing = _infer_capability_from_deliverable(task)
    if existing is not None:
        return existing

    if task.task_type == "contract_review":
        return CapabilityArchetype.REVIEW_CHALLENGE
    if task.task_type == "document_restructuring":
        return CapabilityArchetype.RESTRUCTURE_REFRAME
    if task.task_type == "research_synthesis":
        return CapabilityArchetype.SYNTHESIZE_BRIEF

    signal_text = _build_capability_signal_text(task, decision_context, domain_lenses)
    if _contains_any_keyword(signal_text, ("比較", "方案", "option", "trade-off", "scenario")):
        return CapabilityArchetype.SCENARIO_COMPARISON
    if _contains_any_keyword(signal_text, ("路線圖", "roadmap", "規劃", "milestone", "計畫")):
        return CapabilityArchetype.PLAN_ROADMAP
    if _contains_any_keyword(signal_text, ("風險盤點", "risk surfacing", "法務風險", "合規風險", "risk register")):
        return CapabilityArchetype.RISK_SURFACING
    if _contains_any_keyword(signal_text, ("診斷", "盤點", "assess", "framing", "現況")):
        return CapabilityArchetype.DIAGNOSE_ASSESS
    if len([item for item in domain_lenses if item and item != DEFAULT_DOMAIN_LENS]) > 1:
        return CapabilityArchetype.DECIDE_CONVERGE
    return CapabilityArchetype.DECIDE_CONVERGE if task.task_type == "complex_convergence" else CapabilityArchetype.SYNTHESIZE_BRIEF


def _allow_specialists_for_selection(
    capability: CapabilityArchetype,
    input_entry_mode: InputEntryMode,
    external_research_heavy_candidate: bool,
) -> bool:
    if external_research_heavy_candidate or input_entry_mode == InputEntryMode.ONE_LINE_INQUIRY:
        return False
    return capability in {
        CapabilityArchetype.REVIEW_CHALLENGE,
        CapabilityArchetype.SYNTHESIZE_BRIEF,
        CapabilityArchetype.RESTRUCTURE_REFRAME,
    }


def _build_selected_agent_reason(
    *,
    agent: AgentSpec,
    capability: CapabilityArchetype,
    explicit_agent_ids: list[str],
    selected_domain_packs: list[schemas.SelectedPackRead],
    selected_industry_packs: list[schemas.SelectedPackRead],
) -> str:
    reasons: list[str] = []
    domain_pack_names = {
        pack.pack_id: pack.pack_name for pack in selected_domain_packs
    }
    industry_pack_names = {
        pack.pack_id: pack.pack_name for pack in selected_industry_packs
    }

    if agent.agent_id == "host_agent":
        reasons.append("Host 是唯一 orchestration center，因此本輪固定存在。")
    if agent.agent_id in explicit_agent_ids:
        reasons.append("由使用者或上游流程明確指定覆寫。")
    if capability in agent.supported_capabilities:
        reasons.append(f"對齊 Capability Archetype：{capability.value}。")

    matched_domain_packs = [
        domain_pack_names[pack_id]
        for pack_id in agent.relevant_domain_packs
        if pack_id in domain_pack_names
    ]
    if matched_domain_packs:
        reasons.append(f"對齊 Domain Packs：{join_natural_list(matched_domain_packs)}。")

    matched_industry_packs = [
        industry_pack_names[pack_id]
        for pack_id in agent.relevant_industry_packs
        if pack_id in industry_pack_names
    ]
    if matched_industry_packs:
        reasons.append(f"對齊 Industry Packs：{join_natural_list(matched_industry_packs)}。")

    if not reasons:
        reasons.append("由 Agent Resolver 依目前 DecisionContext 與 readiness 推定。")
    return " ".join(reasons)


def _serialize_selected_agent(
    *,
    agent: AgentSpec,
    capability: CapabilityArchetype,
    explicit_agent_ids: list[str],
    selected_domain_packs: list[schemas.SelectedPackRead],
    selected_industry_packs: list[schemas.SelectedPackRead],
) -> schemas.SelectedAgentRead:
    return schemas.SelectedAgentRead(
        agent_id=agent.agent_id,
        agent_name=agent.agent_name,
        agent_type=agent.agent_type.value,
        description=agent.description,
        supported_capabilities=[item.value for item in agent.supported_capabilities],
        relevant_domain_packs=agent.relevant_domain_packs,
        relevant_industry_packs=agent.relevant_industry_packs,
        reason=_build_selected_agent_reason(
            agent=agent,
            capability=capability,
            explicit_agent_ids=explicit_agent_ids,
            selected_domain_packs=selected_domain_packs,
            selected_industry_packs=selected_industry_packs,
        ),
        runtime_binding=resolve_runtime_agent_binding(agent.agent_id),
        status=agent.status.value,
        version=agent.version,
    )


def resolve_agent_selection_for_task(
    task: models.Task,
    decision_context: schemas.DecisionContextRead | None,
    domain_lenses: list[str],
    pack_resolution: schemas.PackResolutionRead,
    input_entry_mode: InputEntryMode,
    deliverable_class_hint: DeliverableClass,
    external_research_heavy_candidate: bool,
    source_materials: list[schemas.SourceMaterialRead],
    artifacts: list[schemas.ArtifactRead],
) -> schemas.AgentSelectionRead:
    capability = _infer_capability_for_task(task, decision_context, domain_lenses)
    explicit_agent_ids = _extract_explicit_agent_overrides(task.constraints)
    decision_context_clear = bool(
        decision_context
        and decision_context.summary.strip()
        and decision_context.judgment_to_make.strip()
    )
    resolution = AGENT_RESOLVER.resolve(
        AgentResolverInput(
            capability=capability,
            selected_domain_pack_ids=[item.pack_id for item in pack_resolution.selected_domain_packs],
            selected_industry_pack_ids=[item.pack_id for item in pack_resolution.selected_industry_packs],
            decision_context_summary=(
                decision_context.summary if decision_context else task.description or task.title
            ),
            explicit_agent_ids=explicit_agent_ids,
            evidence_count=len(_usable_evidence(task)),
            artifact_count=len(_meaningful_artifacts(artifacts)),
            input_entry_mode=input_entry_mode,
            deliverable_class=deliverable_class_hint,
            decision_context_clear=decision_context_clear,
            external_research_heavy_case=external_research_heavy_candidate,
            allow_specialists=_allow_specialists_for_selection(
                capability,
                input_entry_mode,
                external_research_heavy_candidate,
            ),
        )
    )

    selected_reasoning_agents = [
        _serialize_selected_agent(
            agent=agent,
            capability=capability,
            explicit_agent_ids=explicit_agent_ids,
            selected_domain_packs=pack_resolution.selected_domain_packs,
            selected_industry_packs=pack_resolution.selected_industry_packs,
        )
        for agent_id in resolution.reasoning_agent_ids
        if (agent := EXTENSION_REGISTRY.get_agent(agent_id)) is not None
    ]
    selected_specialist_agents = [
        _serialize_selected_agent(
            agent=agent,
            capability=capability,
            explicit_agent_ids=explicit_agent_ids,
            selected_domain_packs=pack_resolution.selected_domain_packs,
            selected_industry_packs=pack_resolution.selected_industry_packs,
        )
        for agent_id in resolution.specialist_agent_ids
        if (agent := EXTENSION_REGISTRY.get_agent(agent_id)) is not None
    ]
    host_agent = EXTENSION_REGISTRY.get_agent(resolution.host_agent_id)

    selected_agent_names = [
        *[item.agent_name for item in selected_reasoning_agents],
        *[item.agent_name for item in selected_specialist_agents],
    ]
    rationale = [
        f"本輪 capability archetype 已被解析為 {capability.value}。",
    ]
    if pack_resolution.selected_domain_packs:
        rationale.append(
            "Domain / Functional Packs："
            + join_natural_list([item.pack_name for item in pack_resolution.selected_domain_packs])
            + "。"
        )
    if pack_resolution.selected_industry_packs:
        rationale.append(
            "Industry Packs："
            + join_natural_list([item.pack_name for item in pack_resolution.selected_industry_packs])
            + "。"
        )
    if input_entry_mode == InputEntryMode.ONE_LINE_INQUIRY:
        rationale.append("目前屬於一句話 sparse input，因此 agent selection 會維持較保守的最小集合。")
    elif input_entry_mode == InputEntryMode.SINGLE_DOCUMENT_INTAKE:
        rationale.append("目前屬於單文件 intake，因此 agent selection 會偏向 document-centered review / synthesis。")
    else:
        rationale.append("目前屬於多材料案件，因此可啟用較完整的 reasoning 組合。")
    if external_research_heavy_candidate:
        rationale.append("這輪屬於 external-research-heavy sparse case，因此會優先保留外部研究與不確定性 framing 能力。")
    if deliverable_class_hint == DeliverableClass.EXPLORATORY_BRIEF:
        rationale.append("目前交付等級仍偏 exploratory，因此不會假裝已啟用完整 decision-action agent 組合。")
    if resolution.deferred_agent_notes:
        rationale.append("部分相關 agents 已被正式標記為 deferred，待後續補證或工作條件成熟後再啟用。")
    if resolution.escalation_notes:
        rationale.append("本輪也保留了 escalation notes，說明要如何升級到更完整的 agent 組合。")

    return schemas.AgentSelectionRead(
        host_agent=(
            _serialize_selected_agent(
                agent=host_agent,
                capability=capability,
                explicit_agent_ids=explicit_agent_ids,
                selected_domain_packs=pack_resolution.selected_domain_packs,
                selected_industry_packs=pack_resolution.selected_industry_packs,
            )
            if host_agent is not None
            else None
        ),
        selected_reasoning_agents=selected_reasoning_agents,
        selected_specialist_agents=selected_specialist_agents,
        selected_agent_ids=[item.agent_id for item in [*selected_reasoning_agents, *selected_specialist_agents]],
        selected_agent_names=selected_agent_names,
        override_agent_ids=resolution.override_agent_ids,
        resolver_notes=resolution.resolver_notes,
        rationale=rationale,
        omitted_agent_notes=resolution.omitted_agent_notes,
        deferred_agent_notes=resolution.deferred_agent_notes,
        escalation_notes=resolution.escalation_notes,
    )


def resolve_pack_selection_for_task(
    task: models.Task,
    client: schemas.ClientRead | None,
    engagement: schemas.EngagementRead | None,
    workstream: schemas.WorkstreamRead | None,
    decision_context: schemas.DecisionContextRead | None,
    domain_lenses: list[str],
) -> schemas.PackResolutionRead:
    explicit_pack_ids = _extract_explicit_pack_overrides(task.constraints)
    client_type = decision_context.client_type if decision_context else client.client_type if client else None
    client_stage = decision_context.client_stage if decision_context else client.client_stage if client else None
    industry_hints = _build_industry_hints(task, client, engagement, workstream, decision_context)
    resolution = PACK_RESOLVER.resolve(
        PackResolverInput(
            domain_lenses=domain_lenses,
            client_type=client_type,
            client_stage=client_stage,
            decision_context_summary=(
                decision_context.summary if decision_context else task.description or task.title
            ),
            explicit_pack_ids=explicit_pack_ids,
            industry_hints=industry_hints,
        )
    )
    selected_domain_packs = [
        _serialize_selected_pack(
            pack=pack,
            explicit_pack_ids=explicit_pack_ids,
            domain_lenses=domain_lenses,
            client_type=client_type,
            client_stage=client_stage,
            industry_hints=industry_hints,
        )
        for pack_id in resolution.selected_domain_pack_ids
        if (pack := EXTENSION_REGISTRY.get_pack(pack_id))
    ]
    selected_industry_packs = [
        _serialize_selected_pack(
            pack=pack,
            explicit_pack_ids=explicit_pack_ids,
            domain_lenses=domain_lenses,
            client_type=client_type,
            client_stage=client_stage,
            industry_hints=industry_hints,
        )
        for pack_id in resolution.selected_industry_pack_ids
        if (pack := EXTENSION_REGISTRY.get_pack(pack_id))
    ]

    return schemas.PackResolutionRead(
        selected_domain_packs=selected_domain_packs,
        selected_industry_packs=selected_industry_packs,
        override_pack_ids=resolution.override_pack_ids,
        conflicts=resolution.conflicts,
        stack_order=resolution.stack_order,
        resolver_notes=resolution.resolver_notes,
        evidence_expectations=_unique_preserve_order(
            [
                *[
                    item
                    for pack in selected_domain_packs
                    for item in pack.evidence_expectations
                ],
                *[
                    item
                    for pack in selected_industry_packs
                    for item in pack.evidence_expectations
                ],
            ]
        ),
        key_kpis_or_operating_signals=_unique_preserve_order(
            [
                *[
                    item
                    for pack in selected_domain_packs
                    for item in (
                        pack.key_kpis_or_operating_signals or pack.key_kpis
                    )
                ],
                *[
                    item
                    for pack in selected_industry_packs
                    for item in (
                        pack.key_kpis_or_operating_signals or pack.key_kpis
                    )
                ],
            ]
        ),
        key_kpis=_unique_preserve_order(
            [
                *[
                    item
                    for pack in selected_domain_packs
                    for item in (
                        pack.key_kpis_or_operating_signals or pack.key_kpis
                    )
                ],
                *[
                    item
                    for pack in selected_industry_packs
                    for item in (
                        pack.key_kpis_or_operating_signals or pack.key_kpis
                    )
                ],
            ]
        ),
        common_risks=_unique_preserve_order(
            [
                *[
                    item
                    for pack in selected_domain_packs
                    for item in pack.common_risks
                ],
                *[
                    item
                    for pack in selected_industry_packs
                    for item in pack.common_risks
                ],
            ]
        ),
        decision_patterns=_unique_preserve_order(
            [
                *[
                    item
                    for pack in selected_domain_packs
                    for item in pack.decision_patterns
                ],
                *[
                    item
                    for pack in selected_industry_packs
                    for item in pack.decision_patterns
                ],
            ]
        ),
        deliverable_presets=_unique_preserve_order(
            [
                *[
                    item
                    for pack in selected_domain_packs
                    for item in pack.deliverable_presets
                ],
                *[
                    item
                    for pack in selected_industry_packs
                    for item in pack.deliverable_presets
                ],
            ]
        ),
    )


def _resolve_deliverable_class_hint(
    input_entry_mode: InputEntryMode,
    decision_context: schemas.DecisionContextRead | None,
    source_materials: list[schemas.SourceMaterialRead],
    artifacts: list[schemas.ArtifactRead],
    evidence_count: int,
    external_research_heavy_candidate: bool,
) -> DeliverableClass:
    if external_research_heavy_candidate or input_entry_mode == InputEntryMode.ONE_LINE_INQUIRY:
        return DeliverableClass.EXPLORATORY_BRIEF

    if input_entry_mode == InputEntryMode.SINGLE_DOCUMENT_INTAKE:
        return DeliverableClass.ASSESSMENT_REVIEW_MEMO

    if (
        decision_context
        and _normalize_whitespace(decision_context.judgment_to_make)
        and (
            _material_unit_count(source_materials, artifacts) >= 2
            or evidence_count >= 2
        )
    ):
        return DeliverableClass.DECISION_ACTION_DELIVERABLE

    return DeliverableClass.ASSESSMENT_REVIEW_MEMO


def _build_sparse_input_summary(
    input_entry_mode: InputEntryMode,
    deliverable_class_hint: DeliverableClass,
    external_research_heavy_candidate: bool,
) -> str:
    if external_research_heavy_candidate:
        return (
            "目前屬於高時效外部事件導向案例，內部資料仍偏稀疏；系統應先產出 exploratory brief，"
            "避免假裝已具備 company-specific certainty。"
        )
    if input_entry_mode == InputEntryMode.ONE_LINE_INQUIRY:
        return "目前屬於一句話問題進件；Host 會先建立 provisional world，再以 exploratory brief 形成第一輪判斷。"
    if input_entry_mode == InputEntryMode.SINGLE_DOCUMENT_INTAKE:
        return "目前屬於單文件進件；系統可先圍繞該 artifact 形成 assessment / review memo。"
    if deliverable_class_hint == DeliverableClass.DECISION_ACTION_DELIVERABLE:
        return "目前屬於多材料案件，資料密度已開始接近 decision / action deliverable 所需的工作鏈。"
    return "目前屬於多材料案件，但仍建議先以 assessment / review memo 收斂關鍵判斷。"


def _build_decision_context_read_from_row(
    decision_context: models.DecisionContext | None,
    *,
    goals: list[str],
    constraints: list[str],
    assumptions: list[str],
    reference_task_id: str | None = None,
) -> schemas.DecisionContextRead | None:
    if decision_context is None:
        return None

    row_goals = _string_list_from_summary_text(decision_context.goal_summary)
    row_constraints = _string_list_from_summary_text(decision_context.constraint_summary)
    row_assumptions = _string_list_from_summary_text(decision_context.assumption_summary)

    return schemas.DecisionContextRead(
        id=decision_context.id,
        task_id=(
            reference_task_id
            if decision_context.identity_scope == WORLD_AUTHORITY_IDENTITY_SCOPE and reference_task_id
            else decision_context.task_id
        ),
        matter_workspace_id=decision_context.matter_workspace_id,
        identity_scope=decision_context.identity_scope,
        task_reference_role=_task_reference_role_for_identity_scope(decision_context.identity_scope),
        client_id=decision_context.client_id,
        engagement_id=decision_context.engagement_id,
        workstream_id=decision_context.workstream_id,
        title=decision_context.title,
        summary=decision_context.summary,
        judgment_to_make=decision_context.judgment_to_make,
        domain_lenses=decision_context.domain_lenses,
        client_stage=decision_context.client_stage,
        client_type=decision_context.client_type,
        goals=row_goals or goals,
        constraints=row_constraints or constraints,
        assumptions=row_assumptions or assumptions,
        source_priority=decision_context.source_priority or "",
        external_data_policy=decision_context.external_data_policy or "",
        created_at=decision_context.created_at,
    )


def _build_decision_context_read(task: models.Task) -> schemas.DecisionContextRead | None:
    visible_constraints = [
        item.description
        for item in task.constraints
        if item.constraint_type != EXTERNAL_DATA_STRATEGY_CONSTRAINT_TYPE
    ]
    return _build_decision_context_read_from_row(
        _primary_item(task.decision_contexts),
        goals=[item.description for item in task.goals],
        constraints=visible_constraints,
        assumptions=_extract_assumptions(task),
        reference_task_id=task.id,
    )


def _coerce_string_list_from_payload(value: object) -> list[str]:
    if isinstance(value, str):
        return _split_multiline_items(value)
    if isinstance(value, list):
        return _normalize_string_list_payload(
            [str(item).strip() for item in value if str(item).strip()]
        )
    return []


def _string_list_from_summary_text(value: str | None) -> list[str]:
    return _split_multiline_items(value or "")


def _first_nonempty_text(*values: object) -> str:
    for value in values:
        normalized = _coerce_text(value)
        if normalized:
            return normalized
    return ""


def _first_nonempty_list(*values: list[str] | None) -> list[str]:
    for value in values:
        normalized = _normalize_string_list_payload(value)
        if normalized:
            return normalized
    return []


def _assumption_lines_from_payload(
    payload: list[dict] | None,
) -> list[str]:
    lines: list[str] = []
    for item in list(payload or []):
        if not isinstance(item, dict):
            continue
        detail = _normalize_whitespace(str(item.get("detail", "")))
        title = _normalize_whitespace(str(item.get("title", "")))
        if detail and title and detail != title:
            lines.append(f"{title}：{detail}")
        elif detail:
            lines.append(detail)
        elif title:
            lines.append(title)
    return _unique_preserve_order(lines)


def _build_world_decision_context_read(
    *,
    case_world_state: models.CaseWorldState | None,
    authority_decision_context: models.DecisionContext | None,
    fallback: schemas.DecisionContextRead | None,
) -> schemas.DecisionContextRead | None:
    payload = (
        case_world_state.decision_context_payload
        if case_world_state and isinstance(case_world_state.decision_context_payload, dict)
        else {}
    )
    if authority_decision_context is None and fallback is None and not payload:
        return None

    payload_goals = _coerce_string_list_from_payload(payload.get("goals"))
    payload_constraints = _coerce_string_list_from_payload(payload.get("constraints"))
    payload_assumptions = _coerce_string_list_from_payload(payload.get("assumptions"))
    state_assumptions = _assumption_lines_from_payload(
        case_world_state.assumptions_payload if case_world_state else None
    )
    authority_goals = _string_list_from_summary_text(
        authority_decision_context.goal_summary if authority_decision_context is not None else None
    )
    authority_constraints = _string_list_from_summary_text(
        authority_decision_context.constraint_summary if authority_decision_context is not None else None
    )
    authority_assumptions = _string_list_from_summary_text(
        authority_decision_context.assumption_summary if authority_decision_context is not None else None
    )

    return schemas.DecisionContextRead(
        id=(
            authority_decision_context.id
            if authority_decision_context is not None
            else fallback.id if fallback is not None else (case_world_state.decision_context_id if case_world_state else "")
        ),
        task_id=(
            case_world_state.latest_task_id
            if case_world_state and case_world_state.latest_task_id
            else authority_decision_context.task_id
            if authority_decision_context is not None
            else fallback.task_id if fallback is not None else ""
        ),
        matter_workspace_id=(
            authority_decision_context.matter_workspace_id
            if authority_decision_context is not None
            else fallback.matter_workspace_id if fallback is not None else (case_world_state.matter_workspace_id if case_world_state else None)
        ),
        identity_scope=WORLD_AUTHORITY_IDENTITY_SCOPE,
        task_reference_role="compatibility_only",
        client_id=(
            authority_decision_context.client_id
            if authority_decision_context is not None
            else fallback.client_id if fallback is not None else (case_world_state.client_id if case_world_state else None)
        ),
        engagement_id=(
            authority_decision_context.engagement_id
            if authority_decision_context is not None
            else fallback.engagement_id if fallback is not None else (case_world_state.engagement_id if case_world_state else None)
        ),
        workstream_id=(
            authority_decision_context.workstream_id
            if authority_decision_context is not None
            else fallback.workstream_id if fallback is not None else (case_world_state.workstream_id if case_world_state else None)
        ),
        title=(
            _coerce_text(payload.get("title"))
            or (authority_decision_context.title if authority_decision_context is not None else "")
            or (fallback.title if fallback is not None else "")
        ),
        summary=(
            _coerce_text(payload.get("summary"))
            or (authority_decision_context.summary if authority_decision_context is not None else "")
            or (fallback.summary if fallback is not None else "")
        ),
        judgment_to_make=(
            _coerce_text(payload.get("judgment_to_make"))
            or (authority_decision_context.judgment_to_make if authority_decision_context is not None else "")
            or (fallback.judgment_to_make if fallback is not None else "")
        ),
        domain_lenses=(
            _coerce_string_list_from_payload(payload.get("domain_lenses"))
            or (authority_decision_context.domain_lenses if authority_decision_context is not None else [])
            or (fallback.domain_lenses if fallback is not None else [])
        ),
        client_stage=(
            _coerce_text(payload.get("client_stage"))
            or (authority_decision_context.client_stage if authority_decision_context is not None else "")
            or (fallback.client_stage if fallback is not None else None)
        ),
        client_type=(
            _coerce_text(payload.get("client_type"))
            or (authority_decision_context.client_type if authority_decision_context is not None else "")
            or (fallback.client_type if fallback is not None else None)
        ),
        goals=payload_goals or authority_goals or (fallback.goals if fallback is not None else []),
        constraints=payload_constraints or authority_constraints or (fallback.constraints if fallback is not None else []),
        assumptions=(
            payload_assumptions
            or state_assumptions
            or authority_assumptions
            or (fallback.assumptions if fallback is not None else [])
        ),
        source_priority=(
            _coerce_text(payload.get("source_priority"))
            or (authority_decision_context.source_priority if authority_decision_context is not None else "")
            or (fallback.source_priority if fallback is not None else "")
        ),
        external_data_policy=(
            _coerce_text(payload.get("external_data_policy"))
            or (authority_decision_context.external_data_policy if authority_decision_context is not None else "")
            or (fallback.external_data_policy if fallback is not None else "")
        ),
        created_at=(
            authority_decision_context.created_at
            if authority_decision_context is not None
            else fallback.created_at if fallback is not None else (case_world_state.updated_at if case_world_state else models.utc_now())
        ),
    )


def _build_world_preferred_context_reads(
    task: models.Task,
    *,
    client: schemas.ClientRead | None,
    engagement: schemas.EngagementRead | None,
    workstream: schemas.WorkstreamRead | None,
    decision_context: schemas.DecisionContextRead | None,
    domain_lenses: list[str],
) -> tuple[
    schemas.ClientRead | None,
    schemas.EngagementRead | None,
    schemas.WorkstreamRead | None,
    schemas.DecisionContextRead | None,
    list[str],
]:
    db = object_session(task)
    matter_workspace = _get_linked_matter_workspace(task)
    if db is None or matter_workspace is None:
        return client, engagement, workstream, decision_context, domain_lenses

    authority_client_row, authority_engagement_row, authority_workstream_row, authority_decision_context_row = (
        _load_world_authority_spine(
            db,
            matter_workspace=matter_workspace,
            case_world_state=matter_workspace.case_world_state,
        )
    )
    world_client = (
        _build_client_read_from_row(authority_client_row, reference_task_id=task.id)
        if authority_client_row
        else client
    )
    world_engagement = (
        _build_engagement_read_from_row(authority_engagement_row, reference_task_id=task.id)
        if authority_engagement_row
        else engagement
    )
    world_workstream = (
        _build_workstream_read_from_row(authority_workstream_row, reference_task_id=task.id)
        if authority_workstream_row
        else workstream
    )
    world_decision_context = _build_world_decision_context_read(
        case_world_state=matter_workspace.case_world_state,
        authority_decision_context=authority_decision_context_row,
        fallback=decision_context,
    ) or decision_context
    preferred_domain_lenses = (
        world_decision_context.domain_lenses
        if world_decision_context and world_decision_context.domain_lenses
        else world_workstream.domain_lenses if world_workstream else domain_lenses
    )
    return (
        world_client,
        world_engagement,
        world_workstream,
        world_decision_context,
        preferred_domain_lenses,
    )


def _merge_world_decision_context_payload(
    *,
    canonical_payload: dict[str, object] | None,
    canonical_row: models.DecisionContext | None,
    overlay: schemas.DecisionContextRead | None,
    task: models.Task,
) -> dict[str, object]:
    canonical_payload = canonical_payload or {}
    payload_goals = _coerce_string_list_from_payload(canonical_payload.get("goals"))
    payload_constraints = _coerce_string_list_from_payload(canonical_payload.get("constraints"))
    payload_assumptions = _coerce_string_list_from_payload(canonical_payload.get("assumptions"))
    canonical_row_goals = _string_list_from_summary_text(
        canonical_row.goal_summary if canonical_row is not None else None
    )
    canonical_row_constraints = _string_list_from_summary_text(
        canonical_row.constraint_summary if canonical_row is not None else None
    )
    canonical_row_assumptions = _string_list_from_summary_text(
        canonical_row.assumption_summary if canonical_row is not None else None
    )
    task_constraints = [
        item.description
        for item in task.constraints
        if item.constraint_type != EXTERNAL_DATA_STRATEGY_CONSTRAINT_TYPE
    ]

    merged_goals = _first_nonempty_list(
        payload_goals,
        canonical_row_goals,
        overlay.goals if overlay else None,
        [item.description for item in task.goals],
    )
    merged_constraints = _first_nonempty_list(
        payload_constraints,
        canonical_row_constraints,
        overlay.constraints if overlay else None,
        task_constraints,
    )
    merged_assumptions = _first_nonempty_list(
        payload_assumptions,
        canonical_row_assumptions,
        overlay.assumptions if overlay else None,
        _extract_assumptions(task),
    )

    return {
        "title": _first_nonempty_text(
            canonical_payload.get("title"),
            canonical_row.title if canonical_row is not None else None,
            overlay.title if overlay else None,
            task.title,
        ),
        "summary": _first_nonempty_text(
            canonical_payload.get("summary"),
            canonical_row.summary if canonical_row is not None else None,
            overlay.summary if overlay else None,
            task.description,
        ),
        "judgment_to_make": _first_nonempty_text(
            canonical_payload.get("judgment_to_make"),
            canonical_row.judgment_to_make if canonical_row is not None else None,
            overlay.judgment_to_make if overlay else None,
            task.description,
            task.title,
        ),
        "domain_lenses": _first_nonempty_list(
            _coerce_string_list_from_payload(canonical_payload.get("domain_lenses")),
            canonical_row.domain_lenses if canonical_row is not None else None,
            overlay.domain_lenses if overlay else None,
        ),
        "client_stage": _first_nonempty_text(
            canonical_payload.get("client_stage"),
            canonical_row.client_stage if canonical_row is not None else None,
            overlay.client_stage if overlay else None,
        ),
        "client_type": _first_nonempty_text(
            canonical_payload.get("client_type"),
            canonical_row.client_type if canonical_row is not None else None,
            overlay.client_type if overlay else None,
        ),
        "goals": merged_goals,
        "constraints": merged_constraints,
        "assumptions": merged_assumptions,
        "source_priority": _first_nonempty_text(
            canonical_payload.get("source_priority"),
            canonical_row.source_priority if canonical_row is not None else None,
            overlay.source_priority if overlay else None,
        ),
        "external_data_policy": _first_nonempty_text(
            canonical_payload.get("external_data_policy"),
            canonical_row.external_data_policy if canonical_row is not None else None,
            overlay.external_data_policy if overlay else None,
        ),
    }


def _build_slice_decision_context_overlay(
    *,
    overlay: schemas.DecisionContextRead | None,
    world: schemas.DecisionContextRead | None,
) -> schemas.DecisionContextDeltaRead | None:
    if overlay is None:
        return None

    world = world or schemas.DecisionContextRead(
        id=overlay.id,
        task_id=overlay.task_id,
        matter_workspace_id=overlay.matter_workspace_id,
        identity_scope=WORLD_AUTHORITY_IDENTITY_SCOPE,
        client_id=overlay.client_id,
        engagement_id=overlay.engagement_id,
        workstream_id=overlay.workstream_id,
        title="",
        summary="",
        judgment_to_make="",
        domain_lenses=[],
        client_stage=None,
        client_type=None,
        goals=[],
        constraints=[],
        assumptions=[],
        source_priority="",
        external_data_policy="",
        created_at=overlay.created_at,
    )

    def _changed_text(value: str | None, world_value: str | None) -> str | None:
        normalized_value = _normalize_whitespace(value)
        if not normalized_value:
            return None
        normalized_world_value = _normalize_whitespace(world_value)
        return normalized_value if normalized_value != normalized_world_value else None

    def _changed_list(value: list[str], world_value: list[str]) -> list[str]:
        normalized_value = _normalize_string_list_payload(value)
        if not normalized_value:
            return []
        normalized_world_value = _normalize_string_list_payload(world_value)
        return normalized_value if normalized_value != normalized_world_value else []

    text_fields = {
        "title": _changed_text(overlay.title, world.title),
        "summary": _changed_text(overlay.summary, world.summary),
        "judgment_to_make": _changed_text(
            overlay.judgment_to_make,
            world.judgment_to_make,
        ),
        "client_stage": _changed_text(overlay.client_stage, world.client_stage),
        "client_type": _changed_text(overlay.client_type, world.client_type),
        "source_priority": _changed_text(overlay.source_priority, world.source_priority),
        "external_data_policy": _changed_text(
            overlay.external_data_policy,
            world.external_data_policy,
        ),
    }
    list_fields = {
        "domain_lenses": _changed_list(overlay.domain_lenses, world.domain_lenses),
        "goals": _changed_list(overlay.goals, world.goals),
        "constraints": _changed_list(overlay.constraints, world.constraints),
        "assumptions": _changed_list(overlay.assumptions, world.assumptions),
    }
    changed_fields = [
        name
        for name, value in {**text_fields, **list_fields}.items()
        if bool(value)
    ]
    has_meaningful_overlay = bool(changed_fields)
    if not has_meaningful_overlay:
        return None

    return schemas.DecisionContextDeltaRead(
        title=text_fields["title"],
        summary=text_fields["summary"],
        judgment_to_make=text_fields["judgment_to_make"],
        domain_lenses=list_fields["domain_lenses"],
        client_stage=text_fields["client_stage"],
        client_type=text_fields["client_type"],
        goals=list_fields["goals"],
        constraints=list_fields["constraints"],
        assumptions=list_fields["assumptions"],
        source_priority=text_fields["source_priority"],
        external_data_policy=text_fields["external_data_policy"],
        changed_fields=changed_fields,
    )


def _build_ontology_spine_for_task(
    task: models.Task,
) -> tuple[
    schemas.ClientRead | None,
    schemas.EngagementRead | None,
    schemas.WorkstreamRead | None,
    schemas.DecisionContextRead | None,
    list[str],
    list[schemas.SourceMaterialRead],
    list[schemas.ArtifactRead],
]:
    client = _primary_item(task.clients)
    engagement = _primary_item(task.engagements)
    workstream = _primary_item(task.workstreams)
    decision_context = _build_decision_context_read(task)
    domain_lenses = (
        decision_context.domain_lenses
        if decision_context and decision_context.domain_lenses
        else workstream.domain_lenses if workstream else [DEFAULT_DOMAIN_LENS]
    )
    source_materials = _build_source_materials(task)
    artifacts = _build_artifacts(task, source_materials)

    return (
        _build_client_read_from_row(client, reference_task_id=task.id) if client else None,
        _build_engagement_read_from_row(engagement, reference_task_id=task.id) if engagement else None,
        _build_workstream_read_from_row(workstream, reference_task_id=task.id) if workstream else None,
        decision_context,
        domain_lenses,
        source_materials,
        artifacts,
    )


def _build_world_preferred_ontology_spine_for_task(
    task: models.Task,
) -> tuple[
    schemas.ClientRead | None,
    schemas.EngagementRead | None,
    schemas.WorkstreamRead | None,
    schemas.DecisionContextRead | None,
    list[str],
    list[schemas.SourceMaterialRead],
    list[schemas.ArtifactRead],
]:
    client, engagement, workstream, decision_context, domain_lenses, source_materials, artifacts = (
        _build_ontology_spine_for_task(task)
    )
    world_client, world_engagement, world_workstream, world_decision_context, preferred_domain_lenses = (
        _build_world_preferred_context_reads(
            task,
            client=client,
            engagement=engagement,
            workstream=workstream,
            decision_context=decision_context,
            domain_lenses=domain_lenses,
        )
    )
    return (
        world_client,
        world_engagement,
        world_workstream,
        world_decision_context,
        preferred_domain_lenses,
        source_materials,
        artifacts,
    )


def _build_initial_material_snapshot(payload: schemas.TaskCreateRequest) -> InitialMaterialSnapshot:
    normalized_urls = [url.strip() for url in payload.initial_source_urls if url.strip()]
    return InitialMaterialSnapshot(
        url_count=len(normalized_urls),
        file_count=len(payload.initial_file_descriptors),
        pasted_text_present=bool(_normalize_whitespace(payload.initial_pasted_text)),
    )


def _build_provisional_task_for_case_world_seed(
    payload: schemas.TaskCreateRequest,
    *,
    background_brief: str,
    goal_description: str,
    inferred_constraints: list[schemas.ConstraintCreate],
) -> models.Task:
    provisional_task = models.Task(
        title=payload.title,
        description=payload.description,
        task_type=payload.task_type,
        mode=payload.mode.value,
        entry_preset=payload.entry_preset.value,
        status=TaskStatus.READY.value,
    )
    provisional_task.contexts = [
        models.TaskContext(
            task_id="provisional-task",
            summary=background_brief,
            assumptions=payload.assumptions,
            notes=payload.notes,
            version=1,
        )
    ]
    provisional_task.goals = [
        models.Goal(
            task_id="provisional-task",
            goal_type=payload.goal_type,
            description=goal_description,
            success_criteria=_infer_success_criteria(payload),
            priority="high",
        )
    ]
    provisional_task.constraints = [
        models.Constraint(
            task_id="provisional-task",
            constraint_type=item.constraint_type,
            description=item.description,
            severity=item.severity,
        )
        for item in [
            *inferred_constraints,
            _build_external_data_strategy_constraint(payload.external_data_strategy),
        ]
    ]
    if payload.subject_name:
        provisional_task.subjects = [
            models.Subject(
                task_id="provisional-task",
                subject_type=payload.subject_type,
                name=payload.subject_name,
                description=payload.subject_description,
                source_ref=None,
            )
        ]
    return provisional_task


def compile_case_world_seed_from_payload(
    payload: schemas.TaskCreateRequest,
) -> CompiledCaseWorldSeed:
    created_at = models.utc_now()
    background_brief = _infer_background_brief(payload)
    goal_description = _infer_goal_description(payload)
    success_criteria = _infer_success_criteria(payload)
    inferred_constraints = _infer_constraints(payload)
    material_snapshot = _build_initial_material_snapshot(payload)
    input_entry_mode = (
        InputEntryMode.MULTI_MATERIAL_CASE
        if material_snapshot.total_count >= 2
        else InputEntryMode.SINGLE_DOCUMENT_INTAKE
        if material_snapshot.total_count == 1
        else InputEntryMode.ONE_LINE_INQUIRY
    )
    client_name = _infer_client_name(payload)
    client_type = _infer_client_type(payload)
    client_stage = _infer_client_stage(payload)
    domain_lenses = _infer_domain_lenses(payload)
    engagement_name = _infer_engagement_name(payload)
    workstream_name = _infer_workstream_name(payload)
    decision_title = _infer_decision_title(payload)
    constraint_descriptions = [item.description for item in inferred_constraints]
    decision_summary = _normalize_whitespace(payload.decision_summary) or (
        f"本輪要圍繞「{payload.description or payload.title}」形成可採用的顧問判斷。"
    )
    judgment_to_make = _normalize_whitespace(payload.judgment_to_make) or goal_description
    identity = _derive_matter_workspace_identity_from_values(
        task_title=payload.title,
        task_type=payload.task_type,
        client_name=client_name,
        engagement_name=engagement_name,
        workstream_name=workstream_name,
        client_type=client_type,
        client_stage=client_stage,
        domain_lenses=domain_lenses,
        decision_title=decision_title,
        decision_summary=decision_summary,
    )
    client = schemas.ClientRead(
        id="provisional-client",
        task_id="provisional-task",
        name=client_name,
        client_type=client_type,
        client_stage=client_stage,
        description=payload.client_description,
        created_at=created_at,
    )
    engagement = schemas.EngagementRead(
        id="provisional-engagement",
        task_id="provisional-task",
        client_id=client.id,
        name=engagement_name,
        description=payload.engagement_description or payload.description or None,
        created_at=created_at,
    )
    workstream = schemas.WorkstreamRead(
        id="provisional-workstream",
        task_id="provisional-task",
        engagement_id=engagement.id,
        name=workstream_name,
        description=payload.workstream_description or payload.subject_description or None,
        domain_lenses=domain_lenses,
        created_at=created_at,
    )
    decision_context = schemas.DecisionContextRead(
        id="provisional-decision-context",
        task_id="provisional-task",
        client_id=client.id,
        engagement_id=engagement.id,
        workstream_id=workstream.id,
        title=decision_title,
        summary=decision_summary,
        judgment_to_make=judgment_to_make,
        domain_lenses=domain_lenses,
        client_stage=client_stage,
        client_type=client_type,
        goals=[goal_description],
        constraints=constraint_descriptions,
        assumptions=_split_multiline_items(payload.assumptions),
        source_priority=_build_source_priority_from_payload(payload),
        external_data_policy=_build_external_data_policy_text(payload.external_data_strategy),
        created_at=created_at,
    )
    provisional_task = _build_provisional_task_for_case_world_seed(
        payload,
        background_brief=background_brief,
        goal_description=goal_description,
        inferred_constraints=inferred_constraints,
    )
    pack_resolution = resolve_pack_selection_for_task(
        provisional_task,
        client,
        engagement,
        workstream,
        decision_context,
        domain_lenses,
    )
    external_research_heavy_candidate = (
        payload.external_data_strategy != ExternalDataStrategy.STRICT
        and input_entry_mode == InputEntryMode.ONE_LINE_INQUIRY
        and _contains_any_keyword(
            " ".join(
                filter(
                    None,
                    [
                        payload.title,
                        payload.description,
                        decision_summary,
                        judgment_to_make,
                        goal_description,
                    ],
                )
            ),
            EXTERNAL_EVENT_KEYWORDS,
        )
    )
    deliverable_class_hint = _resolve_deliverable_class_hint(
        input_entry_mode,
        decision_context,
        [],
        [],
        0,
        external_research_heavy_candidate,
    )
    agent_selection = resolve_agent_selection_for_task(
        provisional_task,
        decision_context,
        domain_lenses,
        pack_resolution,
        input_entry_mode,
        deliverable_class_hint,
        external_research_heavy_candidate,
        [],
        [],
    )
    canonical_intake_summary = {
        "entry_preset": payload.entry_preset.value,
        "resolved_input_state": input_entry_mode.value,
        "problem_statement": judgment_to_make or payload.description or payload.title,
        "planned_url_count": material_snapshot.url_count,
        "planned_file_count": material_snapshot.file_count,
        "planned_pasted_text": material_snapshot.pasted_text_present,
        "planned_material_count": material_snapshot.total_count,
        "external_research_heavy_candidate": external_research_heavy_candidate,
    }
    facts: list[dict[str, str]] = [
        {"title": "Client", "detail": client.name, "source": "explicit"},
        {"title": "Engagement", "detail": engagement.name, "source": "explicit"},
        {"title": "Workstream", "detail": workstream.name, "source": "explicit"},
        {"title": "DecisionContext", "detail": judgment_to_make, "source": "explicit"},
        {"title": "DomainLens", "detail": join_natural_list(domain_lenses), "source": "inferred"},
    ]
    if material_snapshot.file_count:
        facts.append(
            {
                "title": "Planned file intake",
                "detail": f"建立後預計掛入 {material_snapshot.file_count} 份檔案。",
                "source": "entry_preset",
            }
        )
    if material_snapshot.url_count:
        facts.append(
            {
                "title": "Planned URL intake",
                "detail": f"建立後預計掛入 {material_snapshot.url_count} 筆網址。",
                "source": "entry_preset",
            }
        )
    if material_snapshot.pasted_text_present:
        facts.append(
            {
                "title": "Planned pasted text",
                "detail": "建立後預計掛入一段正式 pasted text。",
                "source": "entry_preset",
            }
        )
    if pack_resolution.selected_domain_packs:
        facts.append(
            {
                "title": "Domain Packs",
                "detail": join_natural_list([item.pack_name for item in pack_resolution.selected_domain_packs]),
                "source": "pack_resolver",
            }
        )
    if pack_resolution.selected_industry_packs:
        facts.append(
            {
                "title": "Industry Packs",
                "detail": join_natural_list([item.pack_name for item in pack_resolution.selected_industry_packs]),
                "source": "pack_resolver",
            }
        )
    assumptions_payload = [
        {"title": "既定假設", "detail": item, "source": "user_provided"}
        for item in _split_multiline_items(payload.assumptions)
    ]
    if client_stage == UNSPECIFIED_LABEL:
        assumptions_payload.append(
            {
                "title": "Client stage 尚未明確",
                "detail": "目前案件世界仍缺少更明確的 client stage，因此 stage-specific heuristics 只能先保守套用。",
                "source": "host_inference",
            }
        )
    evidence_gaps_payload: list[dict[str, object]] = []
    if not _normalize_whitespace(judgment_to_make):
        evidence_gaps_payload.append(
            {
                "gap_key": "decision-context",
                "title": "DecisionContext 仍不夠穩定",
                "description": "這輪仍缺清楚的 judgment_to_make，因此案件世界只能先以 provisional framing 運作。",
                "priority": "high",
                "related_pack_ids": [],
            }
        )
    if material_snapshot.total_count == 0:
        evidence_gaps_payload.append(
            {
                "gap_key": "source-materials",
                "title": "尚缺正式來源材料",
                "description": "目前仍沒有規劃中的檔案、網址或 pasted text，建立後建議先補材料再推進分析。",
                "priority": "high",
                "related_pack_ids": [],
            }
        )
    if material_snapshot.total_count < 2:
        evidence_gaps_payload.append(
            {
                "gap_key": "evidence-thickness",
                "title": "世界證據基礎仍偏薄",
                "description": "目前這個案件世界還是稀疏輸入，建立後應先補件或先以 exploratory / assessment 交付收斂。",
                "priority": "high" if material_snapshot.total_count == 0 else "medium",
                "related_pack_ids": [],
            }
        )
    if external_research_heavy_candidate:
        evidence_gaps_payload.append(
            {
                "gap_key": "research-needed",
                "title": "需要正式外部補完",
                "description": "這輪屬於 external-research-heavy sparse case，Host 應先補 public research，再回到 evidence chain。",
                "priority": "high",
                "related_pack_ids": [
                    *[item.pack_id for item in pack_resolution.selected_domain_packs],
                    *[item.pack_id for item in pack_resolution.selected_industry_packs],
                ],
            }
        )
    for pack in [*pack_resolution.selected_domain_packs, *pack_resolution.selected_industry_packs]:
        if not pack.evidence_expectations:
            continue
        evidence_gaps_payload.append(
            {
                "gap_key": f"pack-{pack.pack_id}-expectation",
                "title": f"{pack.pack_name} 缺少關鍵支撐",
                "description": f"{pack.pack_name} 目前期待 {pack.evidence_expectations[0]} 等材料，但案件世界尚未形成足夠支撐。",
                "priority": "medium",
                "related_pack_ids": [pack.pack_id],
            }
        )
    next_best_actions: list[str] = [
        "先形成案件世界，再把這輪工作掛成第一個 task slice。",
    ]
    if material_snapshot.total_count > 0:
        next_best_actions.append("建立後先把目前材料掛回同一個案件世界，再檢查 evidence readiness。")
    else:
        next_best_actions.append("建立後可先補檔案、網址或 pasted text，或直接讓 Host 先做 exploratory framing。")
    if external_research_heavy_candidate:
        next_best_actions.append("允許 Host 先補正式外部 research，再回到 evidence chain。")
    if payload.engagement_continuity_mode != EngagementContinuityMode.ONE_OFF:
        next_best_actions.append("這輪結果應保留在同一個案件世界內，作為後續 follow-up 的 writeback 基底。")
    task_interpretation = {
        "task_title": payload.title,
        "task_type": payload.task_type,
        "workflow_mode": payload.mode.value,
        "capability": _infer_capability_for_task(provisional_task, decision_context, domain_lenses).value,
        "deliverable_class_hint": deliverable_class_hint.value,
        "external_data_strategy": payload.external_data_strategy.value,
        "summary": _build_sparse_input_summary(
            input_entry_mode,
            deliverable_class_hint,
            external_research_heavy_candidate,
        ),
    }
    return CompiledCaseWorldSeed(
        identity=identity,
        client=client,
        engagement=engagement,
        workstream=workstream,
        decision_context=decision_context,
        domain_lenses=domain_lenses,
        background_brief=background_brief,
        goal_description=goal_description,
        success_criteria=success_criteria,
        inferred_constraints=inferred_constraints,
        input_entry_mode=input_entry_mode,
        external_research_heavy_candidate=external_research_heavy_candidate,
        deliverable_class_hint=deliverable_class_hint,
        pack_resolution=pack_resolution,
        agent_selection=agent_selection,
        canonical_intake_summary=canonical_intake_summary,
        task_interpretation=task_interpretation,
        extracted_objects=[
            {"object_type": "client", "object_id": client.id, "label": client.name},
            {"object_type": "engagement", "object_id": engagement.id, "label": engagement.name},
            {"object_type": "workstream", "object_id": workstream.id, "label": workstream.name},
            {"object_type": "decision_context", "object_id": decision_context.id, "label": decision_context.title},
        ],
        inferred_links=[
            {"from": "case_world", "to": "task_slice", "relation": "will_frame"},
            {"from": "task_slice", "to": "decision_context", "relation": "frames"},
            {"from": "decision_context", "to": "source_material", "relation": "expects_evidence"},
        ],
        facts=facts,
        assumptions_payload=assumptions_payload,
        evidence_gaps_payload=evidence_gaps_payload,
        next_best_actions=_unique_preserve_order(next_best_actions),
    )


def ensure_matter_workspace_for_seed(
    db: Session,
    *,
    seed: CompiledCaseWorldSeed,
    continuity_mode: EngagementContinuityMode,
    writeback_depth: WritebackDepth,
) -> tuple[models.MatterWorkspace, bool]:
    matter_workspace = db.scalars(
        select(models.MatterWorkspace).where(
            models.MatterWorkspace.matter_key == str(seed.identity["matter_key"])
        )
    ).one_or_none()
    changed = matter_workspace is None
    if matter_workspace is None:
        matter_workspace = models.MatterWorkspace(
            **seed.identity,
            engagement_continuity_mode=continuity_mode.value,
            writeback_depth=writeback_depth.value,
        )
        db.add(matter_workspace)
        db.flush()
        return matter_workspace, True

    for field, value in seed.identity.items():
        if field == "title" and matter_workspace.title_override_active:
            continue
        if value and getattr(matter_workspace, field) != value:
            setattr(matter_workspace, field, value)
            changed = True
    if matter_workspace.engagement_continuity_mode != continuity_mode.value:
        matter_workspace.engagement_continuity_mode = continuity_mode.value
        changed = True
    if matter_workspace.writeback_depth != writeback_depth.value:
        matter_workspace.writeback_depth = writeback_depth.value
        changed = True
    if changed:
        db.add(matter_workspace)
        db.flush()
    return matter_workspace, changed


def ensure_matter_workspace_for_task(
    db: Session,
    task: models.Task,
    client: schemas.ClientRead | None,
    engagement: schemas.EngagementRead | None,
    workstream: schemas.WorkstreamRead | None,
    decision_context: schemas.DecisionContextRead | None,
    domain_lenses: list[str],
    continuity_mode: EngagementContinuityMode | None = None,
    writeback_depth: WritebackDepth | None = None,
) -> tuple[models.MatterWorkspace, bool]:
    identity = _derive_matter_workspace_identity(
        task,
        client,
        engagement,
        workstream,
        decision_context,
        domain_lenses,
    )
    matter_workspace = db.scalars(
        select(models.MatterWorkspace).where(
            models.MatterWorkspace.matter_key == str(identity["matter_key"])
        )
    ).one_or_none()
    changed = False
    try:
        fallback_entry_preset = InputEntryMode(task.entry_preset)
    except ValueError:
        fallback_entry_preset = InputEntryMode.ONE_LINE_INQUIRY

    if matter_workspace is None:
        resolved_continuity_mode = _normalize_continuity_mode(
            continuity_mode,
            fallback=fallback_entry_preset,
        )
        resolved_writeback_depth = _normalize_writeback_depth(
            writeback_depth,
            fallback=fallback_entry_preset,
        )
        matter_workspace = models.MatterWorkspace(
            **identity,
            engagement_continuity_mode=resolved_continuity_mode.value,
            writeback_depth=resolved_writeback_depth.value,
        )
        db.add(matter_workspace)
        db.flush()
        changed = True
    else:
        if not matter_workspace.title_override_active:
            incoming_title = identity["title"]
            if incoming_title and matter_workspace.title != incoming_title:
                matter_workspace.title = incoming_title
                changed = True

        for field in (
            "client_name",
            "engagement_name",
            "workstream_name",
            "client_type",
            "client_stage",
            "current_decision_context_title",
            "current_decision_context_summary",
        ):
            incoming = identity[field]
            if incoming and getattr(matter_workspace, field) != incoming:
                setattr(matter_workspace, field, incoming)
                changed = True
        merged_lenses = _unique_preserve_order(
            [*matter_workspace.domain_lenses, *list(identity["domain_lenses"])]
        )
        if merged_lenses != matter_workspace.domain_lenses:
            matter_workspace.domain_lenses = merged_lenses
            changed = True
        if not matter_workspace.engagement_continuity_mode:
            matter_workspace.engagement_continuity_mode = _normalize_continuity_mode(
                continuity_mode,
                fallback=fallback_entry_preset,
            ).value
            changed = True
        if not matter_workspace.writeback_depth:
            matter_workspace.writeback_depth = _normalize_writeback_depth(
                writeback_depth,
                fallback=fallback_entry_preset,
            ).value
            changed = True

    existing_link = next(
        (
            link
            for link in task.matter_workspace_links
            if link.matter_workspace_id == matter_workspace.id
        ),
        None,
    )
    if existing_link is None:
        db.add(
            models.MatterWorkspaceTaskLink(
                matter_workspace_id=matter_workspace.id,
                task_id=task.id,
            )
        )
        changed = True

    if changed:
        db.flush()

    return matter_workspace, changed


def _build_case_world_state_identity_payload(
    matter_workspace: models.MatterWorkspace,
) -> dict[str, object]:
    return {
        "matter_workspace_id": matter_workspace.id,
        "matter_key": matter_workspace.matter_key,
        "title": matter_workspace.title,
        "object_path": _build_matter_workspace_object_path(
            matter_workspace.client_name,
            matter_workspace.engagement_name,
            matter_workspace.workstream_name,
        ),
        "client_name": matter_workspace.client_name,
        "engagement_name": matter_workspace.engagement_name,
        "workstream_name": matter_workspace.workstream_name,
        "client_type": matter_workspace.client_type,
        "client_stage": matter_workspace.client_stage,
        "domain_lenses": matter_workspace.domain_lenses or [DEFAULT_DOMAIN_LENS],
        "current_decision_context_title": matter_workspace.current_decision_context_title,
        "current_decision_context_summary": matter_workspace.current_decision_context_summary,
    }


def _ensure_case_world_state_row(
    db: Session,
    matter_workspace: models.MatterWorkspace,
) -> models.CaseWorldState:
    state = matter_workspace.case_world_state
    if state is not None:
        return state

    state = models.CaseWorldState(
        matter_workspace_id=matter_workspace.id,
        world_status=matter_workspace.status or "active",
        continuity_mode=matter_workspace.engagement_continuity_mode,
        writeback_depth=matter_workspace.writeback_depth,
        world_identity_payload=_build_case_world_state_identity_payload(matter_workspace),
    )
    db.add(state)
    db.flush()
    matter_workspace.case_world_state = state
    return state


def _load_world_authority_object(
    db: Session,
    *,
    model: type[models.Client] | type[models.Engagement] | type[models.Workstream] | type[models.DecisionContext],
    matter_workspace_id: str,
    object_id: str | None,
):
    candidate = db.get(model, object_id) if object_id else None
    if (
        candidate is not None
        and getattr(candidate, "matter_workspace_id", None) == matter_workspace_id
        and getattr(candidate, "identity_scope", LEGACY_TASK_SLICE_IDENTITY_SCOPE) == WORLD_AUTHORITY_IDENTITY_SCOPE
    ):
        return candidate

    return db.scalars(
        select(model)
        .where(model.matter_workspace_id == matter_workspace_id)
        .where(model.identity_scope == WORLD_AUTHORITY_IDENTITY_SCOPE)
        .order_by(model.created_at)
    ).first()


def _assign_if_changed(instance: object, attribute: str, value: object) -> bool:
    if getattr(instance, attribute) == value:
        return False
    setattr(instance, attribute, value)
    return True


def _canonical_overlay_sync_value(
    current_value: object,
    incoming_value: object,
    *,
    allow_overlay_promotion: bool,
) -> object:
    if allow_overlay_promotion:
        return incoming_value
    if isinstance(current_value, list):
        return incoming_value if not current_value and incoming_value else current_value
    if current_value in (None, "", UNSPECIFIED_LABEL) and incoming_value not in (None, "", [], UNSPECIFIED_LABEL):
        return incoming_value
    return current_value


def ensure_world_context_spine_for_task(
    db: Session,
    *,
    matter_workspace: models.MatterWorkspace,
    task: models.Task,
    client: models.Client | None,
    engagement: models.Engagement | None,
    workstream: models.Workstream | None,
    decision_context: models.DecisionContext | None,
    case_world_state: models.CaseWorldState | None = None,
    refresh_compatibility_task_reference: bool = False,
) -> tuple[models.CaseWorldState, bool]:
    state = case_world_state or _ensure_case_world_state_row(db, matter_workspace)
    changed = False

    for item in (client, engagement, workstream, decision_context):
        if item is None:
            continue
        changed = _assign_if_changed(item, "matter_workspace_id", matter_workspace.id) or changed
        changed = _assign_if_changed(item, "identity_scope", SLICE_OVERLAY_IDENTITY_SCOPE) or changed

    canonical_client = _load_world_authority_object(
        db,
        model=models.Client,
        matter_workspace_id=matter_workspace.id,
        object_id=state.client_id,
    )
    if client is not None:
        if canonical_client is None:
            canonical_client = models.Client(
                task_id=task.id,
                matter_workspace_id=matter_workspace.id,
                identity_scope=WORLD_AUTHORITY_IDENTITY_SCOPE,
                name=client.name,
                client_type=client.client_type,
                client_stage=client.client_stage,
                description=client.description,
            )
            db.add(canonical_client)
            db.flush()
            changed = True
        else:
            changed = _assign_if_changed(canonical_client, "matter_workspace_id", matter_workspace.id) or changed
            changed = _assign_if_changed(
                canonical_client,
                "identity_scope",
                WORLD_AUTHORITY_IDENTITY_SCOPE,
            ) or changed
            changed = _assign_if_changed(
                canonical_client,
                "name",
                _canonical_overlay_sync_value(
                    canonical_client.name,
                    client.name,
                    allow_overlay_promotion=refresh_compatibility_task_reference,
                ),
            ) or changed
            changed = _assign_if_changed(
                canonical_client,
                "client_type",
                _canonical_overlay_sync_value(
                    canonical_client.client_type,
                    client.client_type,
                    allow_overlay_promotion=refresh_compatibility_task_reference,
                ),
            ) or changed
            changed = _assign_if_changed(
                canonical_client,
                "client_stage",
                _canonical_overlay_sync_value(
                    canonical_client.client_stage,
                    client.client_stage,
                    allow_overlay_promotion=refresh_compatibility_task_reference,
                ),
            ) or changed
            changed = _assign_if_changed(
                canonical_client,
                "description",
                _canonical_overlay_sync_value(
                    canonical_client.description,
                    client.description,
                    allow_overlay_promotion=refresh_compatibility_task_reference,
                ),
            ) or changed

    canonical_engagement = _load_world_authority_object(
        db,
        model=models.Engagement,
        matter_workspace_id=matter_workspace.id,
        object_id=state.engagement_id,
    )
    if engagement is not None:
        if canonical_engagement is None:
            canonical_engagement = models.Engagement(
                task_id=task.id,
                matter_workspace_id=matter_workspace.id,
                identity_scope=WORLD_AUTHORITY_IDENTITY_SCOPE,
                client_id=canonical_client.id if canonical_client else None,
                name=engagement.name,
                description=engagement.description,
            )
            db.add(canonical_engagement)
            db.flush()
            changed = True
        else:
            changed = _assign_if_changed(canonical_engagement, "matter_workspace_id", matter_workspace.id) or changed
            changed = _assign_if_changed(
                canonical_engagement,
                "identity_scope",
                WORLD_AUTHORITY_IDENTITY_SCOPE,
            ) or changed
            changed = _assign_if_changed(
                canonical_engagement,
                "client_id",
                canonical_client.id if canonical_client else None,
            ) or changed
            changed = _assign_if_changed(
                canonical_engagement,
                "name",
                _canonical_overlay_sync_value(
                    canonical_engagement.name,
                    engagement.name,
                    allow_overlay_promotion=refresh_compatibility_task_reference,
                ),
            ) or changed
            changed = _assign_if_changed(
                canonical_engagement,
                "description",
                _canonical_overlay_sync_value(
                    canonical_engagement.description,
                    engagement.description,
                    allow_overlay_promotion=refresh_compatibility_task_reference,
                ),
            ) or changed

    canonical_workstream = _load_world_authority_object(
        db,
        model=models.Workstream,
        matter_workspace_id=matter_workspace.id,
        object_id=state.workstream_id,
    )
    if workstream is not None:
        if canonical_workstream is None:
            canonical_workstream = models.Workstream(
                task_id=task.id,
                matter_workspace_id=matter_workspace.id,
                identity_scope=WORLD_AUTHORITY_IDENTITY_SCOPE,
                engagement_id=canonical_engagement.id if canonical_engagement else None,
                name=workstream.name,
                description=workstream.description,
                domain_lenses=workstream.domain_lenses,
            )
            db.add(canonical_workstream)
            db.flush()
            changed = True
        else:
            changed = _assign_if_changed(canonical_workstream, "matter_workspace_id", matter_workspace.id) or changed
            changed = _assign_if_changed(
                canonical_workstream,
                "identity_scope",
                WORLD_AUTHORITY_IDENTITY_SCOPE,
            ) or changed
            if refresh_compatibility_task_reference:
                changed = _assign_if_changed(canonical_workstream, "task_id", task.id) or changed
            changed = _assign_if_changed(
                canonical_workstream,
                "engagement_id",
                canonical_engagement.id if canonical_engagement else None,
            ) or changed
            changed = _assign_if_changed(
                canonical_workstream,
                "name",
                _canonical_overlay_sync_value(
                    canonical_workstream.name,
                    workstream.name,
                    allow_overlay_promotion=refresh_compatibility_task_reference,
                ),
            ) or changed
            changed = _assign_if_changed(
                canonical_workstream,
                "description",
                _canonical_overlay_sync_value(
                    canonical_workstream.description,
                    workstream.description,
                    allow_overlay_promotion=refresh_compatibility_task_reference,
                ),
            ) or changed
            changed = _assign_if_changed(
                canonical_workstream,
                "domain_lenses",
                _canonical_overlay_sync_value(
                    canonical_workstream.domain_lenses,
                    workstream.domain_lenses,
                    allow_overlay_promotion=refresh_compatibility_task_reference,
                ),
            ) or changed

    canonical_decision_context = _load_world_authority_object(
        db,
        model=models.DecisionContext,
        matter_workspace_id=matter_workspace.id,
        object_id=state.decision_context_id,
    )
    if decision_context is not None:
        overlay_decision_context = _build_decision_context_read(task)
        canonical_decision_context_payload = _merge_world_decision_context_payload(
            canonical_payload=(
                state.decision_context_payload
                if isinstance(state.decision_context_payload, dict)
                else None
            ),
            canonical_row=canonical_decision_context,
            overlay=overlay_decision_context,
            task=task,
        )
        canonical_goal_summary = "\n".join(canonical_decision_context_payload.get("goals") or [])
        canonical_constraint_summary = "\n".join(
            canonical_decision_context_payload.get("constraints") or []
        )
        canonical_assumption_summary = "\n".join(
            canonical_decision_context_payload.get("assumptions") or []
        )
        if canonical_decision_context is None:
            canonical_decision_context = models.DecisionContext(
                task_id=task.id,
                matter_workspace_id=matter_workspace.id,
                identity_scope=WORLD_AUTHORITY_IDENTITY_SCOPE,
                client_id=canonical_client.id if canonical_client else None,
                engagement_id=canonical_engagement.id if canonical_engagement else None,
                workstream_id=canonical_workstream.id if canonical_workstream else None,
                title=str(canonical_decision_context_payload.get("title") or decision_context.title),
                summary=str(canonical_decision_context_payload.get("summary") or decision_context.summary),
                judgment_to_make=str(
                    canonical_decision_context_payload.get("judgment_to_make")
                    or decision_context.judgment_to_make
                ),
                domain_lenses=list(canonical_decision_context_payload.get("domain_lenses") or decision_context.domain_lenses),
                client_stage=str(canonical_decision_context_payload.get("client_stage") or decision_context.client_stage),
                client_type=str(canonical_decision_context_payload.get("client_type") or decision_context.client_type),
                goal_summary=canonical_goal_summary or decision_context.goal_summary,
                constraint_summary=canonical_constraint_summary or decision_context.constraint_summary,
                assumption_summary=canonical_assumption_summary or decision_context.assumption_summary,
                source_priority=str(canonical_decision_context_payload.get("source_priority") or decision_context.source_priority or ""),
                external_data_policy=str(
                    canonical_decision_context_payload.get("external_data_policy")
                    or decision_context.external_data_policy
                    or ""
                ),
            )
            db.add(canonical_decision_context)
            db.flush()
            changed = True
        else:
            changed = _assign_if_changed(
                canonical_decision_context,
                "matter_workspace_id",
                matter_workspace.id,
            ) or changed
            changed = _assign_if_changed(
                canonical_decision_context,
                "identity_scope",
                WORLD_AUTHORITY_IDENTITY_SCOPE,
            ) or changed
            if refresh_compatibility_task_reference:
                changed = _assign_if_changed(canonical_decision_context, "task_id", task.id) or changed
            changed = _assign_if_changed(
                canonical_decision_context,
                "client_id",
                canonical_client.id if canonical_client else None,
            ) or changed
            changed = _assign_if_changed(
                canonical_decision_context,
                "engagement_id",
                canonical_engagement.id if canonical_engagement else None,
            ) or changed
            changed = _assign_if_changed(
                canonical_decision_context,
                "workstream_id",
                canonical_workstream.id if canonical_workstream else None,
            ) or changed
            changed = _assign_if_changed(
                canonical_decision_context,
                "title",
                str(canonical_decision_context_payload.get("title") or decision_context.title),
            ) or changed
            changed = _assign_if_changed(
                canonical_decision_context,
                "summary",
                str(canonical_decision_context_payload.get("summary") or decision_context.summary),
            ) or changed
            changed = _assign_if_changed(
                canonical_decision_context,
                "judgment_to_make",
                str(
                    canonical_decision_context_payload.get("judgment_to_make")
                    or decision_context.judgment_to_make
                ),
            ) or changed
            changed = _assign_if_changed(
                canonical_decision_context,
                "domain_lenses",
                list(canonical_decision_context_payload.get("domain_lenses") or decision_context.domain_lenses),
            ) or changed
            changed = _assign_if_changed(
                canonical_decision_context,
                "client_stage",
                str(canonical_decision_context_payload.get("client_stage") or decision_context.client_stage),
            ) or changed
            changed = _assign_if_changed(
                canonical_decision_context,
                "client_type",
                str(canonical_decision_context_payload.get("client_type") or decision_context.client_type),
            ) or changed
            changed = _assign_if_changed(
                canonical_decision_context,
                "goal_summary",
                canonical_goal_summary or decision_context.goal_summary,
            ) or changed
            changed = _assign_if_changed(
                canonical_decision_context,
                "constraint_summary",
                canonical_constraint_summary or decision_context.constraint_summary,
            ) or changed
            changed = _assign_if_changed(
                canonical_decision_context,
                "assumption_summary",
                canonical_assumption_summary or decision_context.assumption_summary,
            ) or changed
            changed = _assign_if_changed(
                canonical_decision_context,
                "source_priority",
                str(
                    canonical_decision_context_payload.get("source_priority")
                    or decision_context.source_priority
                    or ""
                ),
            ) or changed
            changed = _assign_if_changed(
                canonical_decision_context,
                "external_data_policy",
                str(
                    canonical_decision_context_payload.get("external_data_policy")
                    or decision_context.external_data_policy
                    or ""
                ),
            ) or changed
        changed = _assign_if_changed(
            state,
            "decision_context_payload",
            canonical_decision_context_payload,
        ) or changed

    if engagement is not None:
        changed = _assign_if_changed(
            engagement,
            "client_id",
            canonical_client.id if canonical_client else engagement.client_id,
        ) or changed
    if workstream is not None:
        changed = _assign_if_changed(
            workstream,
            "engagement_id",
            canonical_engagement.id if canonical_engagement else workstream.engagement_id,
        ) or changed
    if decision_context is not None:
        changed = _assign_if_changed(
            decision_context,
            "client_id",
            canonical_client.id if canonical_client else decision_context.client_id,
        ) or changed
        changed = _assign_if_changed(
            decision_context,
            "engagement_id",
            canonical_engagement.id if canonical_engagement else decision_context.engagement_id,
        ) or changed
        changed = _assign_if_changed(
            decision_context,
            "workstream_id",
            canonical_workstream.id if canonical_workstream else decision_context.workstream_id,
        ) or changed

    changed = _assign_if_changed(state, "client_id", canonical_client.id if canonical_client else None) or changed
    changed = _assign_if_changed(
        state,
        "engagement_id",
        canonical_engagement.id if canonical_engagement else None,
    ) or changed
    changed = _assign_if_changed(
        state,
        "workstream_id",
        canonical_workstream.id if canonical_workstream else None,
    ) or changed
    changed = _assign_if_changed(
        state,
        "decision_context_id",
        canonical_decision_context.id if canonical_decision_context else None,
    ) or changed

    if changed:
        db.add(state)
        db.flush()
    return state, changed


def _load_world_authority_spine(
    db: Session,
    *,
    matter_workspace: models.MatterWorkspace,
    case_world_state: models.CaseWorldState | None,
) -> tuple[
    models.Client | None,
    models.Engagement | None,
    models.Workstream | None,
    models.DecisionContext | None,
]:
    state = case_world_state or matter_workspace.case_world_state
    client = _load_world_authority_object(
        db,
        model=models.Client,
        matter_workspace_id=matter_workspace.id,
        object_id=state.client_id if state else None,
    )
    engagement = _load_world_authority_object(
        db,
        model=models.Engagement,
        matter_workspace_id=matter_workspace.id,
        object_id=state.engagement_id if state else None,
    )
    workstream = _load_world_authority_object(
        db,
        model=models.Workstream,
        matter_workspace_id=matter_workspace.id,
        object_id=state.workstream_id if state else None,
    )
    decision_context = _load_world_authority_object(
        db,
        model=models.DecisionContext,
        matter_workspace_id=matter_workspace.id,
        object_id=state.decision_context_id if state else None,
    )
    return client, engagement, workstream, decision_context


def ensure_case_world_state_from_seed(
    db: Session,
    *,
    matter_workspace: models.MatterWorkspace,
    seed: CompiledCaseWorldSeed,
) -> tuple[models.CaseWorldState, bool]:
    state = _ensure_case_world_state_row(db, matter_workspace)
    changed = False

    def assign(attribute: str, value: object) -> None:
        nonlocal changed
        if getattr(state, attribute) != value:
            setattr(state, attribute, value)
            changed = True

    assign("compiler_status", "compiled")
    assign("world_status", matter_workspace.status or "active")
    assign("entry_preset", str(seed.canonical_intake_summary["entry_preset"]))
    assign("continuity_mode", matter_workspace.engagement_continuity_mode)
    assign("writeback_depth", matter_workspace.writeback_depth)
    assign("world_identity_payload", _build_case_world_state_identity_payload(matter_workspace))
    assign("canonical_intake_summary", seed.canonical_intake_summary)
    assign("decision_context_payload", seed.decision_context.model_dump(mode="json"))
    assign("extracted_objects", seed.extracted_objects)
    assign("inferred_links", seed.inferred_links)
    assign("facts", seed.facts)
    assign("assumptions_payload", seed.assumptions_payload)
    assign("evidence_gaps_payload", seed.evidence_gaps_payload)
    assign("selected_capabilities", [str(seed.task_interpretation["capability"])])
    assign(
        "selected_domain_packs",
        [item.pack_id for item in seed.pack_resolution.selected_domain_packs],
    )
    assign(
        "selected_industry_packs",
        [item.pack_id for item in seed.pack_resolution.selected_industry_packs],
    )
    assign("selected_agent_ids", seed.agent_selection.selected_agent_ids)
    assign("suggested_research_need", seed.external_research_heavy_candidate)
    assign("next_best_actions", seed.next_best_actions)
    if changed:
        db.add(state)
        db.flush()
    return state, changed


def ensure_case_world_state_for_task(
    db: Session,
    *,
    matter_workspace: models.MatterWorkspace,
    task: models.Task,
    case_world_draft: models.CaseWorldDraft,
) -> tuple[models.CaseWorldState, bool]:
    state = _ensure_case_world_state_row(db, matter_workspace)
    changed = False
    active_task_ids = [link.task_id for link in matter_workspace.task_links]
    previous_summary = state.canonical_intake_summary or {}
    actual_source_material_count = int((case_world_draft.canonical_intake_summary or {}).get("source_material_count") or 0)
    planned_material_count = int(previous_summary.get("planned_material_count") or 0)
    merged_canonical_summary = dict(case_world_draft.canonical_intake_summary or {})
    if any(
        key in previous_summary
        for key in ("planned_material_count", "planned_url_count", "planned_file_count", "planned_pasted_text")
    ):
        merged_canonical_summary["planned_material_count"] = planned_material_count
        merged_canonical_summary["planned_url_count"] = int(previous_summary.get("planned_url_count") or 0)
        merged_canonical_summary["planned_file_count"] = int(previous_summary.get("planned_file_count") or 0)
        merged_canonical_summary["planned_pasted_text"] = bool(previous_summary.get("planned_pasted_text"))
    if actual_source_material_count == 0 and planned_material_count > 0:
        merged_canonical_summary["resolved_input_state"] = previous_summary.get(
            "resolved_input_state",
            merged_canonical_summary.get("resolved_input_state"),
        )
    merged_facts = list(case_world_draft.facts or [])
    if actual_source_material_count == 0 and planned_material_count > 0:
        for item in list(state.facts or []):
            if isinstance(item, dict) and item.get("title", "").startswith("Planned "):
                merged_facts.append(item)
    deduped_facts: list[dict] = []
    seen_fact_keys: set[str] = set()
    for item in merged_facts:
        if not isinstance(item, dict):
            continue
        fact_key = f"{item.get('title')}::{item.get('detail')}"
        if fact_key in seen_fact_keys:
            continue
        seen_fact_keys.add(fact_key)
        deduped_facts.append(item)
    merged_facts = deduped_facts
    merged_next_best_actions = list(case_world_draft.next_best_actions or [])
    if actual_source_material_count == 0 and planned_material_count > 0 and state.next_best_actions:
        merged_next_best_actions = _unique_preserve_order(
            [*state.next_best_actions, *merged_next_best_actions]
        )

    def assign(attribute: str, value: object) -> None:
        nonlocal changed
        if getattr(state, attribute) != value:
            setattr(state, attribute, value)
            changed = True

    assign("compiler_status", case_world_draft.compiler_status)
    assign("world_status", matter_workspace.status or "active")
    assign("entry_preset", case_world_draft.entry_preset)
    assign("continuity_mode", matter_workspace.engagement_continuity_mode)
    assign("writeback_depth", matter_workspace.writeback_depth)
    assign("world_identity_payload", _build_case_world_state_identity_payload(matter_workspace))
    assign("canonical_intake_summary", merged_canonical_summary)
    assign("decision_context_payload", case_world_draft.decision_context_payload)
    assign("extracted_objects", case_world_draft.extracted_objects)
    assign("inferred_links", case_world_draft.inferred_links)
    assign("facts", merged_facts)
    assign("assumptions_payload", case_world_draft.assumptions_payload)
    assign("evidence_gaps_payload", case_world_draft.evidence_gaps_payload)
    assign("selected_capabilities", case_world_draft.suggested_capabilities)
    assign("selected_domain_packs", case_world_draft.suggested_domain_packs)
    assign("selected_industry_packs", case_world_draft.suggested_industry_packs)
    assign("selected_agent_ids", case_world_draft.suggested_agents)
    assign("suggested_research_need", case_world_draft.suggested_research_need)
    assign("next_best_actions", merged_next_best_actions)
    assign("active_task_ids", active_task_ids)
    assign("latest_task_id", task.id)
    assign("latest_task_title", task.title)
    if changed:
        db.add(state)
        db.flush()
    return state, changed


def prepare_case_world_follow_up_for_task(
    db: Session,
    *,
    task: models.Task,
    follow_up_summary: str,
) -> models.CaseWorldState | None:
    summary = _normalize_whitespace(follow_up_summary)
    matter_workspace = _get_linked_matter_workspace(task)
    if matter_workspace is None or not summary:
        return None

    state = _ensure_case_world_state_row(db, matter_workspace)
    planned_material_count = int((state.canonical_intake_summary or {}).get("planned_material_count") or 0)
    initial_intake_backfill = (
        state.latest_task_id == task.id
        and not task.source_materials
        and not task.artifacts
        and planned_material_count > 0
        and (state.supplement_count or 0) == 0
    )
    if initial_intake_backfill:
        state.last_supplement_summary = summary
        state.compiler_status = "ingest_pending"
        state.next_best_actions = _unique_preserve_order(
            [
                f"建立時附帶材料正在掛回案件世界：{summary}",
                *state.next_best_actions,
            ]
        )[:8]
    else:
        state.supplement_count = (state.supplement_count or 0) + 1
        state.last_supplement_summary = summary
        state.compiler_status = "supplement_pending"
        state.next_best_actions = _unique_preserve_order(
            [
                f"已收到 follow-up supplement：{summary}",
                *state.next_best_actions,
            ]
        )[:8]
    db.add(state)
    db.flush()
    return state


def _default_matter_workspace_status(matter_workspace: models.MatterWorkspace, related_tasks: list[models.Task]) -> str:
    if matter_workspace.status:
        return matter_workspace.status

    active_task_count = sum(1 for task in related_tasks if task.status != TaskStatus.COMPLETED.value)
    return "active" if active_task_count > 0 else "paused"


def _get_linked_matter_workspace(task: models.Task) -> models.MatterWorkspace | None:
    for link in task.matter_workspace_links:
        if link.matter_workspace is not None:
            return link.matter_workspace
    return None


def resolve_continuity_policy_for_task(
    task: models.Task,
    matter_workspace: models.MatterWorkspace | None = None,
) -> tuple[EngagementContinuityMode, WritebackDepth]:
    linked_workspace = matter_workspace or _get_linked_matter_workspace(task)
    try:
        fallback_entry_preset = InputEntryMode(task.entry_preset)
    except ValueError:
        fallback_entry_preset = InputEntryMode.ONE_LINE_INQUIRY
    continuity_mode = _normalize_continuity_mode(
        linked_workspace.engagement_continuity_mode if linked_workspace else None,
        fallback=fallback_entry_preset,
    )
    writeback_depth = _normalize_writeback_depth(
        linked_workspace.writeback_depth if linked_workspace else None,
        fallback=fallback_entry_preset,
    )
    return continuity_mode, writeback_depth


def _default_deliverable_status(task: models.Task, deliverable: models.Deliverable) -> str:
    if deliverable.status:
        return deliverable.status
    if task.status == TaskStatus.DRAFT.value:
        return "draft"
    if task.status == TaskStatus.COMPLETED.value:
        return "final"
    return "pending_confirmation"


def _resolve_deliverable_summary_text(deliverable: models.Deliverable) -> str:
    if deliverable.summary.strip():
        return deliverable.summary.strip()

    content = deliverable.content_structure if isinstance(deliverable.content_structure, dict) else {}
    for key in ("executive_summary", "summary", "final_recommendation", "recommended_path"):
        value = content.get(key)
        if isinstance(value, str) and value.strip():
            return value.strip()

    return deliverable.title


def _normalize_multiline_text(value: str | None) -> str:
    if not isinstance(value, str):
        return ""
    return value.strip()


def _normalize_string_list_payload(values: list[str] | tuple[str, ...] | None) -> list[str]:
    if not isinstance(values, (list, tuple)):
        return []
    return [item.strip() for item in values if isinstance(item, str) and item.strip()]


def _read_json_sections(value: object) -> dict:
    return value if isinstance(value, dict) else {}


def _normalize_matter_content_sections(
    value: object,
) -> schemas.MatterWorkspaceContentSectionsRead:
    sections = _read_json_sections(value)
    return schemas.MatterWorkspaceContentSectionsRead(
        core_question=_normalize_multiline_text(sections.get("core_question")),
        analysis_focus=_normalize_multiline_text(sections.get("analysis_focus")),
        constraints_and_risks=_normalize_multiline_text(sections.get("constraints_and_risks")),
        next_steps=_normalize_multiline_text(sections.get("next_steps")),
    )


def _serialize_matter_content_sections(
    payload: schemas.MatterWorkspaceContentSectionsRequest,
) -> dict[str, str]:
    return {
        "core_question": payload.core_question.strip(),
        "analysis_focus": payload.analysis_focus.strip(),
        "constraints_and_risks": payload.constraints_and_risks.strip(),
        "next_steps": payload.next_steps.strip(),
    }


def _normalize_deliverable_content_sections(
    value: object,
) -> schemas.DeliverableContentSectionsRead:
    sections = _read_json_sections(value)
    return schemas.DeliverableContentSectionsRead(
        executive_summary=_normalize_multiline_text(sections.get("executive_summary")),
        recommendations=_normalize_string_list_payload(sections.get("recommendations")),
        risks=_normalize_string_list_payload(sections.get("risks")),
        action_items=_normalize_string_list_payload(sections.get("action_items")),
        evidence_basis=_normalize_string_list_payload(sections.get("evidence_basis")),
    )


def _serialize_deliverable_content_sections(
    payload: schemas.DeliverableContentSectionsRequest,
) -> dict[str, object]:
    return {
        "executive_summary": payload.executive_summary.strip(),
        "recommendations": _normalize_string_list_payload(payload.recommendations),
        "risks": _normalize_string_list_payload(payload.risks),
        "action_items": _normalize_string_list_payload(payload.action_items),
        "evidence_basis": _normalize_string_list_payload(payload.evidence_basis),
    }


def _serialize_content_revision_diff_summary(
    value: object,
) -> list[schemas.ContentRevisionDiffItemRead]:
    if not isinstance(value, list):
        return []

    items: list[schemas.ContentRevisionDiffItemRead] = []
    for entry in value:
        if not isinstance(entry, dict):
            continue
        items.append(
            schemas.ContentRevisionDiffItemRead(
                section_key=str(entry.get("section_key", "")).strip(),
                section_label=str(entry.get("section_label", "")).strip(),
                change_type=str(entry.get("change_type", "")).strip(),
                previous_preview=_normalize_multiline_text(entry.get("previous_preview")),
                current_preview=_normalize_multiline_text(entry.get("current_preview")),
            )
        )
    return items


def _build_matter_content_revision_read(
    revision: models.MatterContentRevision,
) -> schemas.MatterContentRevisionRead:
    return schemas.MatterContentRevisionRead(
        id=revision.id,
        matter_workspace_id=revision.matter_workspace_id,
        object_type="matter",
        object_id=revision.matter_workspace_id,
        source=revision.source,
        revision_summary=revision.revision_summary,
        changed_sections=revision.changed_sections or [],
        diff_summary=_serialize_content_revision_diff_summary(revision.diff_summary),
        snapshot=_normalize_matter_content_sections(revision.snapshot),
        rollback_target_revision_id=revision.rollback_target_revision_id,
        created_at=revision.created_at,
    )


def _build_deliverable_content_revision_read(
    revision: models.DeliverableContentRevision,
) -> schemas.DeliverableContentRevisionRead:
    return schemas.DeliverableContentRevisionRead(
        id=revision.id,
        deliverable_id=revision.deliverable_id,
        task_id=revision.task_id,
        object_type="deliverable",
        object_id=revision.deliverable_id,
        source=revision.source,
        revision_summary=revision.revision_summary,
        changed_sections=revision.changed_sections or [],
        diff_summary=_serialize_content_revision_diff_summary(revision.diff_summary),
        snapshot=_normalize_deliverable_content_sections(revision.snapshot),
        version_tag=revision.version_tag,
        deliverable_status=revision.deliverable_status,
        source_version_event_id=revision.source_version_event_id,
        rollback_target_revision_id=revision.rollback_target_revision_id,
        created_at=revision.created_at,
    )


def _load_tasks_for_matter_workspaces(
    db: Session,
    matter_workspace_ids: list[str],
) -> dict[str, list[models.Task]]:
    if not matter_workspace_ids:
        return {}

    links = db.scalars(
        select(models.MatterWorkspaceTaskLink).where(
            models.MatterWorkspaceTaskLink.matter_workspace_id.in_(matter_workspace_ids)
        )
    ).all()
    task_ids = _unique_preserve_order([link.task_id for link in links])
    if not task_ids:
        return {matter_workspace_id: [] for matter_workspace_id in matter_workspace_ids}

    tasks = db.scalars(
        select(models.Task)
        .options(*task_load_options())
        .where(models.Task.id.in_(task_ids))
    ).unique().all()
    task_by_id = {task.id: task for task in tasks}
    tasks_by_workspace: dict[str, list[models.Task]] = {
        matter_workspace_id: [] for matter_workspace_id in matter_workspace_ids
    }

    for link in links:
        task = task_by_id.get(link.task_id)
        if task is not None:
            tasks_by_workspace.setdefault(link.matter_workspace_id, []).append(task)

    for matter_workspace_id, related_tasks in tasks_by_workspace.items():
        tasks_by_workspace[matter_workspace_id] = sorted(
            related_tasks,
            key=lambda item: item.updated_at,
            reverse=True,
        )

    return tasks_by_workspace


def get_primary_task_for_matter(
    db: Session,
    matter_id: str,
) -> models.Task:
    matter_workspace = db.scalars(
        select(models.MatterWorkspace).where(models.MatterWorkspace.id == matter_id)
    ).one_or_none()
    if matter_workspace is None:
        raise HTTPException(status_code=404, detail="找不到指定案件。")

    related_tasks = _load_tasks_for_matter_workspaces(db, [matter_id]).get(matter_id, [])
    if not related_tasks:
        raise HTTPException(status_code=404, detail="這個案件目前還沒有可補件的任務。")

    active_task = next(
        (item for item in related_tasks if item.status in {TaskStatus.READY.value, TaskStatus.RUNNING.value}),
        None,
    )
    return active_task or related_tasks[0]


def _build_matter_workspace_summary_from_tasks(
    matter_workspace: models.MatterWorkspace,
    related_tasks: list[models.Task],
) -> schemas.MatterWorkspaceSummaryRead:
    latest_task = related_tasks[0] if related_tasks else None
    active_task_count = sum(1 for task in related_tasks if task.status != TaskStatus.COMPLETED.value)
    workspace_status = _default_matter_workspace_status(matter_workspace, related_tasks)
    deliverable_count = sum(len(task.deliverables) for task in related_tasks)
    artifact_ids: set[str] = set()
    source_material_ids: set[str] = set()
    shared_material_ids: set[str] = set()
    shared_evidence_ids: set[str] = set()
    selected_pack_names: list[str] = []
    selected_agent_names: list[str] = []
    current_decision_context_title = matter_workspace.current_decision_context_title
    current_decision_context_summary = matter_workspace.current_decision_context_summary

    for task in related_tasks:
        client, engagement, workstream, decision_context, domain_lenses, source_materials, artifacts = (
            _build_world_preferred_ontology_spine_for_task(task)
        )
        artifact_ids.update(item.id for item in _meaningful_artifacts(artifacts))
        source_material_ids.update(item.id for item in _meaningful_source_materials(source_materials))
        shared_material_ids.update(
            item.id
            for item in source_materials
            if item.continuity_scope == WORLD_SHARED_CONTINUITY_SCOPE
        )
        shared_evidence_ids.update(
            item.id
            for item in _serialize_evidence_items(task, source_materials, artifacts)
            if item.continuity_scope == WORLD_SHARED_CONTINUITY_SCOPE
        )

        if not current_decision_context_title and decision_context:
            current_decision_context_title = decision_context.title
            current_decision_context_summary = decision_context.summary

        pack_resolution = resolve_pack_selection_for_task(
            task,
            client,
            engagement,
            workstream,
            decision_context,
            domain_lenses,
        )
        input_entry_mode = _infer_input_entry_mode(task, source_materials, artifacts)
        external_research_heavy_candidate = _is_external_research_heavy_candidate(
            task,
            decision_context,
            source_materials,
            artifacts,
            input_entry_mode,
        )
        deliverable_class_hint = _resolve_deliverable_class_hint(
            input_entry_mode,
            decision_context,
            source_materials,
            artifacts,
            len(_usable_evidence(task)),
            external_research_heavy_candidate,
        )
        agent_selection = resolve_agent_selection_for_task(
            task,
            decision_context,
            domain_lenses,
            pack_resolution,
            input_entry_mode,
            deliverable_class_hint,
            external_research_heavy_candidate,
            source_materials,
            artifacts,
        )
        selected_pack_names.extend(
            [*pack_resolution.selected_domain_packs, *pack_resolution.selected_industry_packs]
        )
        selected_agent_names.extend(agent_selection.selected_agent_names)

    continuity_summary = (
        f"目前這個案件世界採 {matter_workspace.engagement_continuity_mode} / {matter_workspace.writeback_depth} 策略，"
        f"已有 {len(related_tasks)} 個相關 task slices、{deliverable_count} 份 deliverables，"
        f"可直接回看 decision trajectory、最近交付與工作鏈材料；其中 {len(shared_material_ids)} 份共享 materials / "
        f"{len(shared_evidence_ids)} 則共享 evidence 已正式掛在案件世界。"
        if len(related_tasks) > 1 or deliverable_count > 0
        else (
            f"這個案件世界目前仍在第一個 task slice 階段，採 {matter_workspace.engagement_continuity_mode} / "
            f"{matter_workspace.writeback_depth} 策略，但已具備正式的 matter / engagement 工作面。"
        )
    )
    active_work_summary = (
        f"目前有 {active_task_count} 個 active work item；最近更新來自「{latest_task.title}」。"
        if latest_task and active_task_count > 0
        else (
            f"最近一次完成的工作是「{latest_task.title}」。"
            if latest_task
            else "目前尚未有可顯示的案件工作。"
        )
    )

    return schemas.MatterWorkspaceSummaryRead(
        id=matter_workspace.id,
        title=matter_workspace.title,
        workspace_summary=(matter_workspace.summary or "").strip(),
        status=workspace_status,
        object_path=_build_matter_workspace_object_path(
            matter_workspace.client_name,
            matter_workspace.engagement_name,
            matter_workspace.workstream_name,
        ),
        client_name=matter_workspace.client_name,
        engagement_name=matter_workspace.engagement_name,
        workstream_name=matter_workspace.workstream_name,
        client_stage=(
            matter_workspace.client_stage
            if matter_workspace.client_stage != UNSPECIFIED_LABEL
            else None
        ),
        client_type=(
            matter_workspace.client_type
            if matter_workspace.client_type != UNSPECIFIED_LABEL
            else None
        ),
        domain_lenses=matter_workspace.domain_lenses or [DEFAULT_DOMAIN_LENS],
        current_decision_context_title=current_decision_context_title,
        current_decision_context_summary=current_decision_context_summary,
        total_task_count=len(related_tasks),
        active_task_count=active_task_count,
        deliverable_count=deliverable_count,
        artifact_count=len(artifact_ids),
        source_material_count=len(source_material_ids),
        latest_updated_at=(
            max(latest_task.updated_at, matter_workspace.updated_at)
            if latest_task
            else matter_workspace.updated_at
        ),
        continuity_summary=continuity_summary,
        active_work_summary=active_work_summary,
        engagement_continuity_mode=EngagementContinuityMode(
            matter_workspace.engagement_continuity_mode
        ),
        writeback_depth=WritebackDepth(matter_workspace.writeback_depth),
        selected_pack_names=_unique_preserve_order(
            [
                item.pack_name if isinstance(item, schemas.SelectedPackRead) else str(item)
                for item in selected_pack_names
            ]
        )[:6],
        selected_agent_names=_unique_preserve_order(selected_agent_names)[:6],
    )


def _build_matter_workspace_summary_map(
    db: Session,
    matter_workspaces: list[models.MatterWorkspace],
) -> dict[str, schemas.MatterWorkspaceSummaryRead]:
    tasks_by_workspace = _load_tasks_for_matter_workspaces(
        db,
        [item.id for item in matter_workspaces],
    )
    return {
        matter_workspace.id: _build_matter_workspace_summary_from_tasks(
            matter_workspace,
            tasks_by_workspace.get(matter_workspace.id, []),
        )
        for matter_workspace in matter_workspaces
    }


def _build_evidence_object_mappings(
    source_materials: list[schemas.SourceMaterialRead],
    artifacts: list[schemas.ArtifactRead],
) -> tuple[dict[str, str], dict[str, str]]:
    source_material_by_document = {
        item.source_document_id: item.id
        for item in source_materials
        if item.source_document_id
    }
    artifact_by_document = {
        item.source_document_id: item.id
        for item in artifacts
        if item.source_document_id
    }
    return source_material_by_document, artifact_by_document


def _serialize_evidence_items(
    task: models.Task,
    source_materials: list[schemas.SourceMaterialRead],
    artifacts: list[schemas.ArtifactRead],
) -> list[schemas.EvidenceRead]:
    source_material_by_document, artifact_by_document = _build_evidence_object_mappings(
        source_materials,
        artifacts,
    )
    evidence_rows = list(task.evidence)
    db, matter_workspace, matter_workspace_id = _ensure_source_chain_participation_snapshot_ready(task)
    participation_snapshot = _load_task_object_participation_snapshot(
        db,
        matter_workspace_id=matter_workspace_id,
        task_id=task.id,
    )
    if db is not None and matter_workspace is not None:
        world_evidence = db.scalars(
            select(models.Evidence)
            .where(models.Evidence.matter_workspace_id == matter_workspace.id)
            .where(models.Evidence.continuity_scope == WORLD_SHARED_CONTINUITY_SCOPE)
            .order_by(models.Evidence.created_at)
        ).all()
        covered_ids = {item.id for item in evidence_rows}
        for item in world_evidence:
            if item.id in covered_ids:
                continue
            evidence_rows.append(item)
            covered_ids.add(item.id)
    return [
        schemas.EvidenceRead(
            id=item.id,
            task_id=item.task_id,
            matter_workspace_id=item.matter_workspace_id,
            source_document_id=item.source_document_id,
            source_material_id=item.source_material_id
            or (source_material_by_document.get(item.source_document_id) if item.source_document_id else None),
            artifact_id=item.artifact_id
            or (artifact_by_document.get(item.source_document_id) if item.source_document_id else None),
            continuity_scope=item.continuity_scope,
            evidence_type=item.evidence_type,
            source_type=item.source_type,
            source_ref=item.source_ref,
            title=item.title,
            excerpt_or_summary=item.excerpt_or_summary,
            reliability_level=item.reliability_level,
            participation=_build_object_participation_read(
                object_type=OBJECT_TYPE_EVIDENCE,
                object_id=item.id,
                object_task_id=item.task_id,
                current_task_id=task.id,
                matter_workspace_id=item.matter_workspace_id,
                continuity_scope=item.continuity_scope,
                snapshot=participation_snapshot,
            ),
            created_at=item.created_at,
        )
        for item in evidence_rows
    ]


def _build_supporting_evidence_maps(
    task: models.Task,
) -> tuple[dict[str, list[str]], dict[str, list[str]], dict[str, list[str]]]:
    recommendation_map: dict[str, list[str]] = {}
    risk_map: dict[str, list[str]] = {}
    action_item_map: dict[str, list[str]] = {}

    for link in task.recommendation_evidence_links:
        recommendation_map.setdefault(link.recommendation_id, [])
        if link.evidence_id not in recommendation_map[link.recommendation_id]:
            recommendation_map[link.recommendation_id].append(link.evidence_id)

    for link in task.risk_evidence_links:
        risk_map.setdefault(link.risk_id, [])
        if link.evidence_id not in risk_map[link.risk_id]:
            risk_map[link.risk_id].append(link.evidence_id)

    for link in task.action_item_evidence_links:
        action_item_map.setdefault(link.action_item_id, [])
        if link.evidence_id not in action_item_map[link.action_item_id]:
            action_item_map[link.action_item_id].append(link.evidence_id)

    return recommendation_map, risk_map, action_item_map


def _serialize_risks(
    task: models.Task,
    supporting_map: dict[str, list[str]],
) -> list[schemas.RiskRead]:
    return [
        schemas.RiskRead(
            id=item.id,
            task_id=item.task_id,
            title=item.title,
            description=item.description,
            risk_type=item.risk_type,
            impact_level=item.impact_level,
            likelihood_level=item.likelihood_level,
            evidence_refs=item.evidence_refs,
            supporting_evidence_ids=supporting_map.get(item.id, []),
            created_at=item.created_at,
        )
        for item in task.risks
    ]


def _serialize_recommendations(
    task: models.Task,
    supporting_map: dict[str, list[str]],
) -> list[schemas.RecommendationRead]:
    return [
        schemas.RecommendationRead(
            id=item.id,
            task_id=item.task_id,
            summary=item.summary,
            rationale=item.rationale,
            based_on_refs=item.based_on_refs,
            supporting_evidence_ids=supporting_map.get(item.id, []),
            priority=item.priority,
            owner_suggestion=item.owner_suggestion,
            created_at=item.created_at,
        )
        for item in task.recommendations
    ]


def _serialize_action_items(
    task: models.Task,
    supporting_map: dict[str, list[str]],
) -> list[schemas.ActionItemRead]:
    return [
        schemas.ActionItemRead(
            id=item.id,
            task_id=item.task_id,
            description=item.description,
            suggested_owner=item.suggested_owner,
            priority=item.priority,
            due_hint=item.due_hint,
            dependency_refs=item.dependency_refs,
            supporting_evidence_ids=supporting_map.get(item.id, []),
            status=item.status,
            created_at=item.created_at,
        )
        for item in task.action_items
    ]


def _build_object_label_map(
    task: models.Task,
    source_materials: list[schemas.SourceMaterialRead],
    artifacts: list[schemas.ArtifactRead],
    evidence: list[schemas.EvidenceRead],
    recommendations: list[schemas.RecommendationRead],
    risks: list[schemas.RiskRead],
    action_items: list[schemas.ActionItemRead],
    decision_context: schemas.DecisionContextRead | None,
    client: schemas.ClientRead | None,
    engagement: schemas.EngagementRead | None,
    workstream: schemas.WorkstreamRead | None,
) -> dict[str, dict[str, str]]:
    return {
        "client": {client.id: client.name} if client else {},
        "engagement": {engagement.id: engagement.name} if engagement else {},
        "workstream": {workstream.id: workstream.name} if workstream else {},
        "decision_context": (
            {decision_context.id: decision_context.judgment_to_make or decision_context.title}
            if decision_context
            else {}
        ),
        "source_material": {item.id: item.title for item in source_materials},
        "artifact": {item.id: item.title for item in artifacts},
        "evidence": {item.id: item.title for item in evidence},
        "recommendation": {item.id: item.summary for item in recommendations},
        "risk": {item.id: item.title or item.description for item in risks},
        "action_item": {item.id: item.description for item in action_items},
    }


def _serialize_deliverables(
    task: models.Task,
    object_label_map: dict[str, dict[str, str]],
) -> list[schemas.DeliverableRead]:
    deliverables: list[schemas.DeliverableRead] = []
    for item in task.deliverables:
        linked_objects = [
            schemas.DeliverableObjectLinkRead(
                id=link.id,
                task_id=link.task_id,
                deliverable_id=link.deliverable_id,
                object_type=link.object_type,
                object_id=link.object_id,
                object_label=(
                    link.object_label
                    or (
                        object_label_map.get(link.object_type, {}).get(link.object_id, "")
                        if link.object_id
                        else ""
                    )
                    or None
                ),
                relation_type=link.relation_type,
                created_at=link.created_at,
            )
            for link in item.object_links
        ]
        deliverables.append(
            schemas.DeliverableRead(
                id=item.id,
                task_id=item.task_id,
                task_run_id=item.task_run_id,
                deliverable_type=item.deliverable_type,
                title=item.title,
                summary=_resolve_deliverable_summary_text(item),
                status=_default_deliverable_status(task, item),
                version_tag=(item.version_tag or f"v{item.version}"),
                content_structure=item.content_structure,
                version=item.version,
                linked_objects=linked_objects,
                generated_at=item.generated_at,
            )
        )
    return deliverables


def get_loaded_task(db: Session, task_id: str) -> models.Task:
    statement = select(models.Task).options(*task_load_options()).where(models.Task.id == task_id)
    task = db.scalars(statement).unique().one_or_none()
    if task is None:
        raise HTTPException(status_code=404, detail="找不到指定任務。")
    return task


def update_task_extension_overrides(
    db: Session,
    task_id: str,
    payload: schemas.TaskExtensionOverrideRequest,
) -> models.Task:
    task = get_loaded_task(db, task_id)
    pack_override_ids = _validate_pack_override_ids(
        _normalize_override_ids(payload.pack_override_ids)
    )
    agent_override_ids = _validate_agent_override_ids(
        _normalize_override_ids(payload.agent_override_ids)
    )

    for constraint in list(task.constraints):
        if constraint.constraint_type in {PACK_OVERRIDE_CONSTRAINT_TYPE, AGENT_OVERRIDE_CONSTRAINT_TYPE}:
            db.delete(constraint)

    for pack_id in pack_override_ids:
        db.add(
            models.Constraint(
                task_id=task.id,
                constraint_type=PACK_OVERRIDE_CONSTRAINT_TYPE,
                description=pack_id,
                severity="low",
            )
        )

    for agent_id in agent_override_ids:
        db.add(
            models.Constraint(
                task_id=task.id,
                constraint_type=AGENT_OVERRIDE_CONSTRAINT_TYPE,
                description=agent_id,
                severity="low",
            )
        )

    db.commit()
    return get_loaded_task(db, task.id)


def create_task(db: Session, payload: schemas.TaskCreateRequest) -> models.Task:
    logger.info(
        "Creating task title=%s task_type=%s mode=%s",
        payload.title,
        payload.task_type,
        payload.mode.value,
    )
    compiled_seed = compile_case_world_seed_from_payload(payload)
    matter_workspace, _ = ensure_matter_workspace_for_seed(
        db,
        seed=compiled_seed,
        continuity_mode=payload.engagement_continuity_mode,
        writeback_depth=payload.writeback_depth,
    )
    world_state, _ = ensure_case_world_state_from_seed(
        db,
        matter_workspace=matter_workspace,
        seed=compiled_seed,
    )

    task = models.Task(
        title=payload.title,
        description=payload.description,
        task_type=payload.task_type,
        mode=payload.mode.value,
        entry_preset=payload.entry_preset.value,
        status=TaskStatus.READY.value,
    )
    db.add(task)
    db.flush()

    context = models.TaskContext(
        task_id=task.id,
        summary=compiled_seed.background_brief,
        assumptions=payload.assumptions,
        notes=payload.notes,
        version=1,
    )
    db.add(context)
    db.flush()

    client = models.Client(
        task_id=task.id,
        name=compiled_seed.client.name,
        client_type=compiled_seed.client.client_type,
        client_stage=compiled_seed.client.client_stage,
        description=compiled_seed.client.description,
    )
    db.add(client)
    db.flush()

    engagement = models.Engagement(
        task_id=task.id,
        client_id=client.id,
        name=compiled_seed.engagement.name,
        description=compiled_seed.engagement.description,
    )
    db.add(engagement)
    db.flush()

    workstream = models.Workstream(
        task_id=task.id,
        engagement_id=engagement.id,
        name=compiled_seed.workstream.name,
        description=compiled_seed.workstream.description,
        domain_lenses=compiled_seed.domain_lenses,
    )
    db.add(workstream)
    db.flush()

    if payload.subject_name:
        db.add(
            models.Subject(
                task_id=task.id,
                subject_type=payload.subject_type,
                name=payload.subject_name,
                description=payload.subject_description,
                source_ref=None,
            )
        )

    db.add(
        models.Goal(
            task_id=task.id,
            goal_type=payload.goal_type,
            description=compiled_seed.goal_description,
            success_criteria=compiled_seed.success_criteria,
            priority="high",
        )
    )

    for item in compiled_seed.inferred_constraints:
        db.add(
            models.Constraint(
                task_id=task.id,
                constraint_type=item.constraint_type,
                description=item.description,
                severity=item.severity,
            )
        )

    strategy_constraint = _build_external_data_strategy_constraint(payload.external_data_strategy)
    db.add(
        models.Constraint(
            task_id=task.id,
            constraint_type=strategy_constraint.constraint_type,
            description=strategy_constraint.description,
            severity=strategy_constraint.severity,
        )
    )

    constraint_summary = "；".join(item.description for item in compiled_seed.inferred_constraints)
    decision_context = models.DecisionContext(
        task_id=task.id,
        client_id=client.id,
        engagement_id=engagement.id,
        workstream_id=workstream.id,
        title=compiled_seed.decision_context.title,
        summary=compiled_seed.decision_context.summary,
        judgment_to_make=compiled_seed.decision_context.judgment_to_make,
        domain_lenses=compiled_seed.domain_lenses,
        client_stage=client.client_stage,
        client_type=client.client_type,
        goal_summary=compiled_seed.goal_description,
        constraint_summary=constraint_summary or None,
        assumption_summary=payload.assumptions,
        source_priority=compiled_seed.decision_context.source_priority,
        external_data_policy=compiled_seed.decision_context.external_data_policy,
    )
    db.add(decision_context)
    db.flush()

    if compiled_seed.background_brief.strip():
        db.add(
            models.Evidence(
                task_id=task.id,
                evidence_type="background_text",
                source_type="manual_input",
                source_ref=f"task_context:{context.id}",
                title="任務背景摘要",
                excerpt_or_summary=compiled_seed.background_brief.strip()[:1500],
                reliability_level="user_provided",
            )
        )

    ensure_matter_workspace_for_task(
        db,
        task,
        _build_client_read_from_row(client, reference_task_id=task.id),
        _build_engagement_read_from_row(engagement, reference_task_id=task.id),
        _build_workstream_read_from_row(workstream, reference_task_id=task.id),
        schemas.DecisionContextRead(
            id=decision_context.id,
            task_id=decision_context.task_id,
            task_reference_role="slice_linkage",
            client_id=decision_context.client_id,
            engagement_id=decision_context.engagement_id,
            workstream_id=decision_context.workstream_id,
            title=decision_context.title,
            summary=decision_context.summary,
            judgment_to_make=decision_context.judgment_to_make,
            domain_lenses=decision_context.domain_lenses,
            client_stage=decision_context.client_stage,
            client_type=decision_context.client_type,
            goals=[compiled_seed.goal_description],
            constraints=[item.description for item in compiled_seed.inferred_constraints],
            assumptions=_split_multiline_items(payload.assumptions),
            source_priority=decision_context.source_priority or "",
            external_data_policy=decision_context.external_data_policy or "",
            created_at=decision_context.created_at,
        ),
        compiled_seed.domain_lenses,
        payload.engagement_continuity_mode,
        payload.writeback_depth,
    )
    world_state, _ = ensure_world_context_spine_for_task(
        db,
        matter_workspace=matter_workspace,
        task=task,
        client=client,
        engagement=engagement,
        workstream=workstream,
        decision_context=decision_context,
        case_world_state=world_state,
        refresh_compatibility_task_reference=True,
    )
    world_state.active_task_ids = _unique_preserve_order([*world_state.active_task_ids, task.id])
    world_state.latest_task_id = task.id
    world_state.latest_task_title = task.title
    db.add(world_state)

    db.commit()
    return get_loaded_task(db, task.id)


def _build_task_list_item_response(
    task: models.Task,
    matter_workspace: schemas.MatterWorkspaceSummaryRead | None = None,
) -> schemas.TaskListItemResponse:
    latest_deliverable = _latest_deliverable(task)
    client, engagement, workstream, preferred_decision_context, domain_lenses, source_materials, artifacts = (
        _build_world_preferred_ontology_spine_for_task(task)
    )
    pack_resolution = resolve_pack_selection_for_task(
        task,
        client,
        engagement,
        workstream,
        preferred_decision_context,
        domain_lenses,
    )
    input_entry_mode = _infer_input_entry_mode(task, source_materials, artifacts)
    external_research_heavy_candidate = _is_external_research_heavy_candidate(
        task,
        preferred_decision_context,
        source_materials,
        artifacts,
        input_entry_mode,
    )
    deliverable_class_hint = _resolve_deliverable_class_hint(
        input_entry_mode,
        preferred_decision_context,
        source_materials,
        artifacts,
        len(_usable_evidence(task)),
        external_research_heavy_candidate,
    )
    agent_selection = resolve_agent_selection_for_task(
        task,
        preferred_decision_context,
        domain_lenses,
        pack_resolution,
        input_entry_mode,
        deliverable_class_hint,
        external_research_heavy_candidate,
        source_materials,
        artifacts,
    )
    continuity_mode, writeback_depth = resolve_continuity_policy_for_task(task)
    return schemas.TaskListItemResponse(
        id=task.id,
        title=task.title,
        description=task.description,
        task_type=task.task_type,
        mode=task.mode,
        status=task.status,
        created_at=task.created_at,
        updated_at=task.updated_at,
        client_name=client.name if client else None,
        engagement_name=engagement.name if engagement else None,
        workstream_name=workstream.name if workstream else None,
        decision_context_title=preferred_decision_context.title if preferred_decision_context else None,
        client_stage=(
            preferred_decision_context.client_stage
            if preferred_decision_context and preferred_decision_context.client_stage != UNSPECIFIED_LABEL
            else client.client_stage if client and client.client_stage != UNSPECIFIED_LABEL else None
        ),
        client_type=(
            preferred_decision_context.client_type
            if preferred_decision_context and preferred_decision_context.client_type != UNSPECIFIED_LABEL
            else client.client_type if client and client.client_type != UNSPECIFIED_LABEL else None
        ),
        domain_lenses=domain_lenses,
        entry_preset=InputEntryMode(task.entry_preset),
        input_entry_mode=input_entry_mode,
        engagement_continuity_mode=continuity_mode,
        writeback_depth=writeback_depth,
        deliverable_class_hint=deliverable_class_hint,
        external_research_heavy_candidate=external_research_heavy_candidate,
        selected_pack_ids=pack_resolution.stack_order,
        selected_pack_names=[
            *[item.pack_name for item in pack_resolution.selected_domain_packs],
            *[item.pack_name for item in pack_resolution.selected_industry_packs],
        ],
        pack_summary=(
            f"Domain Packs：{join_natural_list([item.pack_name for item in pack_resolution.selected_domain_packs]) or '未選用'}；"
            f"Industry Packs：{join_natural_list([item.pack_name for item in pack_resolution.selected_industry_packs]) or '未選用'}"
        ),
        selected_agent_ids=agent_selection.selected_agent_ids,
        selected_agent_names=agent_selection.selected_agent_names,
        agent_summary=(
            f"Host：{agent_selection.host_agent.agent_name}；"
            f"Selected agents：{join_natural_list(agent_selection.selected_agent_names[:3]) or '尚未選到其他 agents'}"
            if agent_selection.host_agent
            else None
        ),
        evidence_count=len(task.evidence),
        deliverable_count=len(task.deliverables),
        run_count=len(task.runs),
        latest_deliverable_id=latest_deliverable.id if latest_deliverable else None,
        latest_deliverable_title=latest_deliverable.title if latest_deliverable else None,
        latest_deliverable_summary=(
            _resolve_deliverable_summary_text(latest_deliverable)
            if latest_deliverable
            else None
        ),
        latest_deliverable_status=(
            _default_deliverable_status(task, latest_deliverable)
            if latest_deliverable
            else None
        ),
        latest_deliverable_version_tag=(
            latest_deliverable.version_tag or f"v{latest_deliverable.version}"
            if latest_deliverable
            else None
        ),
        matter_workspace=matter_workspace,
    )


def list_tasks(db: Session) -> list[schemas.TaskListItemResponse]:
    statement = select(models.Task).options(*task_load_options()).order_by(models.Task.updated_at.desc())
    tasks = db.scalars(statement).unique().all()
    matter_workspaces: dict[str, models.MatterWorkspace] = {}
    workspace_id_by_task_id: dict[str, str] = {}
    changed = False

    for task in tasks:
        client, engagement, workstream, decision_context, domain_lenses, _, _ = _build_world_preferred_ontology_spine_for_task(task)
        matter_workspace, workspace_changed = ensure_matter_workspace_for_task(
            db,
            task,
            client,
            engagement,
            workstream,
            decision_context,
            domain_lenses,
        )
        matter_workspaces[matter_workspace.id] = matter_workspace
        workspace_id_by_task_id[task.id] = matter_workspace.id
        changed = changed or workspace_changed

    if changed:
        db.commit()

    matter_summary_by_id = _build_matter_workspace_summary_map(db, list(matter_workspaces.values()))

    return [
        _build_task_list_item_response(
            task,
            matter_summary_by_id.get(workspace_id_by_task_id.get(task.id, "")),
        )
        for task in tasks
    ]


def list_matter_workspaces(db: Session) -> list[schemas.MatterWorkspaceSummaryRead]:
    tasks = db.scalars(
        select(models.Task).options(*task_load_options()).order_by(models.Task.updated_at.desc())
    ).unique().all()
    changed = False

    for task in tasks:
        client, engagement, workstream, decision_context, domain_lenses, _, _ = _build_world_preferred_ontology_spine_for_task(task)
        _, workspace_changed = ensure_matter_workspace_for_task(
            db,
            task,
            client,
            engagement,
            workstream,
            decision_context,
            domain_lenses,
        )
        changed = changed or workspace_changed

    if changed:
        db.commit()

    matter_workspaces = db.scalars(
        select(models.MatterWorkspace).order_by(models.MatterWorkspace.updated_at.desc())
    ).all()
    summary_by_id = _build_matter_workspace_summary_map(db, matter_workspaces)
    return sorted(
        summary_by_id.values(),
        key=lambda item: item.latest_updated_at,
        reverse=True,
    )


def get_matter_workspace(db: Session, matter_id: str) -> schemas.MatterWorkspaceResponse:
    matter_workspace = db.scalars(
        select(models.MatterWorkspace)
        .options(selectinload(models.MatterWorkspace.case_world_state))
        .where(models.MatterWorkspace.id == matter_id)
    ).one_or_none()
    if matter_workspace is None:
        raise HTTPException(status_code=404, detail="找不到指定案件工作面。")
    content_revisions = ensure_matter_content_revisions(db, matter_workspace)

    related_tasks = _load_tasks_for_matter_workspaces(db, [matter_workspace.id]).get(
        matter_workspace.id,
        [],
    )
    if not related_tasks:
        raise HTTPException(status_code=404, detail="找不到此案件工作面的相關任務。")

    latest_task = related_tasks[0]
    _, world_spine_changed = ensure_world_context_spine_for_task(
        db,
        matter_workspace=matter_workspace,
        task=latest_task,
        client=_primary_item(latest_task.clients),
        engagement=_primary_item(latest_task.engagements),
        workstream=_primary_item(latest_task.workstreams),
        decision_context=_primary_item(latest_task.decision_contexts),
        case_world_state=matter_workspace.case_world_state,
    )
    if world_spine_changed:
        db.commit()
        matter_workspace = db.scalars(
            select(models.MatterWorkspace)
            .options(selectinload(models.MatterWorkspace.case_world_state))
            .where(models.MatterWorkspace.id == matter_id)
        ).one()
        related_tasks = _load_tasks_for_matter_workspaces(db, [matter_workspace.id]).get(
            matter_workspace.id,
            [],
        )
        if not related_tasks:
            raise HTTPException(status_code=404, detail="找不到此案件工作面的相關任務。")

    summary = _build_matter_workspace_summary_from_tasks(matter_workspace, related_tasks)
    latest_task = related_tasks[0]
    latest_task_spine = _build_world_preferred_ontology_spine_for_task(latest_task)
    client = latest_task_spine[0]
    engagement = latest_task_spine[1]
    workstream = latest_task_spine[2]
    current_decision_context = latest_task_spine[3]

    related_task_items = [
        _build_task_list_item_response(task, summary) for task in related_tasks[:8]
    ]
    decision_trajectory: list[schemas.MatterDecisionPointRead] = []
    related_deliverables: list[schemas.MatterDeliverableSummaryRead] = []
    related_artifacts: list[schemas.MatterMaterialSummaryRead] = []
    related_source_materials: list[schemas.MatterMaterialSummaryRead] = []
    case_world_drafts: list[schemas.CaseWorldDraftRead] = []
    evidence_gap_records: list[schemas.EvidenceGapRead] = []
    research_runs: list[schemas.ResearchRunRead] = []
    decision_records: list[schemas.DecisionRecordRead] = []
    action_plans: list[schemas.ActionPlanRead] = []
    action_executions: list[schemas.ActionExecutionRead] = []
    outcome_records: list[schemas.OutcomeRecordRead] = []
    seen_material_keys: set[str] = set()
    continuity_notes: list[str] = []

    for task in related_tasks:
        if not task.case_world_drafts:
            serialize_task(task)
            task = get_loaded_task(db, task.id)
        _, _, _, decision_context, _, source_materials, artifacts = _build_world_preferred_ontology_spine_for_task(task)
        input_entry_mode = _infer_input_entry_mode(task, source_materials, artifacts)
        external_research_heavy_candidate = _is_external_research_heavy_candidate(
            task,
            decision_context,
            source_materials,
            artifacts,
            input_entry_mode,
        )
        deliverable_class_hint = _resolve_deliverable_class_hint(
            input_entry_mode,
            decision_context,
            source_materials,
            artifacts,
            len(_usable_evidence(task)),
            external_research_heavy_candidate,
        )
        decision_trajectory.append(
            schemas.MatterDecisionPointRead(
                task_id=task.id,
                task_title=task.title,
                task_status=TaskStatus(task.status),
                decision_context_id=decision_context.id if decision_context else None,
                decision_context_title=decision_context.title if decision_context else task.title,
                judgment_to_make=(
                    decision_context.judgment_to_make
                    if decision_context and decision_context.judgment_to_make
                    else task.description or task.title
                ),
                deliverable_class_hint=deliverable_class_hint,
                updated_at=task.updated_at,
            )
        )
        if task.case_world_drafts:
            case_world_drafts.extend(
                filter(
                    None,
                    [_serialize_case_world_draft(item) for item in task.case_world_drafts[:1]],
                )
            )
        evidence_gap_records.extend(
            schemas.EvidenceGapRead.model_validate(item)
            for item in task.evidence_gaps
        )
        research_runs.extend(
            schemas.ResearchRunRead.model_validate(item)
            for item in task.research_runs
        )
        decision_records.extend(
            schemas.DecisionRecordRead.model_validate(item)
            for item in task.decision_records
        )
        action_plans.extend(
            schemas.ActionPlanRead.model_validate(item)
            for item in task.action_plans
        )
        action_executions.extend(
            schemas.ActionExecutionRead.model_validate(item)
            for item in task.action_executions
        )
        outcome_records.extend(
            schemas.OutcomeRecordRead.model_validate(item)
            for item in task.outcome_records
        )
        for deliverable in sorted(task.deliverables, key=lambda item: (item.generated_at, item.version), reverse=True)[:3]:
            related_deliverables.append(
                schemas.MatterDeliverableSummaryRead(
                    deliverable_id=deliverable.id,
                    task_id=task.id,
                    task_title=task.title,
                    title=deliverable.title,
                    summary=_resolve_deliverable_summary_text(deliverable),
                    status=_default_deliverable_status(task, deliverable),
                    version_tag=(deliverable.version_tag or f"v{deliverable.version}"),
                    deliverable_type=deliverable.deliverable_type,
                    version=deliverable.version,
                    generated_at=deliverable.generated_at,
                    decision_context_title=decision_context.title if decision_context else None,
                )
            )
        for artifact in artifacts:
            key = f"artifact:{artifact.id}"
            if key in seen_material_keys:
                continue
            seen_material_keys.add(key)
            related_artifacts.append(
                schemas.MatterMaterialSummaryRead(
                    object_id=artifact.id,
                    task_id=task.id,
                    task_title=task.title,
                    object_type="artifact",
                    title=artifact.title,
                    summary=artifact.description or artifact.artifact_type,
                    created_at=artifact.created_at,
                )
            )
        for source_material in source_materials:
            key = f"source_material:{source_material.id}"
            if key in seen_material_keys:
                continue
            seen_material_keys.add(key)
            related_source_materials.append(
                schemas.MatterMaterialSummaryRead(
                    object_id=source_material.id,
                    task_id=task.id,
                    task_title=task.title,
                    object_type="source_material",
                    title=source_material.title,
                    summary=source_material.summary,
                    file_extension=source_material.file_extension,
                    content_type=source_material.content_type,
                    file_size=source_material.file_size,
                    ingest_status=source_material.ingest_status,
                    support_level=source_material.support_level,
                    retention_policy=source_material.retention_policy,
                    purge_at=source_material.purge_at,
                    availability_state=source_material.availability_state,
                    metadata_only=source_material.metadata_only,
                    created_at=source_material.created_at,
                )
            )

    if summary.total_task_count > 1:
        continuity_notes.append(
            f"這個案件世界目前已串起 {summary.total_task_count} 個 task slices，可回看 decision trajectory 與跨 task deliverables。"
        )
    if summary.active_task_count > 0:
        continuity_notes.append(f"目前仍有 {summary.active_task_count} 個 active work item 正在這個案件世界內推進。")
    if summary.artifact_count == 0 and summary.source_material_count == 0:
        continuity_notes.append("目前案件材料仍偏 sparse，工作面主要依賴 DecisionContext 與 Host framing。")

    readiness_hint = (
        f"最近一次工作屬於「{summary.current_decision_context_title or latest_task.title}」，"
        f"目前累積 {summary.deliverable_count} 份 deliverables、{summary.artifact_count} 份 artifacts、"
        f"{summary.source_material_count} 份 source materials。"
    )
    matter_workspace = db.scalars(
        select(models.MatterWorkspace)
        .options(selectinload(models.MatterWorkspace.case_world_state))
        .where(models.MatterWorkspace.id == matter_id)
    ).one()

    return schemas.MatterWorkspaceResponse(
        summary=summary,
        client=client,
        engagement=engagement,
        workstream=workstream,
        current_decision_context=current_decision_context,
        case_world_state=_serialize_case_world_state(matter_workspace.case_world_state),
        content_sections=_normalize_matter_content_sections(matter_workspace.content_sections),
        content_revisions=[
            _build_matter_content_revision_read(item) for item in content_revisions[:12]
        ],
        decision_trajectory=decision_trajectory[:8],
        related_tasks=related_task_items,
        related_deliverables=related_deliverables[:8],
        related_artifacts=sorted(
            related_artifacts,
            key=lambda item: item.created_at,
            reverse=True,
        )[:8],
        related_source_materials=sorted(
            related_source_materials,
            key=lambda item: item.created_at,
            reverse=True,
        )[:8],
        case_world_drafts=sorted(case_world_drafts, key=lambda item: item.updated_at, reverse=True)[:6],
        evidence_gaps=sorted(evidence_gap_records, key=lambda item: item.updated_at, reverse=True)[:10],
        research_runs=sorted(research_runs, key=lambda item: item.started_at, reverse=True)[:6],
        decision_records=sorted(decision_records, key=lambda item: item.created_at, reverse=True)[:6],
        action_plans=sorted(action_plans, key=lambda item: item.created_at, reverse=True)[:6],
        action_executions=sorted(action_executions, key=lambda item: item.updated_at, reverse=True)[:10],
        outcome_records=sorted(outcome_records, key=lambda item: item.created_at, reverse=True)[:10],
        readiness_hint=readiness_hint,
        continuity_notes=continuity_notes,
    )


def update_matter_workspace(
    db: Session,
    matter_id: str,
    payload: schemas.MatterWorkspaceUpdateRequest,
) -> schemas.MatterWorkspaceResponse:
    matter_workspace = db.scalars(
        select(models.MatterWorkspace).where(models.MatterWorkspace.id == matter_id)
    ).one_or_none()
    if matter_workspace is None:
        raise HTTPException(status_code=404, detail="找不到指定案件工作面。")

    previous_content_sections = _normalize_matter_content_sections(matter_workspace.content_sections)
    next_content_sections_payload = _serialize_matter_content_sections(payload.content_sections)
    next_content_sections = _normalize_matter_content_sections(next_content_sections_payload)

    matter_workspace.title = payload.title.strip()
    matter_workspace.summary = payload.summary.strip()
    matter_workspace.status = payload.status
    matter_workspace.content_sections = next_content_sections_payload
    matter_workspace.title_override_active = True
    db.add(matter_workspace)
    if next_content_sections != previous_content_sections:
        create_matter_content_revision(
            db,
            matter_workspace,
            snapshot=next_content_sections_payload,
            previous_snapshot=_serialize_matter_content_sections(
                schemas.MatterWorkspaceContentSectionsRequest(
                    core_question=previous_content_sections.core_question,
                    analysis_focus=previous_content_sections.analysis_focus,
                    constraints_and_risks=previous_content_sections.constraints_and_risks,
                    next_steps=previous_content_sections.next_steps,
                )
            ),
            source=CONTENT_REVISION_SOURCE_MANUAL_EDIT,
        )
    db.commit()

    return get_matter_workspace(db, matter_id)


def update_matter_workspace_metadata(
    db: Session,
    matter_id: str,
    payload: schemas.MatterWorkspaceMetadataUpdateRequest,
) -> schemas.MatterWorkspaceResponse:
    matter_workspace = db.scalars(
        select(models.MatterWorkspace).where(models.MatterWorkspace.id == matter_id)
    ).one_or_none()
    if matter_workspace is None:
        raise HTTPException(status_code=404, detail="找不到指定案件工作面。")

    return update_matter_workspace(
        db,
        matter_id,
        schemas.MatterWorkspaceUpdateRequest(
            title=payload.title,
            summary=payload.summary,
            status=payload.status,
            content_sections=schemas.MatterWorkspaceContentSectionsRequest(
                core_question=_normalize_matter_content_sections(
                    matter_workspace.content_sections
                ).core_question,
                analysis_focus=_normalize_matter_content_sections(
                    matter_workspace.content_sections
                ).analysis_focus,
                constraints_and_risks=_normalize_matter_content_sections(
                    matter_workspace.content_sections
                ).constraints_and_risks,
                next_steps=_normalize_matter_content_sections(
                    matter_workspace.content_sections
                ).next_steps,
            ),
        ),
    )


def rollback_matter_content_revision(
    db: Session,
    matter_id: str,
    revision_id: str,
) -> schemas.MatterWorkspaceResponse:
    matter_workspace = db.scalars(
        select(models.MatterWorkspace).where(models.MatterWorkspace.id == matter_id)
    ).one_or_none()
    if matter_workspace is None:
        raise HTTPException(status_code=404, detail="找不到指定案件工作面。")

    target_revision = db.scalars(
        select(models.MatterContentRevision).where(
            models.MatterContentRevision.id == revision_id,
            models.MatterContentRevision.matter_workspace_id == matter_id,
        )
    ).one_or_none()
    if target_revision is None:
        raise HTTPException(status_code=404, detail="找不到指定案件正文修訂。")

    current_snapshot = _normalize_matter_content_sections(matter_workspace.content_sections)
    target_snapshot = _normalize_matter_content_sections(target_revision.snapshot)
    if target_snapshot == current_snapshot:
        raise HTTPException(status_code=400, detail="目前案件正文已與目標修訂相同。")

    serialized_target_snapshot = {
        "core_question": target_snapshot.core_question,
        "analysis_focus": target_snapshot.analysis_focus,
        "constraints_and_risks": target_snapshot.constraints_and_risks,
        "next_steps": target_snapshot.next_steps,
    }
    matter_workspace.content_sections = serialized_target_snapshot
    matter_workspace.updated_at = models.utc_now()
    db.add(matter_workspace)
    db.flush()

    create_matter_content_revision(
        db,
        matter_workspace,
        snapshot=serialized_target_snapshot,
        previous_snapshot={
            "core_question": current_snapshot.core_question,
            "analysis_focus": current_snapshot.analysis_focus,
            "constraints_and_risks": current_snapshot.constraints_and_risks,
            "next_steps": current_snapshot.next_steps,
        },
        source=CONTENT_REVISION_SOURCE_ROLLBACK,
        revision_summary=f"回退案件正文到 {target_revision.created_at.isoformat()} 的修訂",
        rollback_target_revision_id=target_revision.id,
    )
    db.commit()
    return get_matter_workspace(db, matter_id)


def _material_role_label(
    linked_evidence_count: int,
    linked_output_count: int,
    ingest_status: str | None = None,
    continuity_scope: str | None = None,
    participation_task_count: int = 0,
) -> tuple[str, PresenceState]:
    if continuity_scope == WORLD_SHARED_CONTINUITY_SCOPE and participation_task_count >= 2:
        return "共享主來源", PresenceState.EXPLICIT
    if continuity_scope == WORLD_SHARED_CONTINUITY_SCOPE:
        return "正式來源", PresenceState.EXPLICIT
    if linked_output_count >= 2 or linked_evidence_count >= 2:
        return "主要來源", PresenceState.EXPLICIT
    if linked_evidence_count >= 1:
        return "補充來源", PresenceState.EXPLICIT
    if ingest_status and ingest_status != "processed":
        return "暫定來源", PresenceState.PROVISIONAL
    return "待補強來源", PresenceState.PROVISIONAL


def _workspace_material_participation_fields(
    *,
    task_id: str,
    continuity_scope: str | None,
    participation: schemas.ObjectParticipationRead | None,
) -> tuple[str | None, int, bool, str | None, str | None, str | None]:
    if participation is not None:
        return (
            participation.participation_type,
            participation.participation_task_count,
            participation.current_task_participation,
            participation.mapping_mode,
            participation.canonical_owner_scope,
            participation.compatibility_task_id,
        )
    if continuity_scope == WORLD_SHARED_CONTINUITY_SCOPE:
        return None, 0, False, None, "world_canonical", None
    return "slice_local_only", 1, True, "slice_local_only", "slice_local", task_id


def _evidence_strength_label(
    linked_output_count: int,
    linked_deliverable_count: int,
    reliability_level: str,
) -> str:
    normalized_reliability = (reliability_level or "").lower()
    if linked_output_count >= 2 or (
        linked_output_count >= 1 and normalized_reliability in {"user_provided", "high", "high_confidence"}
    ):
        return "strong"
    if linked_output_count >= 1 or linked_deliverable_count >= 1:
        return "moderate"
    if normalized_reliability in {"user_provided", "medium"}:
        return "moderate"
    return "thin"


def _build_evidence_support_note(
    linked_recommendations: list[schemas.EvidenceSupportTargetRead],
    linked_risks: list[schemas.EvidenceSupportTargetRead],
    linked_action_items: list[schemas.EvidenceSupportTargetRead],
    linked_deliverables: list[schemas.EvidenceSupportTargetRead],
) -> str:
    parts: list[str] = []
    if linked_recommendations:
        parts.append(f"{len(linked_recommendations)} 項 recommendations")
    if linked_risks:
        parts.append(f"{len(linked_risks)} 項 risks")
    if linked_action_items:
        parts.append(f"{len(linked_action_items)} 項 action items")
    if linked_deliverables:
        parts.append(f"{len(linked_deliverables)} 份 deliverables")

    if parts:
        return f"這則 evidence 目前已正式支撐 {join_natural_list(parts)}。"

    return "這則 evidence 已被保存，但尚未形成足夠清楚的 recommendation / risk / action 支撐鏈。"


def get_artifact_evidence_workspace(
    db: Session,
    matter_id: str,
) -> schemas.ArtifactEvidenceWorkspaceResponse:
    matter_workspace = db.scalars(
        select(models.MatterWorkspace)
        .options(selectinload(models.MatterWorkspace.case_world_state))
        .where(models.MatterWorkspace.id == matter_id)
    ).one_or_none()
    if matter_workspace is None:
        raise HTTPException(status_code=404, detail="找不到指定來源 / 證據工作面。")

    related_tasks = _load_tasks_for_matter_workspaces(db, [matter_workspace.id]).get(
        matter_workspace.id,
        [],
    )
    if not related_tasks:
        raise HTTPException(status_code=404, detail="找不到此來源 / 證據工作面的相關任務。")

    latest_task = related_tasks[0]
    _, world_spine_changed = ensure_world_context_spine_for_task(
        db,
        matter_workspace=matter_workspace,
        task=latest_task,
        client=_primary_item(latest_task.clients),
        engagement=_primary_item(latest_task.engagements),
        workstream=_primary_item(latest_task.workstreams),
        decision_context=_primary_item(latest_task.decision_contexts),
        case_world_state=matter_workspace.case_world_state,
    )
    if world_spine_changed:
        db.commit()
        matter_workspace = db.scalars(
            select(models.MatterWorkspace)
            .options(selectinload(models.MatterWorkspace.case_world_state))
            .where(models.MatterWorkspace.id == matter_id)
        ).one()
        related_tasks = _load_tasks_for_matter_workspaces(db, [matter_workspace.id]).get(
            matter_workspace.id,
            [],
        )
        if not related_tasks:
            raise HTTPException(status_code=404, detail="找不到此來源 / 證據工作面的相關任務。")

    matter_summary = _build_matter_workspace_summary_from_tasks(matter_workspace, related_tasks)
    latest_task = related_tasks[0]
    latest_task_spine = _build_world_preferred_ontology_spine_for_task(latest_task)
    client = latest_task_spine[0]
    engagement = latest_task_spine[1]
    workstream = latest_task_spine[2]
    current_decision_context = latest_task_spine[3]
    related_task_items = [
        _build_task_list_item_response(task, matter_summary) for task in related_tasks[:8]
    ]

    source_material_cards: list[schemas.ArtifactEvidenceMaterialRead] = []
    artifact_cards: list[schemas.ArtifactEvidenceMaterialRead] = []
    evidence_chains: list[schemas.EvidenceWorkspaceEvidenceRead] = []
    evidence_expectations: list[str] = []
    high_impact_gaps: list[str] = []
    evidence_gap_records: list[schemas.EvidenceGapRead] = []
    research_runs: list[schemas.ResearchRunRead] = []
    deliverable_limitations: list[str] = []
    continuity_notes: list[str] = []

    source_material_seen: set[str] = set()
    artifact_seen: set[str] = set()
    evidence_seen: set[str] = set()
    task_ids_with_materials: set[str] = set()
    task_ids_with_evidence: set[str] = set()
    unsupported_recommendation_count = 0
    unsupported_risk_count = 0
    unsupported_action_item_count = 0
    linked_evidence_count = 0
    linked_deliverable_count = 0
    external_research_heavy_detected = False

    for task in related_tasks:
        if not task.case_world_drafts:
            serialize_task(task)
            task = get_loaded_task(db, task.id)
        task_title = task.title
        task_spine = _build_world_preferred_ontology_spine_for_task(task)
        decision_context = task_spine[3]
        domain_lenses = task_spine[4]
        source_materials = task_spine[5]
        artifacts = task_spine[6]
        evidence = _serialize_evidence_items(task, source_materials, artifacts)
        recommendation_support_map, risk_support_map, action_item_support_map = _build_supporting_evidence_maps(
            task
        )
        recommendations = _serialize_recommendations(task, recommendation_support_map)
        risks = _serialize_risks(task, risk_support_map)
        action_items = _serialize_action_items(task, action_item_support_map)
        deliverables = _serialize_deliverables(
            task,
            _build_object_label_map(
                task,
                source_materials,
                artifacts,
                evidence,
                recommendations,
                risks,
                action_items,
                decision_context,
                task_spine[0],
                task_spine[1],
                task_spine[2],
            ),
        )
        input_entry_mode = _infer_input_entry_mode(task, source_materials, artifacts)
        external_research_heavy_candidate = _is_external_research_heavy_candidate(
            task,
            decision_context,
            source_materials,
            artifacts,
            input_entry_mode,
        )
        external_research_heavy_detected = (
            external_research_heavy_detected or external_research_heavy_candidate
        )
        deliverable_class_hint = _resolve_deliverable_class_hint(
            input_entry_mode,
            decision_context,
            source_materials,
            artifacts,
            len(_usable_evidence(task)),
            external_research_heavy_candidate,
        )
        pack_resolution = resolve_pack_selection_for_task(
            task,
            task_spine[0],
            task_spine[1],
            task_spine[2],
            decision_context,
            domain_lenses,
        )
        evidence_expectations.extend(pack_resolution.evidence_expectations)
        evidence_gap_records.extend(
            schemas.EvidenceGapRead.model_validate(item)
            for item in task.evidence_gaps
        )
        research_runs.extend(
            schemas.ResearchRunRead.model_validate(item)
            for item in task.research_runs
        )

        unsupported_recommendation_count += sum(
            1 for item in recommendations if not item.supporting_evidence_ids
        )
        unsupported_risk_count += sum(1 for item in risks if not item.supporting_evidence_ids)
        unsupported_action_item_count += sum(
            1 for item in action_items if not item.supporting_evidence_ids
        )

        recommendation_targets_by_evidence: dict[str, list[schemas.EvidenceSupportTargetRead]] = {}
        risk_targets_by_evidence: dict[str, list[schemas.EvidenceSupportTargetRead]] = {}
        action_targets_by_evidence: dict[str, list[schemas.EvidenceSupportTargetRead]] = {}
        deliverable_targets_by_evidence: dict[str, list[schemas.EvidenceSupportTargetRead]] = {}

        for item in recommendations:
            for evidence_id in item.supporting_evidence_ids:
                recommendation_targets_by_evidence.setdefault(evidence_id, []).append(
                    schemas.EvidenceSupportTargetRead(
                        target_type="recommendation",
                        target_id=item.id,
                        task_id=item.task_id,
                        task_title=task_title,
                        title=item.summary,
                        note=f"優先級：{item.priority}",
                    )
                )
        for item in risks:
            for evidence_id in item.supporting_evidence_ids:
                risk_targets_by_evidence.setdefault(evidence_id, []).append(
                    schemas.EvidenceSupportTargetRead(
                        target_type="risk",
                        target_id=item.id,
                        task_id=item.task_id,
                        task_title=task_title,
                        title=item.title,
                        note=f"影響：{item.impact_level} / 可能性：{item.likelihood_level}",
                    )
                )
        for item in action_items:
            for evidence_id in item.supporting_evidence_ids:
                action_targets_by_evidence.setdefault(evidence_id, []).append(
                    schemas.EvidenceSupportTargetRead(
                        target_type="action_item",
                        target_id=item.id,
                        task_id=item.task_id,
                        task_title=task_title,
                        title=item.description,
                        note=f"優先級：{item.priority}",
                    )
                )
        for deliverable in deliverables:
            for link in deliverable.linked_objects:
                if link.object_type != "evidence" or not link.object_id:
                    continue
                deliverable_targets_by_evidence.setdefault(link.object_id, []).append(
                    schemas.EvidenceSupportTargetRead(
                        target_type="deliverable",
                        target_id=deliverable.id,
                        task_id=deliverable.task_id,
                        task_title=task_title,
                        title=deliverable.title,
                        note=f"v{deliverable.version}",
                    )
                )

        evidence_by_source_material: dict[str, list[schemas.EvidenceRead]] = {}
        evidence_by_artifact: dict[str, list[schemas.EvidenceRead]] = {}
        for item in evidence:
            if item.source_material_id:
                evidence_by_source_material.setdefault(item.source_material_id, []).append(item)
            if item.artifact_id:
                evidence_by_artifact.setdefault(item.artifact_id, []).append(item)

        for source_material in source_materials:
            if source_material.id in source_material_seen:
                continue
            source_material_seen.add(source_material.id)
            linked_evidence = evidence_by_source_material.get(source_material.id, [])
            linked_output_count = sum(
                len(recommendation_targets_by_evidence.get(item.id, []))
                + len(risk_targets_by_evidence.get(item.id, []))
                + len(action_targets_by_evidence.get(item.id, []))
                for item in linked_evidence
            )
            role_label, presence_state = _material_role_label(
                len(linked_evidence),
                linked_output_count,
                source_material.ingest_status,
                source_material.continuity_scope,
                source_material.participation.participation_task_count
                if source_material.participation is not None
                else 0,
            )
            (
                participation_type,
                participation_task_count,
                current_task_participation,
                mapping_mode,
                canonical_owner_scope,
                compatibility_task_id,
            ) = _workspace_material_participation_fields(
                task_id=task.id,
                continuity_scope=source_material.continuity_scope,
                participation=source_material.participation,
            )
            source_material_cards.append(
                schemas.ArtifactEvidenceMaterialRead(
                    object_id=source_material.id,
                    task_id=task.id,
                    task_title=task_title,
                    object_type="source_material",
                    title=source_material.title,
                    summary=source_material.summary,
                    role_label=role_label,
                    presence_state=presence_state,
                    continuity_scope=source_material.continuity_scope,
                    participation_type=participation_type,
                    participation_task_count=participation_task_count,
                    current_task_participation=current_task_participation,
                    canonical_owner_scope=canonical_owner_scope,
                    compatibility_task_id=compatibility_task_id,
                    mapping_mode=mapping_mode,
                    source_type=source_material.source_type,
                    ingest_status=source_material.ingest_status,
                    support_level=source_material.support_level,
                    ingest_strategy=source_material.ingest_strategy,
                    source_ref=source_material.source_ref,
                    file_extension=source_material.file_extension,
                    content_type=source_material.content_type,
                    file_size=source_material.file_size,
                    storage_key=source_material.storage_key,
                    retention_policy=source_material.retention_policy,
                    purge_at=source_material.purge_at,
                    availability_state=source_material.availability_state,
                    metadata_only=source_material.metadata_only,
                    linked_evidence_count=len(linked_evidence),
                    linked_output_count=linked_output_count,
                    created_at=source_material.created_at,
                )
            )
            task_ids_with_materials.add(task.id)

        for artifact in artifacts:
            if artifact.id in artifact_seen:
                continue
            artifact_seen.add(artifact.id)
            linked_evidence = evidence_by_artifact.get(artifact.id, [])
            linked_output_count = sum(
                len(recommendation_targets_by_evidence.get(item.id, []))
                + len(risk_targets_by_evidence.get(item.id, []))
                + len(action_targets_by_evidence.get(item.id, []))
                for item in linked_evidence
            )
            role_label, presence_state = _material_role_label(
                len(linked_evidence),
                linked_output_count,
                None,
                artifact.continuity_scope,
                artifact.participation.participation_task_count
                if artifact.participation is not None
                else 0,
            )
            (
                participation_type,
                participation_task_count,
                current_task_participation,
                mapping_mode,
                canonical_owner_scope,
                compatibility_task_id,
            ) = _workspace_material_participation_fields(
                task_id=task.id,
                continuity_scope=artifact.continuity_scope,
                participation=artifact.participation,
            )
            artifact_cards.append(
                schemas.ArtifactEvidenceMaterialRead(
                    object_id=artifact.id,
                    task_id=task.id,
                    task_title=task_title,
                    object_type="artifact",
                    title=artifact.title,
                    summary=artifact.description or artifact.artifact_type,
                    role_label=role_label,
                    presence_state=presence_state,
                    continuity_scope=artifact.continuity_scope,
                    participation_type=participation_type,
                    participation_task_count=participation_task_count,
                    current_task_participation=current_task_participation,
                    canonical_owner_scope=canonical_owner_scope,
                    compatibility_task_id=compatibility_task_id,
                    mapping_mode=mapping_mode,
                    source_type=None,
                    ingest_status=None,
                    source_ref=None,
                    linked_evidence_count=len(linked_evidence),
                    linked_output_count=linked_output_count,
                    created_at=artifact.created_at,
                )
            )
            task_ids_with_materials.add(task.id)

        source_material_lookup = {item.id: item.title for item in source_materials}
        artifact_lookup = {item.id: item.title for item in artifacts}
        for item in evidence:
            if item.id in evidence_seen:
                continue
            evidence_seen.add(item.id)
            linked_recommendations = recommendation_targets_by_evidence.get(item.id, [])
            linked_risks = risk_targets_by_evidence.get(item.id, [])
            linked_action_items = action_targets_by_evidence.get(item.id, [])
            linked_deliverables = deliverable_targets_by_evidence.get(item.id, [])
            linked_output_total = (
                len(linked_recommendations) + len(linked_risks) + len(linked_action_items)
            )
            if linked_output_total > 0:
                linked_evidence_count += 1
            if linked_deliverables:
                linked_deliverable_count += 1
            task_ids_with_evidence.add(task.id)
            evidence_chains.append(
                schemas.EvidenceWorkspaceEvidenceRead(
                    evidence=item,
                    task_title=task_title,
                    source_material_title=(
                        source_material_lookup.get(item.source_material_id)
                        if item.source_material_id
                        else None
                    ),
                    artifact_title=(
                        artifact_lookup.get(item.artifact_id) if item.artifact_id else None
                    ),
                    strength_label=_evidence_strength_label(
                        linked_output_total,
                        len(linked_deliverables),
                        item.reliability_level,
                    ),
                    sufficiency_note=_build_evidence_support_note(
                        linked_recommendations,
                        linked_risks,
                        linked_action_items,
                        linked_deliverables,
                    ),
                    linked_recommendations=linked_recommendations,
                    linked_risks=linked_risks,
                    linked_action_items=linked_action_items,
                    linked_deliverables=linked_deliverables,
                )
            )

        if deliverable_class_hint != DeliverableClass.DECISION_ACTION_DELIVERABLE:
            deliverable_limitations.append(
                f"任務「{task.title}」目前仍屬 {deliverable_class_hint.value}，代表來源 / 證據鏈尚不足以穩定支撐更高等級的 decision / action deliverable。"
            )
        if external_research_heavy_candidate:
            deliverable_limitations.append(
                f"任務「{task.title}」目前屬 external-research-heavy sparse case，不能假裝已具備 company-specific certainty。"
            )

    if not source_material_cards:
        high_impact_gaps.append("目前案件世界仍缺可直接回看的 source materials。")
    if not artifact_cards:
        high_impact_gaps.append("目前案件世界仍缺可直接回看的 artifacts。")
    if not evidence_chains:
        high_impact_gaps.append("目前尚未形成可正式支撐 recommendation / risk / action 的 evidence chain。")
    if unsupported_recommendation_count > 0:
        high_impact_gaps.append(
            f"目前仍有 {unsupported_recommendation_count} 項 recommendations 缺乏正式 supporting evidence。"
        )
    if unsupported_risk_count > 0:
        high_impact_gaps.append(
            f"目前仍有 {unsupported_risk_count} 項 risks 缺乏正式 supporting evidence。"
        )
    if unsupported_action_item_count > 0:
        high_impact_gaps.append(
            f"目前仍有 {unsupported_action_item_count} 項 action items 缺乏正式 supporting evidence。"
        )

    unique_expectations = _unique_preserve_order(evidence_expectations)[:8]
    if unique_expectations and len(evidence_chains) < max(2, len(unique_expectations) // 2):
        high_impact_gaps.append(
            "目前已選 packs 對 evidence coverage 的期待仍高於現有來源厚度，建議優先補齊核心 source materials 與可引用 evidence。"
        )

    if evidence_chains:
        sufficiency_summary = (
            f"目前案件世界共有 {len(source_material_cards)} 份 source materials、"
            f"{len(artifact_cards)} 份 artifacts、{len(evidence_chains)} 則 evidence；"
            f"其中 {linked_evidence_count} 則 evidence 已正式支撐 recommendation / risk / action，"
            f"{linked_deliverable_count} 則 evidence 已被寫回 deliverables。"
        )
    elif source_material_cards or artifact_cards:
        sufficiency_summary = (
            "目前已形成來源與材料工作面，但 evidence-to-decision 支撐鏈仍偏薄，這輪較適合做 exploratory / review 等級判斷。"
        )
    else:
        sufficiency_summary = (
            "目前案件世界仍偏 sparse input，來源 / 證據工作面主要用來揭示缺口，而不是聲稱已有完整 evidence world。"
        )

    if len(task_ids_with_materials) > 1:
        continuity_notes.append(
            f"目前來源材料已跨 {len(task_ids_with_materials)} 個 tasks 累積，不再只是單次任務的附帶資料。"
        )
    if len(task_ids_with_evidence) > 1:
        continuity_notes.append(
            f"目前 evidence 已跨 {len(task_ids_with_evidence)} 個 tasks 被累積與回看，形成最小 cross-task continuity。"
        )
    if linked_deliverable_count > 0:
        continuity_notes.append(
            f"目前已有 {linked_deliverable_count} 則 evidence 正式回鏈到 deliverables，可直接回看支撐鏈。"
        )
    if external_research_heavy_detected:
        continuity_notes.append(
            "這個案件世界中包含 external-research-heavy sparse case；系統會保留 exploratory 層級誠實性，不假裝 company-specific certainty。"
        )

    source_material_cards = sorted(
        source_material_cards,
        key=lambda item: (item.linked_output_count, item.linked_evidence_count, item.created_at),
        reverse=True,
    )[:10]
    artifact_cards = sorted(
        artifact_cards,
        key=lambda item: (item.linked_output_count, item.linked_evidence_count, item.created_at),
        reverse=True,
    )[:10]
    evidence_chains = sorted(
        evidence_chains,
        key=lambda item: (
            len(item.linked_recommendations) + len(item.linked_risks) + len(item.linked_action_items),
            len(item.linked_deliverables),
            item.evidence.created_at,
        ),
        reverse=True,
    )[:12]

    return schemas.ArtifactEvidenceWorkspaceResponse(
        matter_summary=matter_summary,
        client=client,
        engagement=engagement,
        workstream=workstream,
        current_decision_context=current_decision_context,
        related_tasks=related_task_items,
        artifact_cards=artifact_cards,
        source_material_cards=source_material_cards,
        evidence_chains=evidence_chains,
        evidence_expectations=unique_expectations,
        high_impact_gaps=_unique_preserve_order(high_impact_gaps)[:8],
        evidence_gaps=sorted(evidence_gap_records, key=lambda item: item.updated_at, reverse=True)[:10],
        research_runs=sorted(research_runs, key=lambda item: item.started_at, reverse=True)[:6],
        sufficiency_summary=sufficiency_summary,
        deliverable_limitations=_unique_preserve_order(deliverable_limitations)[:6],
        continuity_notes=_unique_preserve_order(continuity_notes)[:6],
    )


def _get_content_record(deliverable: schemas.DeliverableRead) -> dict:
    return deliverable.content_structure if isinstance(deliverable.content_structure, dict) else {}


def _coerce_text(value: object) -> str:
    if isinstance(value, str):
        return value.strip()
    if isinstance(value, list):
        lines = [item.strip() for item in value if isinstance(item, str) and item.strip()]
        return "\n".join(lines)
    if isinstance(value, dict):
        for key in ("summary", "text", "content", "value"):
            nested = value.get(key)
            if isinstance(nested, str) and nested.strip():
                return nested.strip()
            if isinstance(nested, list):
                lines = [
                    item.strip()
                    for item in nested
                    if isinstance(item, str) and item.strip()
                ]
                if lines:
                    return "\n".join(lines)
    return ""


def _get_content_string(content: dict, key: str) -> str:
    value = content.get(key)
    return value.strip() if isinstance(value, str) else ""


def _get_content_string_list(content: dict, key: str) -> list[str]:
    value = content.get(key)
    if isinstance(value, list):
        return [item.strip() for item in value if isinstance(item, str) and item.strip()]
    return []


def _dedupe_schema_items(items: list):
    seen_ids: set[str] = set()
    deduped: list = []
    for item in items:
        item_id = getattr(item, "id", None)
        if not item_id or item_id in seen_ids:
            continue
        seen_ids.add(item_id)
        deduped.append(item)
    return deduped


def _resolve_deliverable_class_for_workspace(
    task: schemas.TaskAggregateResponse,
    deliverable: schemas.DeliverableRead,
) -> DeliverableClass:
    content = _get_content_record(deliverable)
    operating_state = content.get("sparse_input_operating_state")
    if isinstance(operating_state, dict):
        raw_class = operating_state.get("deliverable_class")
        if isinstance(raw_class, str):
            try:
                return DeliverableClass(raw_class)
            except ValueError:
                pass

    readiness = content.get("readiness_governance")
    if isinstance(readiness, dict):
        raw_class = readiness.get("supported_deliverable_class")
        if isinstance(raw_class, str):
            try:
                return DeliverableClass(raw_class)
            except ValueError:
                pass

    return task.deliverable_class_hint


def _build_deliverable_confidence_summary(
    deliverable_class: DeliverableClass,
    linked_evidence_count: int,
    high_impact_gaps: list[str],
    external_research_heavy_candidate: bool,
) -> str:
    if external_research_heavy_candidate or deliverable_class == DeliverableClass.EXPLORATORY_BRIEF:
        return (
            "這份交付物目前屬於 exploratory level，重點是建立可用判斷與後續補證方向，"
            "不應被誤讀成已完整對齊 company-specific certainty 的正式決策備忘錄。"
        )
    if deliverable_class == DeliverableClass.ASSESSMENT_REVIEW_MEMO:
        return (
            "這份交付物目前屬於 assessment / review memo，判斷主要建立在既有文件、來源材料與可回看的 evidence 上，"
            "適合支撐 review / challenge，但未必足以直接推進完整行動方案。"
        )
    if high_impact_gaps:
        return (
            f"這份 decision / action deliverable 已具備 {linked_evidence_count} 則正式 evidence 支撐，"
            f"但仍有 {len(high_impact_gaps)} 項高影響缺口需要在採行前一併考慮。"
        )
    return (
        f"這份 decision / action deliverable 目前已有 {linked_evidence_count} 則正式 evidence 支撐，"
        "可作為較完整的決策 / 行動級交付物回看。"
    )


def get_deliverable_workspace(
    db: Session,
    deliverable_id: str,
) -> schemas.DeliverableWorkspaceResponse:
    deliverable_row = db.scalars(
        select(models.Deliverable).where(models.Deliverable.id == deliverable_id)
    ).one_or_none()
    if deliverable_row is None:
        raise HTTPException(status_code=404, detail="找不到指定交付物工作面。")

    task = get_loaded_task(db, deliverable_row.task_id)
    task_aggregate = serialize_task(task)
    deliverable = next(
        (item for item in task_aggregate.deliverables if item.id == deliverable_id),
        None,
    )
    if deliverable is None:
        raise HTTPException(status_code=404, detail="找不到指定交付物工作面。")

    version_events = [
        schemas.DeliverableVersionEventRead.model_validate(item)
        for item in ensure_deliverable_version_events(
            db,
            deliverable_row,
            fallback_status=deliverable.status,
        )
    ]
    content_revisions = ensure_deliverable_content_revisions(db, deliverable_row)
    artifact_records, publish_records = ensure_deliverable_release_records(
        db,
        deliverable_row,
        fallback_status=deliverable.status,
    )

    content = _get_content_record(deliverable)
    readiness_governance = content.get("readiness_governance")
    if not isinstance(readiness_governance, dict):
        readiness_governance = {}

    deliverable_class = _resolve_deliverable_class_for_workspace(task_aggregate, deliverable)
    latest_deliverable = max(
        task_aggregate.deliverables,
        key=lambda item: (item.version, item.generated_at),
    )
    is_latest_for_task = latest_deliverable.id == deliverable.id
    workspace_status = "current" if is_latest_for_task else "superseded"

    linked_ids_by_type: dict[str, set[str]] = {}
    for link in deliverable.linked_objects:
        if not link.object_id:
            continue
        linked_ids_by_type.setdefault(link.object_type, set()).add(link.object_id)

    linked_recommendations = [
        item
        for item in task_aggregate.recommendations
        if item.id in linked_ids_by_type.get("recommendation", set())
    ]
    linked_risks = [
        item for item in task_aggregate.risks if item.id in linked_ids_by_type.get("risk", set())
    ]
    linked_action_items = [
        item
        for item in task_aggregate.action_items
        if item.id in linked_ids_by_type.get("action_item", set())
    ]

    if not linked_recommendations:
        linked_recommendations = task_aggregate.recommendations
    if not linked_risks:
        linked_risks = task_aggregate.risks
    if not linked_action_items:
        linked_action_items = task_aggregate.action_items

    evidence_ids = set(linked_ids_by_type.get("evidence", set()))
    for item in [*linked_recommendations, *linked_risks, *linked_action_items]:
        evidence_ids.update(getattr(item, "supporting_evidence_ids", []))

    linked_evidence = [
        item for item in task_aggregate.evidence if item.id in evidence_ids
    ] or task_aggregate.evidence
    linked_evidence = _dedupe_schema_items(linked_evidence)

    source_material_ids = set(linked_ids_by_type.get("source_material", set()))
    artifact_ids = set(linked_ids_by_type.get("artifact", set()))
    for evidence in linked_evidence:
        if evidence.source_material_id:
            source_material_ids.add(evidence.source_material_id)
        if evidence.artifact_id:
            artifact_ids.add(evidence.artifact_id)

    linked_source_materials = [
        item for item in task_aggregate.source_materials if item.id in source_material_ids
    ] or task_aggregate.source_materials
    linked_artifacts = [
        item for item in task_aggregate.artifacts if item.id in artifact_ids
    ] or task_aggregate.artifacts

    high_impact_gaps = _unique_preserve_order(
        [
            *[
                item
                for item in readiness_governance.get("pack_high_impact_gaps", [])
                if isinstance(item, str) and item.strip()
            ],
            *_get_content_string_list(content, "missing_information"),
        ]
    )
    deliverable_guidance = (
        readiness_governance.get("deliverable_guidance", "")
        if isinstance(readiness_governance.get("deliverable_guidance", ""), str)
        else ""
    )
    limitation_notes = _unique_preserve_order(
        [
            deliverable_guidance,
            task_aggregate.sparse_input_summary,
            *high_impact_gaps[:4],
        ]
    )
    confidence_summary = _build_deliverable_confidence_summary(
        deliverable_class,
        len(linked_evidence),
        high_impact_gaps,
        bool(task_aggregate.external_research_heavy_candidate),
    )

    related_deliverables: list[schemas.MatterDeliverableSummaryRead] = []
    continuity_notes: list[str] = []
    if task_aggregate.matter_workspace:
        matter_workspace = get_matter_workspace(db, task_aggregate.matter_workspace.id)
        related_deliverables = [
            item
            for item in matter_workspace.related_deliverables
            if item.deliverable_id != deliverable.id
        ][:6]
        continuity_notes.extend(matter_workspace.continuity_notes)

    if is_latest_for_task:
        continuity_notes.append("這份 deliverable 目前是此 task 最新版本的正式交付物。")
    else:
        continuity_notes.append("這份 deliverable 仍可回看，但已不是此 task 的最新版本。")
    if related_deliverables:
        continuity_notes.append(
            f"同一個案件世界內另外還有 {len(related_deliverables)} 份 related deliverables 可交叉回看。"
        )
    if not linked_evidence:
        continuity_notes.append("這份交付物目前尚未形成厚實 evidence basis，應先回看來源 / 證據工作面再採行。")

    return schemas.DeliverableWorkspaceResponse(
        deliverable=deliverable,
        task=task_aggregate,
        matter_workspace=task_aggregate.matter_workspace,
        deliverable_class=deliverable_class,
        workspace_status=workspace_status,
        is_latest_for_task=is_latest_for_task,
        confidence_summary=confidence_summary,
        deliverable_guidance=deliverable_guidance,
        high_impact_gaps=high_impact_gaps[:8],
        limitation_notes=limitation_notes[:8],
        linked_source_materials=_dedupe_schema_items(linked_source_materials)[:8],
        linked_artifacts=_dedupe_schema_items(linked_artifacts)[:8],
        linked_evidence=_dedupe_schema_items(linked_evidence)[:8],
        linked_recommendations=_dedupe_schema_items(linked_recommendations)[:8],
        linked_risks=_dedupe_schema_items(linked_risks)[:8],
        linked_action_items=_dedupe_schema_items(linked_action_items)[:8],
        related_deliverables=related_deliverables,
        decision_records=task_aggregate.decision_records[:6],
        action_plans=task_aggregate.action_plans[:6],
        action_executions=task_aggregate.action_executions[:10],
        outcome_records=task_aggregate.outcome_records[:10],
        research_runs=task_aggregate.research_runs[:6],
        continuity_notes=_unique_preserve_order(continuity_notes),
        content_sections=_normalize_deliverable_content_sections(deliverable_row.content_sections),
        content_revisions=[
            _build_deliverable_content_revision_read(item) for item in content_revisions[:12]
        ],
        version_events=version_events,
        artifact_records=[
            schemas.DeliverableArtifactRecordRead.model_validate(item)
            for item in artifact_records[:12]
        ],
        publish_records=[
            schemas.DeliverablePublishRecordRead.model_validate(item)
            for item in publish_records[:8]
        ],
    )


def _load_deliverable_and_task(
    db: Session,
    deliverable_id: str,
) -> tuple[models.Deliverable, models.Task]:
    deliverable = db.scalars(
        select(models.Deliverable).where(models.Deliverable.id == deliverable_id)
    ).one_or_none()
    if deliverable is None:
        raise HTTPException(status_code=404, detail="找不到指定交付物工作面。")

    task = db.scalars(select(models.Task).where(models.Task.id == deliverable.task_id)).one_or_none()
    if task is None:
        raise HTTPException(status_code=404, detail="找不到交付物對應的任務。")

    return deliverable, task


def _apply_deliverable_workspace_update(
    db: Session,
    deliverable: models.Deliverable,
    task: models.Task,
    *,
    title: str,
    summary: str,
    status: str,
    version_tag: str,
    content_sections: dict[str, object] | None = None,
    event_note: str = "",
) -> dict[str, object]:
    previous_title = deliverable.title
    previous_summary = deliverable.summary or ""
    previous_status = _default_deliverable_status(task, deliverable)
    previous_version_tag = default_deliverable_version_tag(deliverable)
    previous_content_sections = _normalize_deliverable_content_sections(deliverable.content_sections)

    next_title = title.strip() or deliverable.title
    next_summary = summary.strip()
    next_status = status.strip() or previous_status
    next_version_tag = version_tag.strip() or previous_version_tag
    next_content_sections = _normalize_deliverable_content_sections(content_sections or {})
    next_content_sections_payload = _serialize_deliverable_content_sections(next_content_sections)

    title_changed = next_title != previous_title
    summary_changed = next_summary != previous_summary
    status_changed = next_status != previous_status
    version_tag_changed = next_version_tag != previous_version_tag
    content_changed = next_content_sections_payload != _serialize_deliverable_content_sections(
        previous_content_sections
    )
    has_changes = (
        title_changed
        or summary_changed
        or status_changed
        or version_tag_changed
        or content_changed
        or bool(event_note.strip())
    )

    deliverable.title = next_title
    deliverable.summary = next_summary
    deliverable.status = next_status
    deliverable.version_tag = next_version_tag
    if content_sections is not None:
        deliverable.content_sections = next_content_sections_payload
    if has_changes:
        task.updated_at = models.utc_now()
    db.add_all([deliverable, task])
    db.flush()

    if title_changed or summary_changed:
        changed_fields: list[str] = []
        if title_changed:
            changed_fields.append("標題")
        if summary_changed:
            changed_fields.append("簡短摘要")
        record_deliverable_version_event(
            db,
            deliverable,
            DELIVERABLE_EVENT_CONTENT_UPDATED,
            version_tag=next_version_tag,
            deliverable_status=next_status,
            summary=f"更新交付物{'與'.join(changed_fields)}",
            event_payload={
                "source": DELIVERABLE_EVENT_SOURCE_METADATA_UPDATE,
                "changed_fields": changed_fields,
                "previous_title": previous_title,
                "previous_summary": previous_summary,
                "next_title": next_title,
                "next_summary": next_summary,
            },
        )

    if content_changed:
        changed_sections: list[str] = []
        if next_content_sections.executive_summary != previous_content_sections.executive_summary:
            changed_sections.append("摘要")
        if next_content_sections.recommendations != previous_content_sections.recommendations:
            changed_sections.append("建議")
        if next_content_sections.risks != previous_content_sections.risks:
            changed_sections.append("風險")
        if next_content_sections.action_items != previous_content_sections.action_items:
            changed_sections.append("行動項目")
        if next_content_sections.evidence_basis != previous_content_sections.evidence_basis:
            changed_sections.append("依據來源")
        content_revision = create_deliverable_content_revision(
            db,
            deliverable,
            snapshot=next_content_sections_payload,
            previous_snapshot=_serialize_deliverable_content_sections(previous_content_sections),
            source=CONTENT_REVISION_SOURCE_MANUAL_EDIT,
        )
        content_event = record_deliverable_version_event(
            db,
            deliverable,
            DELIVERABLE_EVENT_CONTENT_UPDATED,
            version_tag=next_version_tag,
            deliverable_status=next_status,
            summary=f"更新交付物正文：{'、'.join(changed_sections)}",
            event_payload={
                "source": DELIVERABLE_EVENT_SOURCE_CONTENT_UPDATE,
                "changed_sections": changed_sections,
                "content_scope": "workspace_sections",
            },
        )
        if content_revision is not None:
            content_revision.source_version_event_id = content_event.id
            db.add(content_revision)
            content_event.event_payload = {
                **(content_event.event_payload if isinstance(content_event.event_payload, dict) else {}),
                "revision_id": content_revision.id,
            }
            db.add(content_event)

    if version_tag_changed:
        record_deliverable_version_event(
            db,
            deliverable,
            DELIVERABLE_EVENT_VERSION_TAG_UPDATED,
            version_tag=next_version_tag,
            deliverable_status=next_status,
            summary=f"版本標記更新為 {next_version_tag}",
            event_payload={
                "source": DELIVERABLE_EVENT_SOURCE_METADATA_UPDATE,
                "previous_version_tag": previous_version_tag,
                "next_version_tag": next_version_tag,
            },
        )

    if status_changed:
        record_deliverable_version_event(
            db,
            deliverable,
            DELIVERABLE_EVENT_STATUS_CHANGED,
            version_tag=next_version_tag,
            deliverable_status=next_status,
            summary=(
                f"狀態由 {label_for_deliverable_status(previous_status)} "
                f"變更為 {label_for_deliverable_status(next_status)}"
            ),
            event_payload={
                "source": DELIVERABLE_EVENT_SOURCE_METADATA_UPDATE,
                "from_status": previous_status,
                "to_status": next_status,
            },
        )

    normalized_event_note = event_note.strip()
    if normalized_event_note:
        record_deliverable_version_event(
            db,
            deliverable,
            DELIVERABLE_EVENT_NOTE_ADDED,
            version_tag=next_version_tag,
            deliverable_status=next_status,
            summary=normalized_event_note[:280],
            event_payload={
                "source": DELIVERABLE_EVENT_SOURCE_METADATA_UPDATE,
                "note": normalized_event_note,
            },
        )

    return {
        "deliverable": deliverable,
        "task": task,
        "version_tag": next_version_tag,
        "status": next_status,
        "status_changed": status_changed,
        "version_tag_changed": version_tag_changed,
        "content_changed": content_changed,
        "event_note": normalized_event_note,
        "previous_status": previous_status,
    }


def update_deliverable_workspace(
    db: Session,
    deliverable_id: str,
    payload: schemas.DeliverableWorkspaceUpdateRequest,
) -> schemas.DeliverableWorkspaceResponse:
    deliverable, task = _load_deliverable_and_task(db, deliverable_id)
    _apply_deliverable_workspace_update(
        db,
        deliverable,
        task,
        title=payload.title,
        summary=payload.summary,
        status=payload.status,
        version_tag=payload.version_tag,
        content_sections=_serialize_deliverable_content_sections(payload.content_sections),
        event_note=payload.event_note,
    )
    db.commit()
    return get_deliverable_workspace(db, deliverable_id)


def update_deliverable_metadata(
    db: Session,
    deliverable_id: str,
    payload: schemas.DeliverableMetadataUpdateRequest,
) -> schemas.DeliverableWorkspaceResponse:
    deliverable, _ = _load_deliverable_and_task(db, deliverable_id)
    current_sections = _normalize_deliverable_content_sections(deliverable.content_sections)
    return update_deliverable_workspace(
        db,
        deliverable_id,
        schemas.DeliverableWorkspaceUpdateRequest(
            title=payload.title,
            summary=payload.summary,
            status=payload.status,
            version_tag=payload.version_tag,
            event_note=payload.event_note,
            content_sections=schemas.DeliverableContentSectionsRequest(
                executive_summary=current_sections.executive_summary,
                recommendations=current_sections.recommendations,
                risks=current_sections.risks,
                action_items=current_sections.action_items,
                evidence_basis=current_sections.evidence_basis,
            ),
        ),
    )


def rollback_deliverable_content_revision(
    db: Session,
    deliverable_id: str,
    revision_id: str,
) -> schemas.DeliverableWorkspaceResponse:
    deliverable, task = _load_deliverable_and_task(db, deliverable_id)
    target_revision = db.scalars(
        select(models.DeliverableContentRevision).where(
            models.DeliverableContentRevision.id == revision_id,
            models.DeliverableContentRevision.deliverable_id == deliverable_id,
        )
    ).one_or_none()
    if target_revision is None:
        raise HTTPException(status_code=404, detail="找不到指定交付物正文修訂。")

    current_snapshot = _normalize_deliverable_content_sections(deliverable.content_sections)
    target_snapshot = _normalize_deliverable_content_sections(target_revision.snapshot)
    if target_snapshot == current_snapshot:
        raise HTTPException(status_code=400, detail="目前交付物正文已與目標修訂相同。")

    serialized_target_snapshot = _serialize_deliverable_content_sections(target_snapshot)
    deliverable.content_sections = serialized_target_snapshot
    task.updated_at = models.utc_now()
    db.add_all([deliverable, task])
    db.flush()

    content_revision = create_deliverable_content_revision(
        db,
        deliverable,
        snapshot=serialized_target_snapshot,
        previous_snapshot=_serialize_deliverable_content_sections(current_snapshot),
        source=CONTENT_REVISION_SOURCE_ROLLBACK,
        revision_summary=f"回退交付物正文到 {target_revision.created_at.isoformat()} 的修訂",
        rollback_target_revision_id=target_revision.id,
    )
    rollback_event = record_deliverable_version_event(
        db,
        deliverable,
        DELIVERABLE_EVENT_CONTENT_ROLLED_BACK,
        version_tag=default_deliverable_version_tag(deliverable),
        deliverable_status=deliverable.status,
        summary=f"回退交付物正文到 {target_revision.created_at.isoformat()} 的修訂",
        event_payload={
            "source": DELIVERABLE_EVENT_SOURCE_CONTENT_ROLLBACK,
            "rollback_target_revision_id": target_revision.id,
            "revision_id": content_revision.id if content_revision else None,
            "changed_sections": content_revision.changed_sections if content_revision else [],
            "content_scope": "workspace_sections",
        },
    )
    if content_revision is not None:
        content_revision.source_version_event_id = rollback_event.id
        db.add(content_revision)

    db.commit()
    return get_deliverable_workspace(db, deliverable_id)


def _build_deliverable_export_filename(
    title: str,
    version_tag: str,
    extension: str,
) -> str:
    filename_base = re.sub(
        r"[^a-z0-9]+",
        "-",
        f"{title}-{version_tag}".lower(),
    ).strip("-")
    return f"{filename_base or 'deliverable-export'}.{extension}"


def _build_resolved_deliverable_content_sections(
    workspace: schemas.DeliverableWorkspaceResponse,
) -> schemas.DeliverableContentSectionsRead:
    stored = workspace.content_sections
    deliverable = workspace.deliverable

    return schemas.DeliverableContentSectionsRead(
        executive_summary=(
            stored.executive_summary
            or _coerce_text(_get_content_record(deliverable).get("executive_summary"))
            or deliverable.summary
            or deliverable.title
        ),
        recommendations=stored.recommendations
        or [item.summary for item in workspace.linked_recommendations[:8]],
        risks=stored.risks
        or [
            *[f"{item.title}：{item.description}" for item in workspace.linked_risks[:8]],
            *[f"高影響缺口：{item}" for item in workspace.high_impact_gaps[:5]],
        ],
        action_items=stored.action_items
        or [item.description for item in workspace.linked_action_items[:8]],
        evidence_basis=stored.evidence_basis
        or [
            f"{item.title}｜{item.source_type or item.evidence_type or 'evidence'}"
            for item in workspace.linked_evidence[:8]
        ],
    )


def _build_deliverable_export_lists(
    workspace: schemas.DeliverableWorkspaceResponse,
) -> tuple[str, str, str, str, str, list[str], list[str], list[str], list[str], list[str]]:
    deliverable = workspace.deliverable
    task = workspace.task
    version_tag = deliverable.version_tag or f"v{deliverable.version}"
    matter_title = workspace.matter_workspace.title if workspace.matter_workspace else "未掛案件"
    matter_path = workspace.matter_workspace.object_path if workspace.matter_workspace else "未掛案件"
    content_sections = _build_resolved_deliverable_content_sections(workspace)
    executive_summary = content_sections.executive_summary

    def _bullet_lines(items: list[str], empty_text: str) -> list[str]:
        if not items:
            return [empty_text]
        return items

    recommendation_lines = _bullet_lines(
        content_sections.recommendations,
        "目前沒有額外的建議項目。",
    )
    risk_lines = _bullet_lines(
        content_sections.risks,
        "目前沒有額外的風險與缺口項目。",
    )
    action_item_lines = _bullet_lines(
        content_sections.action_items,
        "目前沒有額外的行動項目。",
    )
    evidence_lines = _bullet_lines(
        content_sections.evidence_basis,
        "目前沒有額外的依據來源。",
    )
    version_event_lines = _bullet_lines(
        [
            f"{item.created_at.isoformat()}｜{item.version_tag}｜{item.summary}"
            for item in workspace.version_events[:8]
        ],
        "目前沒有額外的版本事件。",
    )
    return (
        version_tag,
        matter_title,
        matter_path,
        executive_summary,
        task.updated_at.isoformat(),
        recommendation_lines,
        risk_lines,
        action_item_lines,
        evidence_lines,
        version_event_lines,
    )


def publish_deliverable_release(
    db: Session,
    deliverable_id: str,
    payload: schemas.DeliverablePublishRequest,
) -> schemas.DeliverableWorkspaceResponse:
    deliverable, task = _load_deliverable_and_task(db, deliverable_id)
    update_result = _apply_deliverable_workspace_update(
        db,
        deliverable,
        task,
        title=payload.title,
        summary=payload.summary,
        status="final",
        version_tag=payload.version_tag,
        content_sections=_serialize_deliverable_content_sections(payload.content_sections),
        event_note="",
    )

    workspace = get_deliverable_workspace(db, deliverable_id)
    version_tag = str(update_result["version_tag"])
    publish_note = payload.publish_note.strip()
    artifact_formats = _unique_preserve_order(
        [item.strip() for item in payload.artifact_formats if item.strip()]
    ) or ["markdown", "docx"]

    rendered_artifacts: list[tuple[str, str, str, bytes]] = []
    markdown_filename, markdown_content, _ = _render_deliverable_markdown_artifact(workspace)
    if "markdown" in artifact_formats:
        rendered_artifacts.append(
            (
                "markdown",
                "text/markdown; charset=utf-8",
                markdown_filename,
                markdown_content.encode("utf-8"),
            )
        )
    if "docx" in artifact_formats:
        docx_filename, docx_content, _ = _render_deliverable_docx_artifact(workspace)
        rendered_artifacts.append(
            (
                "docx",
                "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                docx_filename,
                docx_content,
            )
        )

    publish_event = record_deliverable_version_event(
        db,
        deliverable,
        DELIVERABLE_EVENT_PUBLISHED,
        version_tag=version_tag,
        deliverable_status="final",
        summary=f"正式發布 {version_tag} 定稿版",
        event_payload={
            "source": DELIVERABLE_EVENT_SOURCE_PUBLISH,
            "artifact_formats": artifact_formats,
            "publish_note": publish_note,
        },
    )
    publish_record = create_deliverable_publish_record(
        db,
        deliverable,
        version_tag=version_tag,
        deliverable_status="final",
        publish_note=publish_note,
        artifact_formats=artifact_formats,
        source_version_event=publish_event,
    )

    artifact_record_ids: list[str] = []
    for artifact_format, mime_type, file_name, content_bytes in rendered_artifacts:
        artifact_record = create_deliverable_artifact_record(
            db,
            deliverable,
            artifact_kind=DELIVERABLE_ARTIFACT_KIND_RELEASE,
            artifact_format=artifact_format,
            version_tag=version_tag,
            deliverable_status="final",
            file_name=file_name,
            mime_type=mime_type,
            content_bytes=content_bytes,
            publish_record=publish_record,
        )
        artifact_record_ids.append(artifact_record.id)

    publish_event.event_payload = {
        "source": DELIVERABLE_EVENT_SOURCE_PUBLISH,
        "artifact_formats": artifact_formats,
        "artifact_record_ids": artifact_record_ids,
        "publish_record_id": publish_record.id,
        "publish_note": publish_note,
        "status_changed": bool(update_result["status_changed"]),
        "version_tag_changed": bool(update_result["version_tag_changed"]),
    }
    db.add(publish_event)
    db.commit()

    return get_deliverable_workspace(db, deliverable_id)


def _render_deliverable_markdown_artifact(
    workspace: schemas.DeliverableWorkspaceResponse,
) -> tuple[str, str, str]:
    deliverable = workspace.deliverable
    (
        version_tag,
        matter_title,
        matter_path,
        executive_summary,
        updated_at,
        recommendation_lines,
        risk_lines,
        action_item_lines,
        evidence_lines,
        version_event_lines,
    ) = _build_deliverable_export_lists(workspace)

    publish_record = workspace.publish_records[0] if workspace.publish_records else None
    latest_publish_note = publish_record.publish_note if publish_record else ""

    content = "\n".join(
        [
            f"# {deliverable.title}",
            "",
            "## 匯出資訊",
            f"- 版本：{version_tag}",
            f"- 狀態：{label_for_deliverable_status(deliverable.status)}",
            f"- 交付物類型：{deliverable.deliverable_type}",
            f"- 案件：{matter_title}",
            f"- 路徑：{matter_path}",
            f"- 最近更新：{updated_at}",
            *( [f"- 最近發布說明：{latest_publish_note}"] if latest_publish_note else [] ),
            "",
            "## 簡短摘要",
            executive_summary,
            "",
            "## 一句話結論",
            workspace.confidence_summary,
            "",
            "## 建議",
            *[f"- {item}" for item in recommendation_lines],
            "",
            "## 風險與缺口",
            *[f"- {item}" for item in risk_lines],
            "",
            "## 行動項目",
            *[f"- {item}" for item in action_item_lines],
            "",
            "## 依據來源",
            *[f"- {item}" for item in evidence_lines],
            "",
            "## 版本事件",
            *[f"- {item}" for item in version_event_lines],
            "",
        ]
    )
    filename = _build_deliverable_export_filename(deliverable.title, version_tag, "md")
    return filename, content, version_tag


def build_deliverable_markdown_export(
    db: Session,
    deliverable_id: str,
) -> tuple[str, str, str]:
    workspace = get_deliverable_workspace(db, deliverable_id)
    deliverable_row, _ = _load_deliverable_and_task(db, deliverable_id)
    filename, content, version_tag = _render_deliverable_markdown_artifact(workspace)
    export_event = record_deliverable_version_event(
        db,
        deliverable_row,
        DELIVERABLE_EVENT_EXPORTED,
        version_tag=version_tag,
        deliverable_status=workspace.deliverable.status,
        summary=f"匯出 Markdown｜{version_tag}",
        event_payload={
            "source": DELIVERABLE_EVENT_SOURCE_EXPORT,
            "artifact_kind": DELIVERABLE_ARTIFACT_KIND_EXPORT,
            "artifact_format": "markdown",
            "file_name": filename,
            "mime_type": "text/markdown; charset=utf-8",
        },
    )
    artifact_record = create_deliverable_artifact_record(
        db,
        deliverable_row,
        artifact_kind=DELIVERABLE_ARTIFACT_KIND_EXPORT,
        artifact_format="markdown",
        version_tag=version_tag,
        deliverable_status=workspace.deliverable.status,
        file_name=filename,
        mime_type="text/markdown; charset=utf-8",
        content_bytes=content.encode("utf-8"),
        source_version_event=export_event,
    )
    export_event.event_payload = {
        "source": DELIVERABLE_EVENT_SOURCE_EXPORT,
        "artifact_kind": DELIVERABLE_ARTIFACT_KIND_EXPORT,
        "artifact_format": "markdown",
        "file_name": filename,
        "mime_type": "text/markdown; charset=utf-8",
        "artifact_record_id": artifact_record.id,
        "artifact_key": artifact_record.artifact_key,
    }
    db.add(export_event)
    db.commit()
    return filename, content, version_tag


def _render_deliverable_docx_artifact(
    workspace: schemas.DeliverableWorkspaceResponse,
) -> tuple[str, bytes, str]:
    deliverable = workspace.deliverable
    (
        version_tag,
        matter_title,
        matter_path,
        executive_summary,
        updated_at,
        recommendation_lines,
        risk_lines,
        action_item_lines,
        evidence_lines,
        version_event_lines,
    ) = _build_deliverable_export_lists(workspace)

    try:
        from docx import Document
    except ImportError as error:  # pragma: no cover - depends on runtime env
        raise HTTPException(
            status_code=503,
            detail="目前的後端環境尚未安裝 docx 匯出依賴。",
        ) from error

    document = Document()
    document.add_heading(deliverable.title, level=0)
    metadata_lines = [
        f"版本：{version_tag}",
        f"狀態：{label_for_deliverable_status(deliverable.status)}",
        f"交付物類型：{deliverable.deliverable_type}",
        f"案件：{matter_title}",
        f"路徑：{matter_path}",
        f"最近更新：{updated_at}",
    ]
    for line in metadata_lines:
        document.add_paragraph(line)

    def _add_section(title: str, lines: list[str] | None = None, paragraph: str | None = None) -> None:
        document.add_heading(title, level=1)
        if paragraph:
            document.add_paragraph(paragraph)
        for line in lines or []:
            document.add_paragraph(line, style="List Bullet")

    _add_section("簡短摘要", paragraph=executive_summary)
    _add_section("一句話結論", paragraph=workspace.confidence_summary)
    _add_section("建議與風險", lines=[*recommendation_lines, *risk_lines])
    _add_section("行動項目", lines=action_item_lines)
    _add_section("依據來源", lines=evidence_lines)
    _add_section("版本事件", lines=version_event_lines)

    buffer = BytesIO()
    document.save(buffer)
    filename = _build_deliverable_export_filename(deliverable.title, version_tag, "docx")
    return filename, buffer.getvalue(), version_tag


def build_deliverable_docx_export(
    db: Session,
    deliverable_id: str,
) -> tuple[str, bytes, str]:
    workspace = get_deliverable_workspace(db, deliverable_id)
    deliverable_row, _ = _load_deliverable_and_task(db, deliverable_id)
    filename, content, version_tag = _render_deliverable_docx_artifact(workspace)
    export_event = record_deliverable_version_event(
        db,
        deliverable_row,
        DELIVERABLE_EVENT_EXPORTED,
        version_tag=version_tag,
        deliverable_status=workspace.deliverable.status,
        summary=f"匯出 DOCX｜{version_tag}",
        event_payload={
            "source": DELIVERABLE_EVENT_SOURCE_EXPORT,
            "artifact_kind": DELIVERABLE_ARTIFACT_KIND_EXPORT,
            "artifact_format": "docx",
            "file_name": filename,
            "mime_type": "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
        },
    )
    artifact_record = create_deliverable_artifact_record(
        db,
        deliverable_row,
        artifact_kind=DELIVERABLE_ARTIFACT_KIND_EXPORT,
        artifact_format="docx",
        version_tag=version_tag,
        deliverable_status=workspace.deliverable.status,
        file_name=filename,
        mime_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
        content_bytes=content,
        source_version_event=export_event,
    )
    export_event.event_payload = {
        "source": DELIVERABLE_EVENT_SOURCE_EXPORT,
        "artifact_kind": DELIVERABLE_ARTIFACT_KIND_EXPORT,
        "artifact_format": "docx",
        "file_name": filename,
        "mime_type": "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
        "artifact_record_id": artifact_record.id,
        "artifact_key": artifact_record.artifact_key,
    }
    db.add(export_event)
    db.commit()
    return filename, content, version_tag


def download_deliverable_artifact(
    db: Session,
    deliverable_id: str,
    artifact_id: str,
) -> tuple[str, bytes, str, str, str]:
    artifact_record = db.scalars(
        select(models.DeliverableArtifactRecord).where(
            models.DeliverableArtifactRecord.id == artifact_id,
            models.DeliverableArtifactRecord.deliverable_id == deliverable_id,
        )
    ).one_or_none()
    if artifact_record is None:
        raise HTTPException(status_code=404, detail="找不到指定 artifact。")
    content_bytes = load_artifact_content(artifact_record)
    if artifact_record.availability_state != "available" or content_bytes is None:
        raise HTTPException(
            status_code=409,
            detail="這份 artifact 目前只有 metadata backfill，尚無可下載的正式內容。",
        )

    return (
        artifact_record.file_name,
        content_bytes,
        artifact_record.version_tag,
        artifact_record.artifact_format,
        artifact_record.mime_type,
    )


def _serialize_case_world_draft(
    draft: models.CaseWorldDraft | None,
) -> schemas.CaseWorldDraftRead | None:
    if draft is None:
        return None

    return schemas.CaseWorldDraftRead(
        id=draft.id,
        task_id=draft.task_id,
        matter_workspace_id=draft.matter_workspace_id,
        compiler_status=draft.compiler_status,
        entry_preset=InputEntryMode(draft.entry_preset),
        continuity_mode=EngagementContinuityMode(draft.continuity_mode),
        writeback_depth=WritebackDepth(draft.writeback_depth),
        canonical_intake_summary=draft.canonical_intake_summary or {},
        task_interpretation=draft.task_interpretation or {},
        decision_context=draft.decision_context_payload or {},
        extracted_objects=list(draft.extracted_objects or []),
        inferred_links=list(draft.inferred_links or []),
        facts=[
            schemas.CaseWorldFactRead(**item)
            for item in list(draft.facts or [])
            if isinstance(item, dict)
        ],
        assumptions=[
            schemas.CaseWorldAssumptionRead(**item)
            for item in list(draft.assumptions_payload or [])
            if isinstance(item, dict)
        ],
        evidence_gaps=[
            schemas.CaseWorldGapRead(**item)
            for item in list(draft.evidence_gaps_payload or [])
            if isinstance(item, dict)
        ],
        suggested_capabilities=list(draft.suggested_capabilities or []),
        suggested_domain_packs=list(draft.suggested_domain_packs or []),
        suggested_industry_packs=list(draft.suggested_industry_packs or []),
        suggested_agents=list(draft.suggested_agents or []),
        suggested_research_need=bool(draft.suggested_research_need),
        next_best_actions=list(draft.next_best_actions or []),
        created_at=draft.created_at,
        updated_at=draft.updated_at,
    )


def _serialize_case_world_state(
    state: models.CaseWorldState | None,
) -> schemas.CaseWorldStateRead | None:
    if state is None:
        return None

    return schemas.CaseWorldStateRead(
        id=state.id,
        matter_workspace_id=state.matter_workspace_id,
        compiler_status=state.compiler_status,
        world_status=state.world_status,
        client_id=state.client_id,
        engagement_id=state.engagement_id,
        workstream_id=state.workstream_id,
        decision_context_id=state.decision_context_id,
        entry_preset=InputEntryMode(state.entry_preset),
        continuity_mode=EngagementContinuityMode(state.continuity_mode),
        writeback_depth=WritebackDepth(state.writeback_depth),
        world_identity=state.world_identity_payload or {},
        canonical_intake_summary=state.canonical_intake_summary or {},
        decision_context=state.decision_context_payload or {},
        extracted_objects=list(state.extracted_objects or []),
        inferred_links=list(state.inferred_links or []),
        facts=[
            schemas.CaseWorldFactRead(**item)
            for item in list(state.facts or [])
            if isinstance(item, dict)
        ],
        assumptions=[
            schemas.CaseWorldAssumptionRead(**item)
            for item in list(state.assumptions_payload or [])
            if isinstance(item, dict)
        ],
        evidence_gaps=[
            schemas.CaseWorldGapRead(**item)
            for item in list(state.evidence_gaps_payload or [])
            if isinstance(item, dict)
        ],
        selected_capabilities=list(state.selected_capabilities or []),
        selected_domain_packs=list(state.selected_domain_packs or []),
        selected_industry_packs=list(state.selected_industry_packs or []),
        selected_agent_ids=list(state.selected_agent_ids or []),
        suggested_research_need=bool(state.suggested_research_need),
        next_best_actions=list(state.next_best_actions or []),
        active_task_ids=list(state.active_task_ids or []),
        latest_task_id=state.latest_task_id,
        latest_task_title=state.latest_task_title,
        supplement_count=int(state.supplement_count or 0),
        last_supplement_summary=state.last_supplement_summary or "",
        created_at=state.created_at,
        updated_at=state.updated_at,
    )


def _build_case_world_fact_payloads(
    *,
    client: schemas.ClientRead | None,
    engagement: schemas.EngagementRead | None,
    workstream: schemas.WorkstreamRead | None,
    decision_context: schemas.DecisionContextRead | None,
    domain_lenses: list[str],
    source_materials: list[schemas.SourceMaterialRead],
    artifacts: list[schemas.ArtifactRead],
    usable_evidence_count: int,
    pack_resolution: schemas.PackResolutionRead,
) -> list[dict[str, str]]:
    facts: list[dict[str, str]] = []

    if client and client.name:
        facts.append({"title": "Client", "detail": client.name, "source": "explicit"})
    if engagement and engagement.name:
        facts.append({"title": "Engagement", "detail": engagement.name, "source": "explicit"})
    if workstream and workstream.name:
        facts.append({"title": "Workstream", "detail": workstream.name, "source": "explicit"})
    if decision_context and decision_context.judgment_to_make:
        facts.append(
            {
                "title": "DecisionContext",
                "detail": decision_context.judgment_to_make,
                "source": "explicit",
            }
        )
    if domain_lenses:
        facts.append(
            {
                "title": "DomainLens",
                "detail": join_natural_list(domain_lenses),
                "source": "inferred",
            }
        )
    if source_materials:
        facts.append(
            {
                "title": "SourceMaterial",
                "detail": f"目前已有 {len(source_materials)} 份來源材料。",
                "source": "explicit",
            }
        )
    if artifacts:
        facts.append(
            {
                "title": "Artifact",
                "detail": f"目前已有 {len(artifacts)} 份工作物件。",
                "source": "explicit",
            }
        )
    if usable_evidence_count:
        facts.append(
            {
                "title": "Evidence",
                "detail": f"目前已有 {usable_evidence_count} 則可引用 evidence。",
                "source": "explicit",
            }
        )
    if pack_resolution.selected_domain_packs:
        facts.append(
            {
                "title": "Domain Packs",
                "detail": join_natural_list(
                    [item.pack_name for item in pack_resolution.selected_domain_packs]
                ),
                "source": "pack_resolver",
            }
        )
    if pack_resolution.selected_industry_packs:
        facts.append(
            {
                "title": "Industry Packs",
                "detail": join_natural_list(
                    [item.pack_name for item in pack_resolution.selected_industry_packs]
                ),
                "source": "pack_resolver",
            }
        )

    return facts


def _build_case_world_assumption_payloads(
    *,
    task: models.Task,
    decision_context: schemas.DecisionContextRead | None,
    presence_state_summary: schemas.PresenceStateSummaryRead,
) -> list[dict[str, str]]:
    assumptions = [
        {
            "title": "既定假設",
            "detail": item,
            "source": "user_provided",
        }
        for item in _extract_assumptions(task)
    ]

    if presence_state_summary.decision_context.state == PresenceState.PROVISIONAL:
        assumptions.append(
            {
                "title": "DecisionContext 仍偏 provisional",
                "detail": "目前判斷主軸仍有一部分是 Host 依任務文字先建立的 provisional framing。",
                "source": "host_inference",
            }
        )
    if (
        decision_context
        and decision_context.client_stage
        and decision_context.client_stage == UNSPECIFIED_LABEL
    ):
        assumptions.append(
            {
                "title": "Client stage 尚未明確",
                "detail": "目前案件世界仍缺少更明確的 client stage，因此 stage-specific heuristics 只能先保守套用。",
                "source": "host_inference",
            }
        )

    return assumptions


def _build_case_world_gap_payloads(
    *,
    task: models.Task,
    decision_context: schemas.DecisionContextRead | None,
    input_entry_mode: InputEntryMode,
    source_materials: list[schemas.SourceMaterialRead],
    artifacts: list[schemas.ArtifactRead],
    usable_evidence_count: int,
    pack_resolution: schemas.PackResolutionRead,
    external_research_heavy_candidate: bool,
) -> list[dict[str, object]]:
    gaps: list[dict[str, object]] = []

    if not decision_context or not _normalize_whitespace(decision_context.judgment_to_make):
        gaps.append(
            {
                "gap_key": "decision-context",
                "title": "DecisionContext 仍不夠穩定",
                "description": "這輪仍缺清楚的 judgment_to_make，因此案件世界只能先以 provisional framing 運作。",
                "priority": "high",
                "related_pack_ids": [],
            }
        )

    if not source_materials:
        gaps.append(
            {
                "gap_key": "source-materials",
                "title": "尚缺正式來源材料",
                "description": "目前仍沒有可掛回案件世界的正式來源材料，建議先補檔案、網址或 pasted text。",
                "priority": "high" if input_entry_mode == InputEntryMode.ONE_LINE_INQUIRY else "medium",
                "related_pack_ids": [],
            }
        )

    if usable_evidence_count < 2:
        gaps.append(
            {
                "gap_key": "evidence-thickness",
                "title": "證據鏈仍偏薄",
                "description": "目前 evidence 仍不足以穩定支撐更高等級的 decision / action deliverable。",
                "priority": "high" if not artifacts else "medium",
                "related_pack_ids": [],
            }
        )

    if external_research_heavy_candidate:
        gaps.append(
            {
                "gap_key": "research-needed",
                "title": "需要正式外部補完",
                "description": "這輪屬於 external-research-heavy sparse case，Host 應先補 public research，再回到 evidence chain。",
                "priority": "high",
                "related_pack_ids": [
                    *[item.pack_id for item in pack_resolution.selected_domain_packs],
                    *[item.pack_id for item in pack_resolution.selected_industry_packs],
                ],
            }
        )

    for pack in [*pack_resolution.selected_domain_packs, *pack_resolution.selected_industry_packs]:
        if not pack.evidence_expectations:
            continue
        expectation = pack.evidence_expectations[0]
        gaps.append(
            {
                "gap_key": f"pack-{pack.pack_id}-expectation",
                "title": f"{pack.pack_name} 缺少關鍵支撐",
                "description": f"{pack.pack_name} 目前期待 {expectation} 等材料，但案件世界尚未形成足夠的可引用證據。",
                "priority": "medium",
                "related_pack_ids": [pack.pack_id],
            }
        )

    unique_by_key: dict[str, dict[str, object]] = {}
    for item in gaps:
        gap_key = str(item["gap_key"])
        if gap_key not in unique_by_key:
            unique_by_key[gap_key] = item
    return list(unique_by_key.values())


def _upsert_evidence_gap_rows(
    db: Session,
    *,
    task: models.Task,
    matter_workspace: models.MatterWorkspace | None,
    gap_payloads: list[dict[str, object]],
) -> tuple[list[models.EvidenceGap], bool]:
    existing_by_key = {
        item.gap_key: item for item in task.evidence_gaps if item.source == "case_world_compiler"
    }
    active_keys = {str(item["gap_key"]) for item in gap_payloads}
    rows: list[models.EvidenceGap] = []
    changed = False

    for payload in gap_payloads:
        gap_key = str(payload["gap_key"])
        row = existing_by_key.get(gap_key)
        if row is None:
            row = models.EvidenceGap(
                task_id=task.id,
                matter_workspace_id=matter_workspace.id if matter_workspace else None,
                gap_key=gap_key,
                source="case_world_compiler",
            )
            db.add(row)
            changed = True
        previous_state = (
            row.matter_workspace_id,
            row.title,
            row.gap_type,
            row.description,
            row.priority,
            row.status,
            list(row.supporting_pack_ids or []),
        )
        row.matter_workspace_id = matter_workspace.id if matter_workspace else None
        row.title = str(payload["title"])
        row.gap_type = "evidence_gap"
        row.description = str(payload["description"])
        row.priority = str(payload.get("priority") or "medium")
        row.status = "open"
        row.supporting_pack_ids = list(payload.get("related_pack_ids") or [])
        if previous_state != (
            row.matter_workspace_id,
            row.title,
            row.gap_type,
            row.description,
            row.priority,
            row.status,
            list(row.supporting_pack_ids or []),
        ):
            changed = True
        rows.append(row)

    for gap_key, row in existing_by_key.items():
        if gap_key in active_keys:
            continue
        if row.status != "resolved":
            row.status = "resolved"
            changed = True
        rows.append(row)

    db.flush()
    ordered = sorted(
        {item.id: item for item in rows}.values(),
        key=lambda item: item.updated_at.isoformat() if item.updated_at else "",
        reverse=True,
    )
    return ordered, changed


def ensure_case_world_draft_for_task(
    db: Session,
    *,
    task: models.Task,
    matter_workspace: models.MatterWorkspace | None,
    client: schemas.ClientRead | None,
    engagement: schemas.EngagementRead | None,
    workstream: schemas.WorkstreamRead | None,
    decision_context: schemas.DecisionContextRead | None,
    domain_lenses: list[str],
    input_entry_mode: InputEntryMode,
    deliverable_class_hint: DeliverableClass,
    external_research_heavy_candidate: bool,
    presence_state_summary: schemas.PresenceStateSummaryRead,
    pack_resolution: schemas.PackResolutionRead,
    agent_selection: schemas.AgentSelectionRead,
    source_materials: list[schemas.SourceMaterialRead],
    artifacts: list[schemas.ArtifactRead],
) -> tuple[models.CaseWorldDraft, list[models.EvidenceGap], bool]:
    continuity_mode, writeback_depth = resolve_continuity_policy_for_task(task, matter_workspace)
    usable_evidence_count = len(_usable_evidence(task))
    capability = _infer_capability_for_task(task, decision_context, domain_lenses)
    fact_payloads = _build_case_world_fact_payloads(
        client=client,
        engagement=engagement,
        workstream=workstream,
        decision_context=decision_context,
        domain_lenses=domain_lenses,
        source_materials=source_materials,
        artifacts=artifacts,
        usable_evidence_count=usable_evidence_count,
        pack_resolution=pack_resolution,
    )
    assumption_payloads = _build_case_world_assumption_payloads(
        task=task,
        decision_context=decision_context,
        presence_state_summary=presence_state_summary,
    )
    gap_payloads = _build_case_world_gap_payloads(
        task=task,
        decision_context=decision_context,
        input_entry_mode=input_entry_mode,
        source_materials=source_materials,
        artifacts=artifacts,
        usable_evidence_count=usable_evidence_count,
        pack_resolution=pack_resolution,
        external_research_heavy_candidate=external_research_heavy_candidate,
    )
    evidence_gap_rows, evidence_gap_changed = _upsert_evidence_gap_rows(
        db,
        task=task,
        matter_workspace=matter_workspace,
        gap_payloads=gap_payloads,
    )

    extracted_objects = [
        {"object_type": "client", "object_id": client.id if client else None, "label": client.name if client else "未指定 client"},
        {"object_type": "engagement", "object_id": engagement.id if engagement else None, "label": engagement.name if engagement else "未指定 engagement"},
        {"object_type": "workstream", "object_id": workstream.id if workstream else None, "label": workstream.name if workstream else "未指定 workstream"},
        {
            "object_type": "decision_context",
            "object_id": decision_context.id if decision_context else None,
            "label": decision_context.title if decision_context else task.title,
        },
    ]
    inferred_links = [
        {"from": "task", "to": "decision_context", "relation": "frames"},
        {"from": "decision_context", "to": "source_material", "relation": "expects_evidence"},
        {"from": "source_material", "to": "evidence", "relation": "produces"},
        {"from": "deliverable", "to": "decision_record", "relation": "writes_back"},
    ]
    next_best_actions: list[str] = []
    if not source_materials:
        next_best_actions.append("先補至少一份檔案、網址或 pasted text。")
    if external_research_heavy_candidate:
        next_best_actions.append("允許 Host 先補正式外部 research，再回到 evidence chain。")
    if usable_evidence_count < 2:
        next_best_actions.append("先執行第一版分析，或補更多 evidence 之後再收斂。")
    else:
        next_best_actions.append("直接執行分析並寫回正式交付物。")
    if continuity_mode != EngagementContinuityMode.ONE_OFF:
        next_best_actions.append("把這輪結果寫回案件世界，保留 decision checkpoint。")

    existing = task.case_world_drafts[0] if task.case_world_drafts else None
    changed = existing is None
    if existing is None:
        existing = models.CaseWorldDraft(
            task_id=task.id,
            matter_workspace_id=matter_workspace.id if matter_workspace else None,
        )
        db.add(existing)

    canonical_intake_summary = {
        "entry_preset": task.entry_preset,
        "resolved_input_state": input_entry_mode.value,
        "problem_statement": (
            decision_context.judgment_to_make
            if decision_context and decision_context.judgment_to_make
            else task.description or task.title
        ),
        "source_material_count": len(source_materials),
        "artifact_count": len(artifacts),
        "usable_evidence_count": usable_evidence_count,
        "external_research_heavy_candidate": external_research_heavy_candidate,
    }
    task_interpretation = {
        "task_title": task.title,
        "task_type": task.task_type,
        "workflow_mode": task.mode,
        "capability": capability.value,
        "deliverable_class_hint": deliverable_class_hint.value,
        "external_data_strategy": get_external_data_strategy_for_task(task).value,
        "summary": _build_sparse_input_summary(
            input_entry_mode,
            deliverable_class_hint,
            external_research_heavy_candidate,
        ),
    }
    canonical_decision_context_payload = None
    canonical_decision_context_row = None
    if matter_workspace and matter_workspace.case_world_state:
        state_payload = matter_workspace.case_world_state.decision_context_payload
        if isinstance(state_payload, dict):
            canonical_decision_context_payload = state_payload
        _, _, _, canonical_decision_context_row = _load_world_authority_spine(
            db,
            matter_workspace=matter_workspace,
            case_world_state=matter_workspace.case_world_state,
        )
    if canonical_decision_context_payload is None and existing is not None and isinstance(existing.decision_context_payload, dict):
        canonical_decision_context_payload = existing.decision_context_payload
    decision_context_payload = _merge_world_decision_context_payload(
        canonical_payload=canonical_decision_context_payload,
        canonical_row=canonical_decision_context_row,
        overlay=decision_context,
        task=task,
    )
    suggested_domain_packs = [item.pack_id for item in pack_resolution.selected_domain_packs]
    suggested_industry_packs = [item.pack_id for item in pack_resolution.selected_industry_packs]
    suggested_research_need = external_research_heavy_candidate or any(
        item["gap_key"] == "research-needed" for item in gap_payloads
    )

    def assign(attribute: str, value: object) -> None:
        nonlocal changed
        if getattr(existing, attribute) != value:
            setattr(existing, attribute, value)
            changed = True

    assign("matter_workspace_id", matter_workspace.id if matter_workspace else None)
    assign("compiler_status", "compiled")
    assign("entry_preset", task.entry_preset)
    assign("continuity_mode", continuity_mode.value)
    assign("writeback_depth", writeback_depth.value)
    assign("canonical_intake_summary", canonical_intake_summary)
    assign("task_interpretation", task_interpretation)
    assign("decision_context_payload", decision_context_payload)
    assign("extracted_objects", extracted_objects)
    assign("inferred_links", inferred_links)
    assign("facts", fact_payloads)
    assign("assumptions_payload", assumption_payloads)
    assign("evidence_gaps_payload", gap_payloads)
    assign("suggested_capabilities", [capability.value])
    assign("suggested_domain_packs", suggested_domain_packs)
    assign("suggested_industry_packs", suggested_industry_packs)
    assign("suggested_agents", agent_selection.selected_agent_ids)
    assign("suggested_research_need", suggested_research_need)
    assign("next_best_actions", _unique_preserve_order(next_best_actions))
    db.flush()
    return existing, evidence_gap_rows, changed or evidence_gap_changed


def serialize_task(task: models.Task) -> schemas.TaskAggregateResponse:
    db = object_session(task)
    if db is None:
        raise RuntimeError("Task must be attached to an active session before serialization.")

    task_client_row = _primary_item(task.clients)
    task_engagement_row = _primary_item(task.engagements)
    task_workstream_row = _primary_item(task.workstreams)
    task_decision_context_row = _primary_item(task.decision_contexts)
    client, engagement, workstream, decision_context, domain_lenses, source_materials, artifacts = (
        _build_ontology_spine_for_task(task)
    )
    matter_workspace, workspace_changed = ensure_matter_workspace_for_task(
        db=db,
        task=task,
        client=client,
        engagement=engagement,
        workstream=workstream,
        decision_context=decision_context,
        domain_lenses=domain_lenses,
    )
    if workspace_changed:
        db.commit()
    case_world_state_row, world_spine_changed = ensure_world_context_spine_for_task(
        db,
        matter_workspace=matter_workspace,
        task=task,
        client=task_client_row,
        engagement=task_engagement_row,
        workstream=task_workstream_row,
        decision_context=task_decision_context_row,
        case_world_state=matter_workspace.case_world_state,
    )
    _, _, _, world_authority_decision_context_row = (
        _load_world_authority_spine(
            db,
            matter_workspace=matter_workspace,
            case_world_state=case_world_state_row,
        )
    )
    world_client, world_engagement, world_workstream, preferred_decision_context, preferred_domain_lenses = (
        _build_world_preferred_context_reads(
            task,
            client=client,
            engagement=engagement,
            workstream=workstream,
            decision_context=decision_context,
            domain_lenses=domain_lenses,
        )
    )
    world_decision_context = _build_world_decision_context_read(
        case_world_state=case_world_state_row,
        authority_decision_context=world_authority_decision_context_row,
        fallback=preferred_decision_context,
    )
    preferred_decision_context = world_decision_context or preferred_decision_context
    slice_decision_context = _build_slice_decision_context_overlay(
        overlay=decision_context,
        world=preferred_decision_context,
    )
    matter_workspace_summary = _build_matter_workspace_summary_map(db, [matter_workspace]).get(
        matter_workspace.id
    )
    evidence = _serialize_evidence_items(task, source_materials, artifacts)
    recommendation_support_map, risk_support_map, action_item_support_map = _build_supporting_evidence_maps(
        task
    )
    recommendations = _serialize_recommendations(task, recommendation_support_map)
    risks = _serialize_risks(task, risk_support_map)
    action_items = _serialize_action_items(task, action_item_support_map)
    deliverables = _serialize_deliverables(
        task,
        _build_object_label_map(
            task,
            source_materials,
            artifacts,
            evidence,
            recommendations,
            risks,
            action_items,
            preferred_decision_context,
            world_client,
            world_engagement,
            world_workstream,
        ),
    )
    input_entry_mode = _infer_input_entry_mode(task, source_materials, artifacts)
    external_research_heavy_candidate = _is_external_research_heavy_candidate(
        task,
        preferred_decision_context,
        source_materials,
        artifacts,
        input_entry_mode,
    )
    presence_state_summary = _build_presence_state_summary(
        task,
        world_client,
        world_engagement,
        world_workstream,
        preferred_decision_context,
        preferred_domain_lenses,
        source_materials,
        artifacts,
        input_entry_mode,
        external_research_heavy_candidate,
    )
    pack_resolution = resolve_pack_selection_for_task(
        task,
        world_client,
        world_engagement,
        world_workstream,
        preferred_decision_context,
        preferred_domain_lenses,
    )
    deliverable_class_hint = _resolve_deliverable_class_hint(
        input_entry_mode,
        preferred_decision_context,
        source_materials,
        artifacts,
        len(_usable_evidence(task)),
        external_research_heavy_candidate,
    )
    agent_selection = resolve_agent_selection_for_task(
        task,
        preferred_decision_context,
        preferred_domain_lenses,
        pack_resolution,
        input_entry_mode,
        deliverable_class_hint,
        external_research_heavy_candidate,
        source_materials,
        artifacts,
    )
    continuity_mode, writeback_depth = resolve_continuity_policy_for_task(task, matter_workspace)
    case_world_draft_row, evidence_gap_rows, case_world_changed = ensure_case_world_draft_for_task(
        db,
        task=task,
        matter_workspace=matter_workspace,
        client=world_client,
        engagement=world_engagement,
        workstream=world_workstream,
        decision_context=preferred_decision_context,
        domain_lenses=preferred_domain_lenses,
        input_entry_mode=input_entry_mode,
        deliverable_class_hint=deliverable_class_hint,
        external_research_heavy_candidate=external_research_heavy_candidate,
        presence_state_summary=presence_state_summary,
        pack_resolution=pack_resolution,
        agent_selection=agent_selection,
        source_materials=source_materials,
        artifacts=artifacts,
    )
    case_world_state_row, case_world_state_changed = ensure_case_world_state_for_task(
        db,
        matter_workspace=matter_workspace,
        task=task,
        case_world_draft=case_world_draft_row,
    )
    if world_spine_changed or case_world_changed or case_world_state_changed:
        db.commit()
        task = get_loaded_task(db, task.id)
        client, engagement, workstream, decision_context, domain_lenses, source_materials, artifacts = (
            _build_ontology_spine_for_task(task)
        )
        matter_workspace = _get_linked_matter_workspace(task) or matter_workspace
        matter_workspace_summary = _build_matter_workspace_summary_map(db, [matter_workspace]).get(
            matter_workspace.id
        )
        evidence = _serialize_evidence_items(task, source_materials, artifacts)
        recommendation_support_map, risk_support_map, action_item_support_map = _build_supporting_evidence_maps(
            task
        )
        recommendations = _serialize_recommendations(task, recommendation_support_map)
        risks = _serialize_risks(task, risk_support_map)
        action_items = _serialize_action_items(task, action_item_support_map)
        case_world_draft_row = task.case_world_drafts[0] if task.case_world_drafts else case_world_draft_row
        case_world_state_row = matter_workspace.case_world_state
        evidence_gap_rows = list(task.evidence_gaps)
        _, _, _, world_authority_decision_context_row = (
            _load_world_authority_spine(
                db,
                matter_workspace=matter_workspace,
                case_world_state=case_world_state_row,
            )
        )
        world_client, world_engagement, world_workstream, preferred_decision_context, preferred_domain_lenses = (
            _build_world_preferred_context_reads(
                task,
                client=client,
                engagement=engagement,
                workstream=workstream,
                decision_context=decision_context,
                domain_lenses=domain_lenses,
            )
        )
        world_decision_context = _build_world_decision_context_read(
            case_world_state=case_world_state_row,
            authority_decision_context=world_authority_decision_context_row,
            fallback=preferred_decision_context,
        )
        preferred_decision_context = world_decision_context or preferred_decision_context
        slice_decision_context = _build_slice_decision_context_overlay(
            overlay=decision_context,
            world=preferred_decision_context,
        )

    return schemas.TaskAggregateResponse(
        id=task.id,
        title=task.title,
        description=task.description,
        task_type=task.task_type,
        mode=task.mode,
        external_data_strategy=get_external_data_strategy_for_task(task),
        status=task.status,
        created_at=task.created_at,
        updated_at=task.updated_at,
        client=world_client,
        engagement=world_engagement,
        workstream=world_workstream,
        decision_context=preferred_decision_context,
        slice_decision_context=slice_decision_context,
        world_decision_context=world_decision_context,
        client_stage=(
            world_decision_context.client_stage
            if world_decision_context and world_decision_context.client_stage
            else decision_context.client_stage if decision_context else client.client_stage if client else None
        ),
        client_type=(
            world_decision_context.client_type
            if world_decision_context and world_decision_context.client_type
            else decision_context.client_type if decision_context else client.client_type if client else None
        ),
        domain_lenses=domain_lenses,
        assumptions=(
            preferred_decision_context.assumptions
            if preferred_decision_context and preferred_decision_context.assumptions
            else _extract_assumptions(task)
        ),
        entry_preset=InputEntryMode(task.entry_preset),
        input_entry_mode=input_entry_mode,
        engagement_continuity_mode=continuity_mode,
        writeback_depth=writeback_depth,
        deliverable_class_hint=deliverable_class_hint,
        external_research_heavy_candidate=external_research_heavy_candidate,
        sparse_input_summary=_build_sparse_input_summary(
            input_entry_mode,
            deliverable_class_hint,
            external_research_heavy_candidate,
        ),
        presence_state_summary=presence_state_summary,
        pack_resolution=pack_resolution,
        agent_selection=agent_selection,
        source_materials=source_materials,
        artifacts=artifacts,
        contexts=[schemas.TaskContextRead.model_validate(item) for item in task.contexts],
        subjects=[schemas.SubjectRead.model_validate(item) for item in task.subjects],
        goals=[schemas.GoalRead.model_validate(item) for item in task.goals],
        constraints=[schemas.ConstraintRead.model_validate(item) for item in task.constraints],
        uploads=_build_source_documents(task),
        evidence=evidence,
        insights=[schemas.InsightRead.model_validate(item) for item in task.insights],
        risks=risks,
        options=[schemas.OptionRead.model_validate(item) for item in task.options],
        recommendations=recommendations,
        action_items=action_items,
        deliverables=deliverables,
        runs=[schemas.TaskRunRead.model_validate(item) for item in task.runs],
        case_world_draft=_serialize_case_world_draft(case_world_draft_row),
        case_world_state=_serialize_case_world_state(case_world_state_row),
        world_work_slice_summary=(
            f"這筆工作是案件世界「{matter_workspace.title}」中的 work slice；"
            f"目前同一世界內共有 {len(case_world_state_row.active_task_ids)} 個 task slices。"
            if matter_workspace and case_world_state_row
            else "目前這筆工作仍未掛回正式案件世界。"
        ),
        evidence_gaps=[
            schemas.EvidenceGapRead.model_validate(item)
            for item in sorted(evidence_gap_rows, key=lambda gap: gap.updated_at, reverse=True)
        ],
        research_runs=[schemas.ResearchRunRead.model_validate(item) for item in task.research_runs],
        decision_records=[schemas.DecisionRecordRead.model_validate(item) for item in task.decision_records],
        action_plans=[schemas.ActionPlanRead.model_validate(item) for item in task.action_plans],
        action_executions=[
            schemas.ActionExecutionRead.model_validate(item) for item in task.action_executions
        ],
        outcome_records=[schemas.OutcomeRecordRead.model_validate(item) for item in task.outcome_records],
        matter_workspace=matter_workspace_summary,
    )


def get_task_history(db: Session, task_id: str) -> schemas.TaskHistoryResponse:
    task = get_loaded_task(db, task_id)
    client, engagement, workstream, decision_context, _, source_materials, artifacts = _build_world_preferred_ontology_spine_for_task(
        task
    )
    evidence = _serialize_evidence_items(task, source_materials, artifacts)
    recommendation_support_map, risk_support_map, action_item_support_map = _build_supporting_evidence_maps(
        task
    )
    recommendations = _serialize_recommendations(task, recommendation_support_map)
    action_items = _serialize_action_items(task, action_item_support_map)
    risks = _serialize_risks(task, risk_support_map)
    deliverables = _serialize_deliverables(
        task,
        _build_object_label_map(
            task,
            source_materials,
            artifacts,
            evidence,
            recommendations,
            risks,
            action_items,
            decision_context,
            client,
            engagement,
            workstream,
        ),
    )
    return schemas.TaskHistoryResponse(
        task_id=task.id,
        runs=[schemas.TaskRunRead.model_validate(item) for item in task.runs],
        deliverables=deliverables,
        recommendations=recommendations,
        action_items=action_items,
        evidence=evidence,
    )
